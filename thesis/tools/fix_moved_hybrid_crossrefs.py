from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
doc = Document(path)

replacements = {
    "In Equation (13),": (
        "In Equation (12), g(·) denotes the learned fusion MLP and square brackets denote vector "
        "concatenation. A linear head maps the fused representation to class scores. On CIFAR-10, "
        "mean, mean-MLP and concatenation-MLP fusion contain 1,600,778, 1,732,746 and 1,798,538 "
        "trainable parameters, compared with 801,034 for one fixed-PE encoder. Most of the increase "
        "comes from duplicating the encoder. The MLP variants add further parameters in the fusion stage."
    ),
    "In Equation (14),": (
        "In Equation (13), W_Q, W_K and W_V are learned query, key and value projections, and "
        "d_h = D/4 = 32 is the width of each attention head. Four heads perform this exchange in "
        "parallel. Their outputs are concatenated and projected back to the model dimension. One block "
        "updates row tokens from column context and another updates column tokens from row context. Both "
        "use pre-normalisation, residual connections and an MLP. The updated classification tokens are "
        "concatenated and passed to either a linear or MLP head. These variants contain 1,999,114 and "
        "2,031,242 parameters, so accuracy is interpreted alongside model size."
    ),
}

for paragraph in doc.paragraphs:
    for prefix, new_text in replacements.items():
        if paragraph.text.startswith(prefix):
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            run = paragraph.add_run(new_text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{attr}"), "Arial")

tmp = path.with_suffix(".crossrefs.tmp.docx")
doc.save(tmp)
tmp.replace(path)
