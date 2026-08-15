from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DOCX_PATH = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
BACKUP_PATH = DOCX_PATH.with_name("Yikai_Zhao_MSc_Dissertation.pre_references_backup.docx")


REFERENCES = [
    (
        "[1] A. Dosovitskiy et al., “An Image Is Worth 16 × 16 Words: Transformers for Image "
        "Recognition at Scale,” in International Conference on Learning Representations, 2021. "
        "[Online]. Available: https://openreview.net/forum?id=YicbFdNTTy"
    ),
    (
        "[2] A. Vaswani et al., “Attention Is All You Need,” in Advances in Neural Information "
        "Processing Systems, vol. 30, 2017. [Online]. Available: "
        "https://papers.nips.cc/paper/7181-attention-is-all-you-need"
    ),
    (
        "[3] L. Yuan et al., “Tokens-to-Token ViT: Training Vision Transformers From Scratch on "
        "ImageNet,” in Proceedings of the IEEE/CVF International Conference on Computer Vision, "
        "2021, pp. 558–567."
    ),
    (
        "[4] F. Wang, Y. Yu, G. Wei, W. Shao, Y. Zhou, A. Yuille, and C. Xie, “Scaling Laws in "
        "Patchification: An Image Is Worth 50,176 Tokens And More,” arXiv:2502.03738, 2025, "
        "doi: 10.48550/arXiv.2502.03738."
    ),
    (
        "[5] P. Dufter, M. Schmitt, and H. Schütze, “Position Information in Transformers: An "
        "Overview,” Computational Linguistics, vol. 48, no. 3, pp. 733–763, 2022, "
        "doi: 10.1162/coli_a_00445."
    ),
    (
        "[6] M. A. M. Chowdhury, M. R. U. Rahman, and A. A. Taki, “LOOPE: Learnable Optimal "
        "Patch Order in Positional Embeddings for Vision Transformers,” arXiv:2504.14386, 2025, "
        "doi: 10.48550/arXiv.2504.14386."
    ),
    (
        "[7] D. Kutscher, D. M. Chan, Y. Bai, T. Darrell, and R. Gupta, “REOrdering Patches "
        "Improves Vision Models,” in Advances in Neural Information Processing Systems, vol. 38, "
        "2025. [Online]. Available: https://papers.nips.cc/paper_files/paper/2025/hash/"
        "a751eb81265da168df0e765b5bc874c9-Abstract-Conference.html"
    ),
    (
        "[8] H. Touvron, M. Cord, M. Douze, F. Massa, A. Sablayrolles, and H. Jégou, “Training "
        "Data-Efficient Image Transformers & Distillation through Attention,” in Proceedings of "
        "the 38th International Conference on Machine Learning, vol. 139, 2021, pp. 10347–10357."
    ),
    (
        "[9] A. Hassani, S. Walton, N. Shah, A. Abuduweili, J. Li, and H. Shi, “Escaping the Big "
        "Data Paradigm with Compact Transformers,” arXiv:2104.05704, 2021, "
        "doi: 10.48550/arXiv.2104.05704."
    ),
    (
        "[10] K. Wu, H. Peng, M. Chen, J. Fu, and H. Chao, “Rethinking and Improving Relative "
        "Position Encoding for Vision Transformer,” in Proceedings of the IEEE/CVF International "
        "Conference on Computer Vision, 2021, pp. 10033–10041."
    ),
    (
        "[11] X. Chu, Z. Tian, B. Zhang, X. Wang, and C. Shen, “Conditional Positional Encodings "
        "for Vision Transformers,” in International Conference on Learning Representations, 2023. "
        "[Online]. Available: https://openreview.net/forum?id=3KWnuT-R1bh"
    ),
    (
        "[12] B. Heo, S. Park, D. Han, and S. Yun, “Rotary Position Embedding for Vision "
        "Transformer,” in Computer Vision – ECCV 2024, vol. 15068, 2024, pp. 289–305, "
        "doi: 10.1007/978-3-031-72684-2_17."
    ),
    (
        "[13] A. Krizhevsky, “Learning Multiple Layers of Features from Tiny Images,” University "
        "of Toronto, Toronto, ON, Canada, Tech. Rep., 2009. [Online]. Available: "
        "https://www.cs.toronto.edu/~kriz/learning-features-2009-TR.pdf"
    ),
    (
        "[14] I. Loshchilov and F. Hutter, “Decoupled Weight Decay Regularization,” in "
        "International Conference on Learning Representations, 2019. [Online]. Available: "
        "https://openreview.net/forum?id=Bkg6RiCqY7"
    ),
]


def set_run_arial_black(run, size: Pt | None = None) -> None:
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = size
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")


def ensure_reference_style(document: Document):
    styles = document.styles
    if "Dissertation Reference" in styles:
        style = styles["Dissertation Reference"]
    else:
        style = styles.add_style("Dissertation Reference", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]

    style.font.name = "Arial"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")

    fmt = style.paragraph_format
    fmt.left_indent = Cm(0.75)
    fmt.first_line_indent = Cm(-0.75)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.keep_with_next = False
    fmt.keep_together = False
    fmt.widow_control = True
    return style


def remove_existing_reference_section(document: Document) -> None:
    start = None
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "References":
            start = i
            break
    if start is None:
        return
    for paragraph in list(document.paragraphs[start:]):
        paragraph._element.getparent().remove(paragraph._element)


def replace_patchification_claim(document: Document) -> bool:
    old = (
        "A recent preprint on patchification scaling indicates that the benefits of increasing "
        "token count depend on the model and compute regime rather than following a single "
        "universal rule [4]."
    )
    new = (
        "A recent preprint on patchification scaling reports that smaller patches can improve "
        "predictive performance, although they also increase sequence length and computation [4]."
    )
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            for run in paragraph.runs:
                set_run_arial_black(run, Pt(11))
            return True
    return False


def add_reference_section(document: Document) -> None:
    reference_style = ensure_reference_style(document)
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.keep_with_next = True
    run = heading.add_run("References")
    set_run_arial_black(run, Pt(16))

    for entry in REFERENCES:
        paragraph = document.add_paragraph(style=reference_style)
        run = paragraph.add_run(entry)
        set_run_arial_black(run, Pt(10))


def set_update_fields(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def audit(document: Document) -> None:
    body_text = "\n".join(p.text for p in document.paragraphs if p.text.strip() != "References")
    cited = set()
    for match in re.finditer(r"\[(\d+)(?:[\]–-]|$)", body_text):
        cited.add(int(match.group(1)))
    for match in re.finditer(r"\[(\d+)\]", body_text):
        cited.add(int(match.group(1)))
    expected = set(range(1, len(REFERENCES) + 1))
    missing = sorted(expected - cited)
    extra = sorted(cited - expected)
    if missing or extra:
        raise RuntimeError(f"Citation audit failed. Missing in text: {missing}; extra in text: {extra}")

    headings = [p.text.strip() for p in document.paragraphs if p.style.name.startswith("Heading")]
    if headings[-1] != "References":
        raise RuntimeError("References is not the final heading.")

    tail = [p.text.strip() for p in document.paragraphs[-len(REFERENCES):]]
    for index, entry in enumerate(tail, start=1):
        if not entry.startswith(f"[{index}]"):
            raise RuntimeError(f"Reference numbering failed at [{index}].")


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    document = Document(DOCX_PATH)
    remove_existing_reference_section(document)
    changed_claim = replace_patchification_claim(document)
    add_reference_section(document)
    set_update_fields(document)
    audit(document)
    document.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")
    print(f"Backup: {BACKUP_PATH}")
    print(f"References added: {len(REFERENCES)}")
    print(f"Patchification claim aligned: {changed_claim}")


if __name__ == "__main__":
    main()
