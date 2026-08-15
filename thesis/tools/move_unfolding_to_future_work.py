from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")


def set_run_font(run, size=11):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Arial")


def replace_paragraph(paragraph, text):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    set_run_font(run)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def replace_prefix(doc, prefix, text):
    replace_paragraph(find_paragraph(doc, prefix), text)


def body_text(element):
    return "".join(element.xpath(".//w:t/text() | .//m:t/text()"))


def delete_section(doc, start_heading, next_heading):
    body = doc.element.body
    children = list(body.iterchildren())
    start = next(i for i, child in enumerate(children) if body_text(child).strip() == start_heading)
    end = next(i for i, child in enumerate(children) if i > start and body_text(child).strip() == next_heading)
    for child in children[start:end]:
        body.remove(child)


def equation_paragraph(doc, marker):
    for paragraph in doc.paragraphs:
        text = "".join(paragraph._p.xpath(".//m:t/text()"))
        if marker in text:
            return paragraph
    raise ValueError(f"Equation not found: {marker}")


def change_equation_number(paragraph, old, new):
    for node in reversed(paragraph._p.xpath(".//m:t")):
        if node.text == str(old):
            node.text = str(new)
            return
    raise ValueError(f"Equation number {old} not found")


doc = Document(DOCX)

# Remove the unfolding experiment from the main literature, method and results.
delete_section(doc, "2.4 Patch Ordering and Patch-to-Position Assignment", "2.5 Data-Efficient and Data-Limited ViT Training")
delete_section(doc, "3.5 Patch Ordering and Position Assignment", "3.6 Dual-Branch Fusion Extensions")
delete_section(doc, "4.4 Patch Ordering and Position Assignment", "4.5 Performance under Limited Training Data")

# Front matter and Introduction.
replace_prefix(
    doc,
    "Fixed and learnable 2D encodings,",
    "Fixed and learnable 2D encodings, data availability, and model extensions on CIFAR-10 and CIFAR-100",
)
replace_prefix(
    doc,
    "The original ViT uses a learned vector",
    "The original ViT uses a learned vector for each token position [1]. These vectors can change during training and adapt to the task. However, they do not directly use the two-dimensional row and column structure of the image. A fixed two-dimensional encoding instead creates each positional vector from the row and column of the patch. It does not add trainable positional parameters and directly encodes the row and column position of each patch. Unlike learned PE, however, its values remain fixed during training. This raises the question of whether learned and fixed PE behave differently as the amount of available training data changes.",
)
replace_prefix(
    doc,
    "Many previous studies compare positional encoding methods",
    "Many previous studies compare positional encoding methods using different model architectures or training settings [5], [10]–[12]. This makes it difficult to know whether a change in performance comes from the positional encoding itself or from other differences in the model. This project therefore keeps the patch size, ViT backbone, and training process the same. It studies whether positional encoding improves the model, which fixed PE methods perform best, and whether the comparison changes with the amount of training data. Dual-branch models are tested separately because they change the architecture and number of parameters.",
)
replace_prefix(
    doc,
    "The study is organised around three research questions",
    "The study is organised around three research questions. RQ1 asks whether positional encoding is necessary for the evaluated ViT and how learned absolute PE compares with fixed alternatives. RQ2 examines which fixed constructions, including axis-wise, additive, multiplicative, shifted, squared and radial designs, are most competitive. RQ3 considers how the amount of training data affects the comparison between learned and fixed PE, and whether the main test-accuracy ordering is repeated on CIFAR-100.",
)
replace_prefix(
    doc,
    "Chapter 2 reviews previous work",
    "Chapter 2 reviews previous work on positional information in Transformers and explains the research gap addressed by this project. Chapter 3 describes the datasets, model architecture, positional encoding methods, fusion extensions, and evaluation process. Chapter 4 presents the experiments and results. Chapter 5 discusses the findings, their limitations, and how they answer the research questions. Chapter 6 summarises the main conclusions and suggests possible directions for future work.",
)

# Renumber the remaining Literature Review sections and narrow the gap statement.
replace_paragraph(find_paragraph(doc, "2.5 Data-Efficient and Data-Limited ViT Training"), "2.4 Data-Efficient and Data-Limited ViT Training")
replace_paragraph(find_paragraph(doc, "2.6 Related-Work Synthesis and Gap"), "2.5 Related-Work Synthesis and Gap")
replace_prefix(
    doc,
    "The practical gap is therefore one of attribution",
    "The practical gap is therefore one of attribution, not a lack of proposed encodings. Results from different studies may reflect changes in architecture, scale, data, augmentation or optimisation as well as the position method. This project instead uses one ViT and a common protocol to separate three factors: PE presence and form, limited training data, and the cost of multi-branch complexity. This controlled comparison does not claim a universally superior encoding or a new general-purpose Transformer.",
)

# Methodology overview and section numbering.
replace_prefix(
    doc,
    "This chapter explains how the experiments were constructed",
    "This chapter explains how the experiments were constructed so that each reported comparison can be reproduced and interpreted. Every image passes through the same broad pipeline: dataset-specific preprocessing, division into patches, assignment of positional information, and processing by a shared Transformer encoder. Validation performance selects the checkpoint. The official test set is used only after selection. The following sections define the controlled design, data, shared architecture, positional encodings, fusion extensions, training protocol and statistical analysis.",
)
replace_prefix(
    doc,
    "The study changes one main factor at a time",
    "The study changes one main factor at a time rather than introducing a new general-purpose Transformer. Within each comparison, the ViT backbone, data split, optimiser and checkpoint-selection rule remain fixed. The factors that do change are the positional encoding and, in the architecture extension, the use of dual-branch fusion components. This design reduces the number of competing explanations for an observed accuracy difference.",
)
replace_prefix(
    doc,
    "The main evaluation uses CIFAR-10 and five training seeds",
    "The main evaluation uses CIFAR-10 and five training seeds. No PE provides a genuine no-position baseline, while learnable absolute PE provides the standard learned reference. Fixed row-column encodings then test different ways to represent the two-dimensional grid without trainable positional parameters. Fusion models are analysed separately because they change more than the core positional vector. CIFAR-100 tests whether the observed ranking is retained on a finer-grained classification problem, and reduced-data CIFAR-10 tests whether the relationship changes when fewer labelled images are available.",
)
replace_prefix(
    doc,
    "Equation (9) combines the axis-specific mappings",
    "Equation (9) combines the axis-specific mappings using the same addition and multiplication rules as before. The superscripts label the frequency schedule used for each axis. They are not powers. The term shifted therefore refers only to the one-index offset between the two schedules. It does not move a patch or change its row or column.",
)
replace_prefix(
    doc,
    "The core comparison is deliberately limited",
    "The core comparison is deliberately limited to methods that produce a 128-component positional term and add it at the same point in the shared ViT. Other approaches reviewed in Sections 2.3 and 2.5, including relative-position, conditional and rotary encodings [10]–[12], either modify the attention calculation or introduce learned position-generating components. Including them would change more than the positional vector itself. The dual-branch models also alter parameter count and architecture, so they are evaluated separately in Section 3.5.",
)
replace_paragraph(find_paragraph(doc, "3.6 Dual-Branch Fusion Extensions"), "3.5 Dual-Branch Fusion Extensions")
replace_paragraph(find_paragraph(doc, "3.6.1 Latent Fusion"), "3.5.1 Latent Fusion")
replace_paragraph(find_paragraph(doc, "3.6.2 Bidirectional Cross-Attention Fusion"), "3.5.2 Bidirectional Cross-Attention Fusion")
replace_paragraph(find_paragraph(doc, "3.7 Training and Checkpoint-Selection Protocol"), "3.6 Training and Checkpoint-Selection Protocol")
replace_paragraph(find_paragraph(doc, "3.8 Evaluation Metrics and Statistical Analysis"), "3.7 Evaluation Metrics and Statistical Analysis")
replace_paragraph(find_paragraph(doc, "3.9 Implementation and Reproducibility"), "3.8 Implementation and Reproducibility")
replace_prefix(
    doc,
    "The reduced-data study draws subsets",
    "The reduced-data study draws subsets of 1,000, 5,000 and 10,000 images from the fixed 45,000-image CIFAR-10 training pool. It compares no PE, learnable absolute PE and shifted multiplicative PE using the common learning rate in Section 3.6. The full-data results use the same optimisation settings and are included as a fourth data-size condition. Validation and test sets do not change. Within each seed, all three models receive the same subset. The subset changes across seeds, so the reported variation includes both stochastic training and the sampled training images.",
)

# Equation 11 was removed with the unfolding method; close the numbering gap.
latent_eq = equation_paragraph(doc, "hmean=")
cross_eq = equation_paragraph(doc, "CA(Q,C)=")
change_equation_number(latent_eq, 12, 11)
change_equation_number(cross_eq, 13, 12)
replace_prefix(
    doc,
    "Latent fusion uses two ViT encoders",
    "Latent fusion uses two ViT encoders on the same image: one with row PE and one with column PE. Their classification-token vectors are h_r and h_c. One variant averages them, one applies an MLP after averaging, and one concatenates them before a layer-normalised MLP. Equation (11) summarises the basic rules.",
)
replace_prefix(
    doc,
    "In Equation (12), g(·)",
    "In Equation (11), g(·) denotes the learned fusion MLP and square brackets denote vector concatenation. A linear head maps the fused representation to class scores. On CIFAR-10, mean, mean-MLP and concatenation-MLP fusion contain 1,600,778, 1,732,746 and 1,798,538 trainable parameters, compared with 801,034 for one fixed-PE encoder. Most of the increase comes from duplicating the encoder. The MLP variants add further parameters in the fusion stage.",
)
replace_prefix(
    doc,
    "Cross-attention allows one branch",
    "Cross-attention allows one branch to request information from the other before classification. The models therefore retain the complete row and column token sequences rather than only their final classification tokens. For one attention head, Q is the sequence being updated and C is the other branch used as context. Equation (12) gives the operation.",
)
replace_prefix(
    doc,
    "In Equation (13), W_Q",
    "In Equation (12), W_Q, W_K and W_V are learned query, key and value projections, and d_h = D/4 = 32 is the width of each attention head. Four heads perform this exchange in parallel. Their outputs are concatenated and projected back to the model dimension. One block updates row tokens from column context and another updates column tokens from row context. Both use pre-normalisation, residual connections and an MLP. The updated classification tokens are concatenated and passed to either a linear or MLP head. These variants contain 1,999,114 and 2,031,242 parameters, so accuracy is interpreted alongside model size.",
)

# Results numbering, table/figure captions, and chapter story.
replace_prefix(
    doc,
    "This chapter presents the experiments in the order needed",
    "This chapter presents the experiments in the order needed to answer the study questions. It begins with the core positional-encoding comparison, then isolates the shifted schedule, reduces the available training data, moves selected models to CIFAR-100, and finally examines fusion and fixed-PE extensions. Each section states what changes and what remains controlled before reporting the selected-checkpoint test results. Mechanistic explanations and practical recommendations are reserved for Chapter 5 so that observation and interpretation remain distinct.",
)
replace_paragraph(find_paragraph(doc, "4.5 Performance under Limited Training Data"), "4.4 Performance under Limited Training Data")
replace_paragraph(find_paragraph(doc, "Table 9. Reduced-data PE comparison"), "Table 8. Reduced-data PE comparison on CIFAR-10.")
replace_paragraph(find_paragraph(doc, "Figure 7. Mean CIFAR-10 validation accuracy"), "Figure 6. Mean CIFAR-10 validation accuracy for three PE settings and four training-set sizes. Shading shows pointwise 95% t confidence intervals across five seeds.")
replace_prefix(doc, "The test ranking changes with the amount of training data", find_paragraph(doc, "The test ranking changes with the amount of training data").text.replace("Figure 7", "Figure 6"))
replace_paragraph(find_paragraph(doc, "4.6 Generalisation from CIFAR-10 to CIFAR-100"), "4.5 Generalisation from CIFAR-10 to CIFAR-100")
replace_paragraph(find_paragraph(doc, "Table 10. Selected PE configurations"), "Table 9. Selected PE configurations on CIFAR-10 and CIFAR-100.")
replace_paragraph(find_paragraph(doc, "Figure 8. Mean CIFAR-100 validation accuracy"), "Figure 7. Mean CIFAR-100 validation accuracy and loss across five seeds. Shading shows pointwise 95% t confidence intervals. Final comparisons use the test results in Table 9.")
replace_prefix(doc, "The four models kept the same test-accuracy order", find_paragraph(doc, "The four models kept the same test-accuracy order").text.replace("Figure 8", "Figure 7"))
replace_paragraph(find_paragraph(doc, "4.7 Fusion and Other Extensions"), "4.6 Fusion and Other Extensions")
replace_paragraph(find_paragraph(doc, "Table 11. Fusion and fixed-PE extensions"), "Table 10. Fusion and fixed-PE extensions.")
replace_paragraph(find_paragraph(doc, "Figure 9. Mean CIFAR-10 validation accuracy"), "Figure 8. Mean CIFAR-10 validation accuracy for single-branch references and fusion models. Shading shows pointwise 95% t confidence intervals across five seeds. Fusion models are not parameter matched.")
replace_prefix(
    doc,
    "The best dual-branch result came from cross-attention",
    "The best dual-branch result came from cross-attention with an MLP head at 77.74 ± 0.44%, but it remained 0.86 points below learnable PE in the paired test comparison while increasing the parameter count from 809,354 to 2,031,242. Figure 8 separates aggregation-based and cross-attention fusion. The validation trajectories show similar convergence, but Table 10 confirms that none of the larger fusion models improves on the learned single-branch reference at the selected checkpoint.",
)
replace_paragraph(find_paragraph(doc, "4.8 Summary of Empirical Findings"), "4.7 Summary of Empirical Findings")
replace_prefix(
    doc,
    "Across the experiments, positional information improved full-data performance",
    "Across the experiments, positional information improved full-data performance over no PE. Learnable PE had the highest core mean on CIFAR-10, while shifted multiplicative PE was the strongest fixed design. The fixed method led at 1,000 and 5,000 examples, then fell slightly behind learnable PE at 10,000 and full data. The four selected methods retained the same test-accuracy order on CIFAR-100. Dual-branch fusion did not improve on the learned single-branch reference.",
)
replace_prefix(
    doc,
    "The evidence now covers two datasets",
    "The evidence now covers two datasets, four training-set sizes, and both single-branch and dual-branch models. It remains bounded by one small ViT architecture, one split-construction rule, and five seeds. Chapter 5 discusses what these results support and where further evidence is still needed.",
)

# Discussion and Conclusion: RQ3 now concerns data availability and cross-dataset evidence.
replace_paragraph(find_paragraph(doc, "5.3 RQ3: Patch Assignment and Data Availability"), "5.3 RQ3: Data Availability and Cross-Dataset Consistency")
for prefix in (
    "The assignment experiment shows why sequence order",
    "Learnable PE changed by only 0.06 percentage points",
):
    paragraph = find_paragraph(doc, prefix)
    paragraph._element.getparent().remove(paragraph._element)
replace_prefix(
    doc,
    "Training-set size produced a second interaction",
    "Training-set size changed the ordering between the two leading methods. Shifted multiplicative PE led learnable PE at 1,000 and 5,000 examples, while the learned method had the higher mean at 10,000 examples and full data. The four pre-selected methods also retained the same accuracy ranking on CIFAR-100. These findings answer RQ3 within the evaluated settings. They suggest that a deterministic spatial signal can be competitive when fewer examples are available, but they do not identify a universal data threshold or establish broad cross-domain generalisation.",
)
replace_prefix(
    doc,
    "A shared optimisation protocol reduces confounding",
    "A shared optimisation protocol reduces confounding, but it may not be the best setting for every PE method. Checkpoints were selected only with validation accuracy and the test set was evaluated after selection. This limits direct test-set leakage, although examining many variants can still influence which comparisons receive attention.",
)
for prefix in (
    "The four patch mappings passed a deterministic coordinate audit",
):
    paragraph = find_paragraph(doc, prefix)
    paragraph._element.getparent().remove(paragraph._element)
replace_prefix(
    doc,
    "This dissertation examined positional encoding as a controlled design choice",
    "This dissertation examined positional encoding as a controlled design choice in a small Vision Transformer trained from scratch. The patch representation, encoder, optimisation and checkpoint rule were held constant so that the main comparisons could be linked to PE design. The evaluation covered CIFAR-10, CIFAR-100, four CIFAR-10 training-set sizes, and several dual-branch extensions.",
)
replace_prefix(
    doc,
    "The experiments answer the three research questions within this setting",
    "The experiments answer the three research questions within this setting. Positional information was important in full-data training, and learnable absolute PE achieved the highest core mean test accuracy. Shifted multiplicative PE was the strongest fixed design. It led learnable PE at 1,000 and 5,000 training examples, but the learned method had the higher mean at 10,000 examples and full data. The selected methods kept the same test-accuracy order on CIFAR-100, while the dual-branch models did not improve on the learned single-branch reference.",
)

# Move the unresolved conceptual question to Future Work.
future_anchor = find_paragraph(doc, "The assignment and architecture studies can also be made more precise")
replace_paragraph(
    future_anchor,
    "A separate future direction concerns patch unfolding and the meaning of positional coordinates. The preliminary unfolding study reordered patches into row-major, column-major or serpentine traversals and then applied fixed PE by sequence slot. This followed the intended experimental design, but it also changed which positional vector was attached to some physical patches. The preliminary results did not show a consistent benefit, so they are not used as core evidence in this dissertation. Future work should distinguish traversal-based positional reassignment from a coordinate-aligned permutation in which each patch keeps the PE derived from its physical row and column. The latter should first be checked with a deterministic forward-equivalence test, because global self-attention is expected to remain unchanged when patch–PE pairs are permuted together. Related ordering studies such as LOOPE and REOrder [6], [7] could guide a more formal evaluation. Fusion should also be revisited only with parameter-matched single- and dual-branch models so that the fusion operation can be separated from model capacity.",
)
replace_paragraph(find_paragraph(doc, "Table 12. Preliminary learnable and multiplicative PE combination"), "Table 11. Preliminary learnable and multiplicative PE combination.")

# Apply the requested Arial/black formatting to ordinary text runs.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_run_font(run, run.font.size.pt if run.font.size else 11)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, run.font.size.pt if run.font.size else 9)

tmp = DOCX.with_suffix(".unfolding.tmp.docx")
doc.save(tmp)
tmp.replace(DOCX)
print(DOCX)
