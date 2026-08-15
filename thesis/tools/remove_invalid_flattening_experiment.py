from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor


DOCX = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
BACKUP = DOCX.with_name("Yikai_Zhao_MSc_Dissertation.pre_invalid_experiment_removal.docx")
LOW_DATA_FIGURE = Path(r"D:\code\Postgraduate-dissertation\thesis\assets\low_data_validation_accuracy_epoch_valid.png")
FUSION_FIGURE = Path(r"D:\code\Postgraduate-dissertation\thesis\assets\fusion_validation_accuracy_epoch_valid.png")


def set_run_style(run, size=11, bold=None):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text, size=11, bold=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_style(run, size=size, bold=bold)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def heading_level(style_name):
    match = re.fullmatch(r"Heading (\d+)", style_name)
    return int(match.group(1)) if match else None


def remove_heading_section(document, heading_prefix):
    paragraphs = list(document.paragraphs)
    start = next(i for i, p in enumerate(paragraphs) if p.text.strip().startswith(heading_prefix))
    start_level = heading_level(paragraphs[start].style.name)
    if start_level is None:
        raise ValueError(f"Not a heading: {heading_prefix}")
    end = len(paragraphs)
    for i in range(start + 1, len(paragraphs)):
        level = heading_level(paragraphs[i].style.name)
        if level is not None and level <= start_level:
            end = i
            break
    for paragraph in paragraphs[start:end]:
        remove_paragraph(paragraph)


def find_table(document, first_header, second_header=None):
    for table in document.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers and headers[0] == first_header and (second_header is None or len(headers) > 1 and headers[1] == second_header):
            return table
    raise ValueError(f"Table not found: {first_header} / {second_header}")


def remove_table(table):
    element = table._element
    element.getparent().remove(element)


def remove_rows_where(table, column_index, predicate):
    for row in list(table.rows[1:]):
        if predicate(row.cells[column_index].text.strip()):
            table._element.remove(row._tr)


def replace_image(paragraph, image_path, width_emu):
    paragraph.clear()
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Emu(width_emu))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_text_nodes(document, replacements):
    for node in document.element.body.iter(qn("w:t")):
        if not node.text:
            continue
        text = node.text
        for old, new in replacements:
            text = text.replace(old, new)
        node.text = text


def renumber_citations(document):
    mapping = {8: 6, 9: 7, 10: 8, 11: 9, 12: 10, 13: 11, 14: 12}
    pattern = re.compile(r"\[(\d+)\]")
    for node in document.element.body.iter(qn("w:t")):
        if node.text:
            node.text = pattern.sub(lambda m: f"[{mapping.get(int(m.group(1)), int(m.group(1)))}]", node.text)


def rewrite(document):
    # Front matter, abstract, and introduction.
    set_paragraph_text(find_paragraph(document, "Fixed and learnable 2D encodings,"),
                       "Fixed and learnable 2D encodings, data availability, and model extensions on CIFAR-10 and CIFAR-100")
    abstract = find_paragraph(document, "Vision Transformers (ViTs) are widely used")
    set_paragraph_text(abstract,
        "Vision Transformers (ViTs) are widely used for image recognition, but self-attention alone does not represent the spatial arrangement of image patches. Positional encoding (PE) supplies this missing information. This dissertation evaluates learned and fixed PE in a small ViT trained from scratch, using CIFAR-10 as the main benchmark and five training seeds for each comparison. Positional encoding improved full-data performance over a model without PE. Learnable absolute PE achieved the highest mean accuracy in the core comparison, while shifted multiplicative PE was the strongest fixed method. Under reduced training data, shifted multiplicative PE led at 1,000 and 5,000 examples, but learnable PE recovered the advantage with more data. The selected methods kept the same test-accuracy ranking on CIFAR-100, while dual-branch models did not provide a clear benefit over the simpler learned model. These results show that PE design matters, but its value depends on the data setting and does not necessarily increase with architectural complexity.")
    set_paragraph_text(find_paragraph(document, "The original ViT uses a learned vector"),
        "The original ViT uses a learned vector for each token position [1]. These vectors can change during training and adapt to the task. However, they do not directly use the two-dimensional row and column structure of the image. A fixed two-dimensional encoding instead creates each positional vector from the row and column of the patch. It does not add trainable positional parameters and directly encodes the position of each patch. Unlike learned PE, however, its values remain fixed during training. This raises the question of whether learned and fixed PE behave differently as the amount of training data or classification difficulty changes.")
    set_paragraph_text(find_paragraph(document, "Many previous studies compare positional encoding"),
        "Many previous studies compare positional encoding methods using different model architectures or training settings [5], [10]–[12]. This makes it difficult to know whether a change in performance comes from the positional encoding itself or from other differences in the model. This project therefore keeps the patch size, ViT backbone, and training process the same. It studies whether positional encoding improves the model, which fixed PE methods perform best, and whether the results change with data availability or classification difficulty. Dual-branch models are tested separately because they also change the architecture and number of parameters.")
    set_paragraph_text(find_paragraph(document, "The study is organised around three research questions."),
        "The study is organised around three research questions. RQ1 asks whether positional encoding is necessary for the evaluated ViT and how learned absolute PE compares with fixed alternatives. RQ2 examines which fixed constructions, including axis-wise, additive, multiplicative, shifted, squared and radial designs, are most competitive. RQ3 considers how the comparison between learned and fixed PE changes with the amount of training data and with the more difficult CIFAR-100 classification task.")
    set_paragraph_text(find_paragraph(document, "The experiments show that positional information"),
        "The experiments show that positional information is important for the ViT used in this project. Learnable absolute PE has the highest mean accuracy in the full-data core comparison, while shifted multiplicative PE is the strongest fixed method. The fixed method leads at 1,000 and 5,000 training examples, but not at 10,000 examples or with the full training set. The same four-model test-accuracy ranking is observed on CIFAR-100. Dual-branch models add complexity without a clear improvement over the learned single-branch model.")

    # Remove literature devoted to the invalid experiment, then close the review around valid evidence.
    remove_heading_section(document, "2.4 Patch Ordering")
    set_paragraph_text(find_paragraph(document, "2.5 Data-Efficient"), "2.4 Data-Efficient and Data-Limited ViT Training", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "2.6 Related-Work"), "2.5 Related-Work Synthesis and Gap", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "The practical gap is therefore one of attribution"),
        "The practical gap is therefore one of attribution, not a lack of proposed encodings. Results from different studies may reflect changes in architecture, scale, data, augmentation or optimisation as well as the position method. This project instead uses one ViT and a common protocol to separate three factors: PE presence and form, the amount of available training data, and the cost of dual-branch complexity. This controlled comparison does not claim a universally superior encoding or a new general-purpose Transformer.")

    # Method overview and low-data protocol.
    set_paragraph_text(find_paragraph(document, "This chapter explains how the experiments were constructed"),
        "This chapter explains how the experiments were constructed so that each reported comparison can be reproduced and interpreted. Every image passes through the same broad pipeline: dataset-specific preprocessing, division into patches, addition of positional information, and processing by a shared Transformer encoder. Validation performance selects the checkpoint. The official test set is used only after selection. The following sections define the controlled design, data, shared architecture, positional encodings, dual-branch extensions, training protocol and statistical analysis.")
    set_paragraph_text(find_paragraph(document, "The study changes one main factor at a time"),
        "The study changes one main factor at a time rather than introducing a new general-purpose Transformer. Within each comparison, the ViT backbone, data split, optimiser and checkpoint-selection rule remain fixed. The main factor that changes is the positional encoding. Dual-branch fusion models are treated as extensions because they also change the architecture and parameter count. This design reduces the number of competing explanations for an observed accuracy difference.")
    set_paragraph_text(find_paragraph(document, "The main evaluation uses CIFAR-10 and five training seeds"),
        "The main evaluation uses CIFAR-10 and five training seeds. No PE provides a genuine no-position baseline, while learnable absolute PE provides the standard learned reference. Fixed row-column encodings then test different ways to represent the two-dimensional grid without trainable positional parameters. Fusion models are analysed separately because they change more than the positional vector. CIFAR-100 tests whether the observed ranking is retained on a finer-grained classification problem, and reduced-data CIFAR-10 tests whether the relationship changes when fewer labelled images are available.")
    set_paragraph_text(find_paragraph(document, "The reduced-data study draws subsets"),
        "The reduced-data study draws subsets of 1,000, 5,000 and 10,000 images from the fixed 45,000-image CIFAR-10 training pool. It compares no PE, learnable absolute PE and shifted multiplicative PE using the common learning rate in Section 3.6. The full-data results use the same optimisation settings and are included as a fourth data-size condition. Validation and test sets do not change. Within each seed, all three models receive the same subset. The subset changes across seeds, so the reported variation includes both stochastic training and the sampled training images.")
    set_paragraph_text(find_paragraph(document, "Equation (9) combines"),
        "Equation (9) combines the axis-specific mappings using the same addition and multiplication rules as before. The superscripts label the frequency schedule used for each axis. They are not powers. The term shifted therefore refers only to the one-index offset between the two schedules. It does not move a patch, change its row or column, or alter token order.")
    set_paragraph_text(find_paragraph(document, "The core comparison is deliberately limited"),
        "The core comparison is deliberately limited to methods that produce a 128-component positional term and add it at the same point in the shared ViT. Other approaches reviewed in Sections 2.3 and 2.5, including relative-position, conditional and rotary encodings [10]–[12], either modify the attention calculation or introduce learned position-generating components. Including them would change more than the positional vector itself. The dual-branch models also alter parameter count and architecture, so they are evaluated separately in Section 3.5.")

    # Remove invalid method and the hybrid that inherited the same incorrect alignment.
    remove_heading_section(document, "3.5 Patch Ordering")
    remove_heading_section(document, "3.6.1 Learnable-Fixed Hybrid")
    set_paragraph_text(find_paragraph(document, "3.6 Hybrid and Dual-Branch"), "3.5 Dual-Branch Fusion Extensions", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "The extension models ask whether extra positional"),
        "The extension models ask whether architectural complexity improves on the single-branch comparison. Each model processes row and column positional signals in separate ViT branches and then fuses their representations. Because these designs duplicate the encoder and add fusion components, they are reported separately from the core comparison and interpreted alongside their parameter counts.")
    set_paragraph_text(find_paragraph(document, "3.6.2 Latent Fusion"), "3.5.1 Latent Fusion", size=12, bold=True)
    set_paragraph_text(find_paragraph(document, "3.6.3 Bidirectional"), "3.5.2 Bidirectional Cross-Attention Fusion", size=12, bold=True)
    set_paragraph_text(find_paragraph(document, "3.7 Training and"), "3.6 Training and Checkpoint-Selection Protocol", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "3.8 Evaluation Metrics"), "3.7 Evaluation Metrics and Statistical Analysis", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "3.9 Implementation"), "3.8 Implementation and Reproducibility", size=13, bold=True)

    replace_text_nodes(document, [
        ("Equation (13)", "Equation (11)"),
        ("Equation (14)", "Equation (12)"),
        ("#13", "#11"),
        ("#14", "#12"),
    ])

    # Remove the invalid results table/section and update valid result groups.
    remove_table(find_table(document, "PE family", "Assignment"))
    remove_heading_section(document, "4.4 Patch Ordering")
    set_paragraph_text(find_paragraph(document, "This chapter presents the experiments"),
        "This chapter presents the experiments in the order needed to answer the study questions. It begins with the core positional-encoding comparison, then isolates the shifted schedule, reduces the available training data, moves selected models to CIFAR-100, and finally examines dual-branch fusion extensions. Each section states what changes and what remains controlled before reporting the selected-checkpoint test results. Mechanistic explanations and practical recommendations are reserved for Chapter 5 so that observation and interpretation remain distinct.")
    set_paragraph_text(find_paragraph(document, "4.5 Performance under Limited"), "4.4 Performance under Limited Training Data", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "The reduced-data experiment asks"),
        "The reduced-data experiment asks whether a fixed spatial prior becomes more useful when fewer labelled images are available. Three models are compared at 1,000, 5,000, 10,000 and the full 45,000-example training split: no PE, learnable PE and shifted multiplicative PE. Within each data size, all three models use the same sampled images within a seed and share the architecture, learning rate, optimisation, checkpoint rule and seed set. The paired comparisons therefore isolate the PE choice as far as this protocol allows.")
    set_paragraph_text(find_paragraph(document, "Table 9. Reduced-data"), "Table 8. Reduced-data PE comparison on CIFAR-10.", size=10)
    low_table = find_table(document, "Training examples", "Model")
    remove_rows_where(low_table, 1, lambda value: "Learnable + multiplicative" in value)
    set_paragraph_text(find_paragraph(document, "The test ranking changes with the amount"),
        "The test ranking changes with the amount of training data. Shifted multiplicative PE exceeded learnable PE by 3.37 percentage points at 1,000 examples with a paired 95% CI from 2.55 to 4.19. It also led by 1.12 points at 5,000 examples with an interval from 0.46 to 1.78. At 10,000 examples, the difference was −0.65 points with an interval from −1.31 to 0.02. The full-data comparison favoured learnable PE by 0.52 points with an interval from 0.15 to 0.89. Figure 6 shows the validation trajectories under the four data sizes. The fixed-versus-learnable separation is most visible at 1,000 and 5,000 examples.")
    replace_image(find_paragraph(document, "Figure 7. Mean CIFAR-10 validation accuracy for four models" )._element.getprevious() if False else document.paragraphs[document.paragraphs.index(find_paragraph(document, "Figure 7. Mean CIFAR-10 validation accuracy for four models"))-1], LOW_DATA_FIGURE, 5394960)
    set_paragraph_text(find_paragraph(document, "Figure 7. Mean CIFAR-10 validation accuracy for four models"),
        "Figure 6. Mean CIFAR-10 validation accuracy for three models and four training-set sizes. Shading shows pointwise 95% t confidence intervals across five seeds.", size=10)

    set_paragraph_text(find_paragraph(document, "4.6 Generalisation"), "4.5 Generalisation from CIFAR-10 to CIFAR-100", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "Table 10. Selected PE"), "Table 9. Selected PE configurations on CIFAR-10 and CIFAR-100.", size=10)
    replace_text_nodes(document, [("Figure 8 shows", "Figure 7 shows")])
    set_paragraph_text(find_paragraph(document, "Figure 8. Mean CIFAR-100"),
        "Figure 7. Mean CIFAR-100 validation accuracy and loss across five seeds. Shading shows pointwise 95% t confidence intervals. Final comparisons use the test results in Table 9.", size=10)

    set_paragraph_text(find_paragraph(document, "4.7 Hybrid, Fusion"), "4.6 Fusion and Other Extensions", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "The final group asks whether extra complexity"),
        "The final group asks whether extra architectural complexity earns better performance than the single-branch models. Fusion models are shown with parameter count because they duplicate the encoder or add learned fusion layers. Squared and radial PE remain in the table as targeted fixed-encoding tests, even though they were not among the strongest core methods.")
    set_paragraph_text(find_paragraph(document, "Table 11. Hybrid, fusion"), "Table 10. Fusion and fixed-PE extensions.", size=10)
    fusion_table = find_table(document, "Model", "Category")
    remove_rows_where(fusion_table, 0, lambda value: value in {"Order-matched learned PE", "Hybrid learned + fixed PE"})
    set_paragraph_text(find_paragraph(document, "The hybrid differed from its order-matched"),
        "The best dual-branch result came from cross-attention with an MLP head at 77.74 ± 0.44%, but it remained 0.86 points below learnable PE in the paired test comparison while increasing the parameter count from 809,354 to 2,031,242. Figure 8 separates aggregation-based and cross-attention fusion. The validation trajectories show similar convergence, but Table 10 confirms that none of the larger fusion models improves on the learned single-branch reference at the selected checkpoint.")
    fusion_caption = find_paragraph(document, "Figure 9. Mean CIFAR-10 validation accuracy")
    fusion_image = document.paragraphs[document.paragraphs.index(fusion_caption) - 1]
    replace_image(fusion_image, FUSION_FIGURE, 5623560)
    set_paragraph_text(fusion_caption,
        "Figure 8. Mean CIFAR-10 validation accuracy for the single-branch reference and fusion models. Shading shows pointwise 95% t confidence intervals across five seeds. Fusion models are not parameter matched.", size=10)
    set_paragraph_text(find_paragraph(document, "4.8 Summary of"), "4.7 Summary of Empirical Findings", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "Across the experiments, positional information improved"),
        "Across the experiments, positional information improved full-data performance over no PE. Learnable PE had the highest core mean on CIFAR-10, while shifted multiplicative PE was the strongest fixed design. The fixed method led at 1,000 and 5,000 examples, then fell slightly behind learnable PE at 10,000 examples and full data. The four selected methods retained the same test-accuracy order on CIFAR-100. Dual-branch models did not provide a clear improvement over the learned single-branch reference.")
    set_paragraph_text(find_paragraph(document, "The evidence now covers two datasets"),
        "The evidence now covers two datasets, four training-set sizes, and both single-branch and multi-branch models. It remains bounded by one small ViT architecture, one split-construction rule, and five seeds. Chapter 5 discusses what these results support and where further evidence is still needed.")

    # Discussion: remove all conclusions drawn from invalid assignment and hybrid results.
    remove_paragraph(find_paragraph(document, "The assignment experiment shows"))
    remove_paragraph(find_paragraph(document, "Learnable PE changed by only"))
    set_paragraph_text(find_paragraph(document, "5.3 RQ3:"), "5.3 RQ3: Data Availability and Classification Difficulty", size=13, bold=True)
    generalisation = find_paragraph(document, "CIFAR-100 provides a harder classification setting")
    practical_heading = find_paragraph(document, "5.4 Generalisation and practical implications")
    practical_heading._element.addprevious(generalisation._element)
    set_paragraph_text(practical_heading, "5.4 Practical Implications", size=13, bold=True)
    set_paragraph_text(find_paragraph(document, "The hybrid and fusion results"),
        "The fusion results also show that extra complexity is not enough on its own. The best fusion model remained below learnable PE even though it used about 2.5 times as many trainable parameters. Because the fusion models were not parameter matched, the experiment cannot isolate the effect of the fusion operation from the change in capacity. It can still support the practical conclusion that the added computation did not produce a better result under the tested design.")
    internal = find_paragraph(document, "A shared optimisation protocol reduces confounding")
    set_paragraph_text(internal,
        "A shared optimisation protocol reduces confounding, but it may not be the best setting for every PE method. Checkpoints were selected only with validation accuracy and the test set was evaluated after selection. This limits direct test-set leakage, although examining many variants can still influence which comparisons receive attention.")
    remove_paragraph(find_paragraph(document, "The four patch mappings passed"))

    # Conclusion and future work.
    set_paragraph_text(find_paragraph(document, "This dissertation examined positional encoding as a controlled"),
        "This dissertation examined positional encoding as a controlled design choice in a small Vision Transformer trained from scratch. The patch representation, encoder, optimisation and checkpoint rule were held constant so that the main comparisons could be linked to PE design. The evaluation covered CIFAR-10, CIFAR-100, four CIFAR-10 training-set sizes, several fixed PE extensions, and dual-branch fusion models.")
    set_paragraph_text(find_paragraph(document, "The experiments answer the three research questions"),
        "The experiments answer the three research questions within this setting. Positional information was important in full-data training, and learnable absolute PE achieved the highest core mean test accuracy. Shifted multiplicative PE was the strongest fixed design. It led learnable PE at 1,000 and 5,000 training examples, but the learned method had the higher mean at 10,000 examples and full data. The selected methods kept the same test-accuracy order on CIFAR-100, while the dual-branch models did not provide a clear improvement over the learned single-branch reference.")
    set_paragraph_text(find_paragraph(document, "The main contribution is therefore empirical"),
        "The main contribution is therefore empirical rather than a new general-purpose ViT architecture. The study shows that PE choices can be compared under a shared protocol, that the amount of training data can change the ordering between learned and fixed PE, and that more complex fusion mechanisms do not automatically improve classification. Learned PE is the stronger full-data default for this model, while shifted multiplicative PE is a credible option when labelled data are limited. These conclusions remain bounded by the low-resolution datasets, model scale, fixed data split and five-seed protocol.")
    set_paragraph_text(find_paragraph(document, "The assignment and architecture studies can also"),
        "The fusion study should be revisited with parameter-matched single- and dual-branch models. This would separate the effect of the fusion operation from the increase in model capacity and provide a fairer test of whether row-column branches offer useful complementary information.")

    # Remove LOOPE and REOrder, which were included solely for the invalid experiment.
    remove_paragraph(find_paragraph(document, "[6] M. A. M. Chowdhury"))
    remove_paragraph(find_paragraph(document, "[7] D. Kutscher"))
    renumber_citations(document)


def audit(document):
    full_text = "\n".join(p.text for p in document.paragraphs)
    forbidden = (
        "patch-to-position", "Patch Ordering", "patch assignment", "patch mappings",
        "normal_col", "proper_row", "proper_col", "serpentine", "LOOPE", "REOrder",
        "hybrid", "Learnable + multiplicative", "order-matched learned",
    )
    hits = [term for term in forbidden if term.lower() in full_text.lower()]
    if hits:
        raise RuntimeError(f"Forbidden invalid-experiment terms remain: {hits}")
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    for removed in ("2.4 Patch Ordering", "3.5 Patch Ordering", "4.4 Patch Ordering"):
        if any(h.startswith(removed) for h in headings):
            raise RuntimeError(f"Removed heading remains: {removed}")
    if len(document.tables) != 10:
        raise RuntimeError(f"Expected 10 tables, found {len(document.tables)}")
    low = find_table(document, "Training examples", "Model")
    if len(low.rows) != 13:
        raise RuntimeError(f"Expected 12 low-data results plus header, found {len(low.rows)} rows")
    fusion = find_table(document, "Model", "Category")
    if len(fusion.rows) != 10:
        raise RuntimeError(f"Expected 9 valid extension results plus header, found {len(fusion.rows)} rows")
    refs = [p.text for p in document.paragraphs if re.match(r"^\[\d+\]", p.text)]
    if len(refs) != 12 or not all(ref.startswith(f"[{i}]") for i, ref in enumerate(refs, 1)):
        raise RuntimeError("Reference list is not consecutively numbered from [1] to [12].")
    body = full_text.split("References", 1)[0]
    cited = set()
    for group in re.findall(r"\[([0-9,\s–-]+)\]", body):
        for part in group.split(","):
            part = part.strip()
            match = re.fullmatch(r"(\d+)\s*[–-]\s*(\d+)", part)
            if match:
                cited.update(range(int(match.group(1)), int(match.group(2)) + 1))
            elif part.isdigit():
                cited.add(int(part))
    if cited != set(range(1, 13)):
        raise RuntimeError(f"Citation coverage mismatch: {sorted(cited)}")


def main():
    if not LOW_DATA_FIGURE.exists() or not FUSION_FIGURE.exists():
        raise FileNotFoundError("Regenerated valid figures are missing.")
    shutil.copy2(DOCX, BACKUP)
    document = Document(DOCX)
    rewrite(document)
    audit(document)
    document.save(DOCX)
    print(f"Saved: {DOCX}")
    print(f"Backup: {BACKUP}")
    print("Invalid flattening/assignment experiment and affected hybrid evidence removed.")


if __name__ == "__main__":
    main()
