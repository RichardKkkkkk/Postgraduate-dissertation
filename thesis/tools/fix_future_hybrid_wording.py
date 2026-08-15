from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
doc = Document(path)
old = (
    "The reduced-data results motivate a more complete study of learned and fixed PE combinations. "
    "A preliminary model added an unshifted multiplicative encoding to the learnable table through one "
    "trainable global scale. Table 12 reports this initial attempt alongside the two components that "
    "motivated it. The hybrid remained close to learnable PE at every training size and did not provide "
    "a clear improvement. It is included here as evidence that the idea was explored, rather than as a "
    "main method or conclusion of this dissertation."
)
new = (
    "The reduced-data results motivate a more complete study of learned and fixed PE combinations. "
    "A preliminary model added an unshifted multiplicative encoding to the learnable table through one "
    "trainable global scale. Table 12 reports this initial attempt alongside the learned and shifted "
    "multiplicative reference methods. The hybrid remained close to learnable PE at every training size "
    "and did not provide a clear improvement. It is included here as evidence that the idea was explored, "
    "rather than as a main method or conclusion of this dissertation."
)
for paragraph in doc.paragraphs:
    if paragraph.text == old:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)
        run = paragraph.add_run(new)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Arial")
        break
else:
    raise SystemExit("Target wording not found")
tmp = path.with_suffix(".wording.tmp.docx")
doc.save(tmp)
tmp.replace(path)
