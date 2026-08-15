from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
doc = Document(path)

replacements = {
    "The reduced-data results motivate a more complete study": (
        "The reduced-data results motivate a more complete study of learned and fixed PE combinations. "
        "A preliminary model added an unshifted multiplicative encoding to the learnable table through one "
        "trainable global scale. Table 11 reports this initial attempt alongside the learned and shifted "
        "multiplicative reference methods. The hybrid remained close to learnable PE at every training size "
        "and did not provide a clear improvement. It is included here as evidence that the idea was explored, "
        "rather than as a main method or conclusion of this dissertation."
    ),
    "The preliminary model does not test the most promising combination directly": (
        "The preliminary model does not test the most promising combination directly. Its fixed component "
        "was the unshifted multiplicative PE, while the reduced-data advantage was observed for shifted "
        "multiplicative PE. Future work should therefore test learnable PE with shifted multiplicative PE, "
        "compare different initial values for the scale, and replace the single global scale with layer-wise "
        "or dimension-wise gates. The study should also record the scale throughout training. These variants "
        "require a new, pre-defined experiment rather than further interpretation of the preliminary result."
    ),
}

for paragraph in doc.paragraphs:
    for prefix, text in replacements.items():
        if paragraph.text.startswith(prefix):
            for child in list(paragraph._p):
                if child.tag != qn("w:pPr"):
                    paragraph._p.remove(child)
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 0, 0)
            rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{attr}"), "Arial")

tmp = path.with_suffix(".residuals.tmp.docx")
doc.save(tmp)
tmp.replace(path)
