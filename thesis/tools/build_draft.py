from __future__ import annotations

import csv
import json
import math
import re
import statistics
import textwrap
from collections import defaultdict
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
EXPERIMENT = ROOT / "results" / "cifar10_final_vit_models_5seeds"
METRICS = EXPERIMENT / "metrics"
REPORT = EXPERIMENT / "reports" / "thesis_core"
FIGURES = REPORT / "figures"
OUTPUT = ROOT / "thesis" / "Yikai_Zhao_Dissertation_Structured_Draft_v0.1.docx"

INK = "1A1A1A"
MUTED = "555555"
LIGHT = "F1F3F5"
GRID = "B8BDC4"
ACCENT = "23395B"
BLUE = "2F5597"
PALE_BLUE = "EAF0F8"


LABELS = {
    "vit_baseline": "No positional encoding",
    "vit_learnable_position": "Learned absolute PE",
    "vit_row_sinusoidal": "Row sinusoidal",
    "vit_col_sinusoidal": "Column sinusoidal",
    "vit_additive_sinusoidal": "Additive 2D sinusoidal",
    "vit_additive_sinusoidal_shifted": "Shifted additive 2D",
    "vit_multiplicative_sinusoidal": "Multiplicative 2D sinusoidal",
    "vit_multiplicative_sinusoidal_shifted": "Shifted multiplicative 2D",
    "vit_squared_multiplicative_sinusoidal": "Squared multiplicative 2D",
    "vit_squared_multiplicative_sinusoidal_shifted": "Shifted squared multiplicative",
    "vit_radial_sinusoidal": "Radial sinusoidal",
    "vit_normal_col_learnable_multiplicative_sinusoidal": "Hybrid learned + multiplicative",
    "vit_row_col_latent_fusion": "Row/column concat + MLP",
    "vit_row_col_mean_fusion": "Row/column mean",
    "vit_row_col_mean_mlp_fusion": "Row/column mean + MLP",
    "vit_row_col_cross_attention_fusion": "Bidirectional cross-attention",
    "vit_row_col_cross_attention_mlp_head_fusion": "Cross-attention + MLP head",
}


PE_SUFFIX_LABELS = {
    "baseline": "No PE",
    "learnable_position": "Learned PE",
    "row_sinusoidal": "Row sinusoidal",
    "col_sinusoidal": "Column sinusoidal",
    "multiplicative_sinusoidal": "Multiplicative",
}


ORDER_PREFIXES = {
    "normal_row": "Row-major",
    "normal_col": "Column-major",
    "proper_row": "Serpentine rows",
    "proper_col": "Serpentine columns",
}


REFERENCES = [
    "A. Vaswani et al., \"Attention Is All You Need,\" in Advances in Neural Information Processing Systems, vol. 30, 2017.",
    "A. Dosovitskiy et al., \"An Image Is Worth 16x16 Words: Transformers for Image Recognition at Scale,\" in International Conference on Learning Representations, 2021.",
    "A. Krizhevsky, \"Learning Multiple Layers of Features from Tiny Images,\" University of Toronto, Technical Report, 2009.",
    "H. Touvron et al., \"Training Data-Efficient Image Transformers and Distillation through Attention,\" in Proceedings of the 38th International Conference on Machine Learning, pp. 10347-10357, 2021.",
    "P. Shaw, J. Uszkoreit, and A. Vaswani, \"Self-Attention with Relative Position Representations,\" in Proceedings of NAACL-HLT, pp. 464-468, 2018, doi: 10.18653/v1/N18-2074.",
    "K. Wu, H. Peng, M. Chen, J. Fu, and H. Chao, \"Rethinking and Improving Relative Position Encoding for Vision Transformer,\" in Proceedings of the IEEE/CVF International Conference on Computer Vision, pp. 10033-10041, 2021.",
    "X. Chu, Z. Tian, B. Zhang, X. Wang, and C. Shen, \"Conditional Positional Encodings for Vision Transformers,\" arXiv:2102.10882, 2021.",
    "J. Su, Y. Lu, S. Pan, B. Murtadha, B. Wen, and Y. Liu, \"RoFormer: Enhanced Transformer with Rotary Position Embedding,\" arXiv:2104.09864, 2021.",
    "Z. Liu et al., \"Swin Transformer: Hierarchical Vision Transformer Using Shifted Windows,\" in Proceedings of the IEEE/CVF International Conference on Computer Vision, pp. 10012-10022, 2021.",
    "C.-F. Chen, Q. Fan, and R. Panda, \"CrossViT: Cross-Attention Multi-Scale Vision Transformer for Image Classification,\" in Proceedings of the IEEE/CVF International Conference on Computer Vision, pp. 357-366, 2021.",
    "I. Loshchilov and F. Hutter, \"Decoupled Weight Decay Regularization,\" in International Conference on Learning Representations, 2019.",
    "M. A. M. Chowdhury, M. R. U. Rahman, and A. A. Taki, \"LOOPE: Learnable Optimal Patch Order in Positional Embeddings for Vision Transformers,\" arXiv:2504.14386, 2025.",
    "D. P. Kingma and J. Ba, \"Adam: A Method for Stochastic Optimization,\" in International Conference on Learning Representations, 2015.",
]


def load_results():
    rows = []
    for path in sorted(METRICS.glob("*/*_summary.json")):
        data = json.loads(path.read_text(encoding="utf-8"))
        selected = data["selected_model"]
        cfg = data["config"]
        rows.append(
            {
                "model": cfg["model"],
                "seed": int(cfg["seed"]),
                "test_acc": float(selected["test_acc"]),
                "test_macro_f1": float(selected["test_macro_f1"]),
                "val_acc": float(selected["val_acc"]),
                "selected_epoch": int(selected["epoch"]),
                "protocol": data.get("test_evaluation_protocol"),
                "path": str(path.relative_to(ROOT)),
            }
        )
    if len(rows) != 160:
        raise RuntimeError(f"Expected 160 final summaries, found {len(rows)}")
    if {r["protocol"] for r in rows} != {"selected_checkpoint_only"}:
        raise RuntimeError("Final result protocol is not uniform")
    frame = pd.DataFrame(rows)
    if frame.groupby("model").size().nunique() != 1 or int(frame.groupby("model").size().iloc[0]) != 5:
        raise RuntimeError("Expected five seeds for every model")

    agg = (
        frame.groupby("model")
        .agg(
            mean_acc=("test_acc", "mean"),
            sd_acc=("test_acc", "std"),
            mean_f1=("test_macro_f1", "mean"),
            sd_f1=("test_macro_f1", "std"),
            mean_epoch=("selected_epoch", "mean"),
        )
        .reset_index()
        .sort_values("mean_acc", ascending=False)
        .reset_index(drop=True)
    )
    agg["rank"] = np.arange(1, len(agg) + 1)
    agg["label"] = agg["model"].map(lambda x: LABELS.get(x, readable_model_name(x)))
    agg.to_csv(REPORT / "aggregate_results.csv", index=False)
    frame.to_csv(REPORT / "per_seed_results.csv", index=False)

    learned = frame[frame.model == "vit_learnable_position"][["seed", "test_acc"]].rename(
        columns={"test_acc": "learned_acc"}
    )
    paired = frame.merge(learned, on="seed")
    paired["delta_vs_learned_pp"] = 100 * (paired["test_acc"] - paired["learned_acc"])
    paired.to_csv(REPORT / "paired_deltas_vs_learned.csv", index=False)
    return frame, agg


def readable_model_name(name):
    core = name.removeprefix("vit_")
    for prefix, order in ORDER_PREFIXES.items():
        token = prefix + "_"
        if core.startswith(token):
            suffix = core[len(token):]
            return f"{order}: {PE_SUFFIX_LABELS.get(suffix, suffix.replace('_', ' '))}"
    return core.replace("_", " ").title()


def aggregate_lookup(agg):
    return {row.model: row for row in agg.itertuples(index=False)}


def pct(x):
    return 100.0 * float(x)


def acc_text(row):
    return f"{pct(row.mean_acc):.3f}% +/- {pct(row.sd_acc):.3f} pp"


def get_font(size, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def save_image_pair(image, stem):
    png = FIGURES / f"{stem}.png"
    pdf = FIGURES / f"{stem}.pdf"
    image.convert("RGB").save(png, dpi=(300, 300))
    page_w = image.width * 72.0 / 300.0
    page_h = image.height * 72.0 / 300.0
    pdf_canvas = canvas.Canvas(str(pdf), pagesize=(page_w, page_h))
    pdf_canvas.drawImage(ImageReader(str(png)), 0, 0, width=page_w, height=page_h)
    pdf_canvas.showPage()
    pdf_canvas.save()
    return png


def draw_overview():
    w, h = 2200, 980
    image = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(image)
    title = get_font(58, True)
    body = get_font(34)
    small = get_font(27)
    d.text((w / 2, 55), "Experimental study design", font=title, fill="#17233A", anchor="ma")
    boxes = [
        (80, 260, 410, 590, "CIFAR-10", "45k train\n5k validation\n10k test"),
        (520, 260, 930, 590, "Base ViT", "8 x 8 patch grid\n128-d tokens\n4 blocks, 4 heads"),
        (1040, 170, 1500, 350, "Positional encodings", "fixed, learned, 2D, hybrid"),
        (1040, 405, 1500, 585, "Patch order", "row, column, serpentine"),
        (1040, 640, 1500, 820, "Row-column fusion", "mean, MLP, cross-attention"),
        (1630, 260, 2100, 590, "Final evaluation", "5 training seeds\nvalidation-selected checkpoint\none holdout test evaluation"),
    ]
    for x1, y1, x2, y2, heading, sub in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, fill="#F3F6FA", outline="#385B88", width=4)
        d.text(((x1 + x2) / 2, y1 + 52), heading, font=body, fill="#18324F", anchor="ma")
        d.multiline_text(((x1 + x2) / 2, y1 + 125), sub, font=small, fill="#333333", anchor="ma", align="center", spacing=10)
    arrows = [((410, 425), (520, 425)), ((930, 425), (1040, 260)), ((930, 425), (1040, 495)), ((930, 425), (1040, 730)), ((1500, 260), (1630, 425)), ((1500, 495), (1630, 425)), ((1500, 730), (1630, 425))]
    for start, end in arrows:
        d.line((start, end), fill="#6B7785", width=6)
        ex, ey = end
        sx, sy = start
        angle = math.atan2(ey - sy, ex - sx)
        size = 20
        points = [(ex, ey), (ex - size * math.cos(angle - 0.5), ey - size * math.sin(angle - 0.5)), (ex - size * math.cos(angle + 0.5), ey - size * math.sin(angle + 0.5))]
        d.polygon(points, fill="#6B7785")
    d.text((w / 2, 915), "All comparisons use the same data split and training configuration.", font=small, fill="#555555", anchor="ma")
    return save_image_pair(image, "study_overview")


def draw_errorbar_chart(agg, models, stem, title_text, x_min=None):
    selected = agg.set_index("model").loc[models].sort_values("mean_acc")
    w = 2200
    h = 250 + 115 * len(selected)
    image = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(image)
    title = get_font(54, True)
    label = get_font(31)
    tick = get_font(28)
    d.text((w / 2, 45), title_text, font=title, fill="#17233A", anchor="ma")
    left, right, top, bottom = 780, 2100, 150, h - 100
    vals = 100 * selected.mean_acc.to_numpy()
    errs = 100 * selected.sd_acc.to_numpy()
    lo = min(vals - errs)
    hi = max(vals + errs)
    xmin = x_min if x_min is not None else math.floor((lo - 0.4) * 2) / 2
    xmax = math.ceil((hi + 0.4) * 2) / 2
    if xmax - xmin < 2:
        xmax = xmin + 2
    d.line((left, top, left, bottom), fill="#444444", width=3)
    d.line((left, bottom, right, bottom), fill="#444444", width=3)
    for t in np.linspace(xmin, xmax, 5):
        x = left + (t - xmin) / (xmax - xmin) * (right - left)
        d.line((x, top, x, bottom), fill="#E1E4E8", width=2)
        d.text((x, bottom + 24), f"{t:.1f}%", font=tick, fill="#444444", anchor="ma")
    colors = ["#4C78A8", "#F58518", "#54A24B", "#E45756", "#72B7B2", "#B279A2", "#FF9DA6", "#9D755D", "#BAB0AC", "#2F4B7C", "#A05195"]
    for idx, (model, row) in enumerate(selected.iterrows()):
        y = top + (idx + 0.5) * (bottom - top) / len(selected)
        mean = 100 * row.mean_acc
        err = 100 * row.sd_acc
        x = left + (mean - xmin) / (xmax - xmin) * (right - left)
        xl = left + (mean - err - xmin) / (xmax - xmin) * (right - left)
        xr = left + (mean + err - xmin) / (xmax - xmin) * (right - left)
        d.text((left - 30, y), LABELS.get(model, readable_model_name(model)), font=label, fill="#222222", anchor="rm")
        d.line((xl, y, xr, y), fill="#555555", width=5)
        d.line((xl, y - 12, xl, y + 12), fill="#555555", width=4)
        d.line((xr, y - 12, xr, y + 12), fill="#555555", width=4)
        d.ellipse((x - 13, y - 13, x + 13, y + 13), fill=colors[idx % len(colors)], outline="#222222", width=2)
        d.text((min(xr + 20, right - 10), y), f"{mean:.2f}", font=tick, fill="#333333", anchor="lm")
    d.text(((left + right) / 2, h - 25), "Selected-checkpoint test accuracy; point = mean, error bar = +/- 1 SD across five seeds", font=tick, fill="#555555", anchor="ms")
    return save_image_pair(image, stem)


def order_model_name(order, suffix):
    if order == "normal_row":
        return {
            "baseline": "vit_baseline",
            "learnable_position": "vit_learnable_position",
            "row_sinusoidal": "vit_row_sinusoidal",
            "col_sinusoidal": "vit_col_sinusoidal",
            "multiplicative_sinusoidal": "vit_multiplicative_sinusoidal",
        }[suffix]
    return f"vit_{order}_{suffix}"


def draw_order_heatmap(agg):
    lookup = agg.set_index("model")
    orders = list(ORDER_PREFIXES)
    suffixes = list(PE_SUFFIX_LABELS)
    matrix = np.array([[100 * lookup.loc[order_model_name(o, s), "mean_acc"] for s in suffixes] for o in orders])
    w, h = 2200, 1150
    image = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(image)
    title = get_font(54, True)
    label = get_font(31)
    value_font = get_font(34, True)
    note = get_font(27)
    d.text((w / 2, 45), "Patch order and positional encoding", font=title, fill="#17233A", anchor="ma")
    left, top = 520, 230
    cell_w, cell_h = 310, 175
    vmin, vmax = matrix.min(), matrix.max()
    for j, suffix in enumerate(suffixes):
        d.multiline_text((left + (j + 0.5) * cell_w, top - 35), PE_SUFFIX_LABELS[suffix].replace(" ", "\n"), font=label, fill="#333333", anchor="ms", align="center", spacing=3)
    for i, order in enumerate(orders):
        d.text((left - 25, top + (i + 0.5) * cell_h), ORDER_PREFIXES[order], font=label, fill="#333333", anchor="rm")
        for j in range(len(suffixes)):
            v = matrix[i, j]
            t = (v - vmin) / max(vmax - vmin, 1e-6)
            color = (int(235 - 115 * t), int(242 - 80 * t), int(248 - 35 * t))
            x1, y1 = left + j * cell_w, top + i * cell_h
            d.rectangle((x1, y1, x1 + cell_w, y1 + cell_h), fill=color, outline="#FFFFFF", width=5)
            d.text((x1 + cell_w / 2, y1 + cell_h / 2), f"{v:.2f}%", font=value_font, fill="#17233A", anchor="mm")
    d.text((w / 2, 1025), "Each cell is the mean held-out test accuracy across seeds 42-46.", font=note, fill="#555555", anchor="ma")
    d.text((w / 2, 1075), f"Colour scale spans {vmin:.2f}% to {vmax:.2f}%; compare values, not colour alone.", font=note, fill="#555555", anchor="ma")
    return save_image_pair(image, "patch_order_heatmap")


def load_curve(model, metric="val_acc"):
    curves = []
    for path in sorted((METRICS / model).glob("*_metrics.csv")):
        frame = pd.read_csv(path)
        curves.append(frame[metric].to_numpy(dtype=float))
    common = min(len(c) for c in curves)
    array = np.vstack([c[:common] for c in curves])
    return np.arange(1, common + 1), array.mean(axis=0), array.std(axis=0, ddof=1)


def draw_learning_curve():
    models = [
        "vit_learnable_position",
        "vit_multiplicative_sinusoidal_shifted",
        "vit_normal_col_learnable_multiplicative_sinusoidal",
    ]
    colors = ["#4C78A8", "#F58518", "#54A24B"]
    w, h = 2200, 1350
    image = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(image, "RGBA")
    title = get_font(54, True)
    label = get_font(31)
    tick = get_font(27)
    d.text((w / 2, 45), "Validation accuracy during training", font=title, fill="#17233A", anchor="ma")
    left, right, top, bottom = 190, 2070, 170, 1120
    d.line((left, top, left, bottom), fill="#444444", width=3)
    d.line((left, bottom, right, bottom), fill="#444444", width=3)
    y_min, y_max = 40.0, 82.0
    for yv in np.arange(40, 83, 10):
        y = bottom - (yv - y_min) / (y_max - y_min) * (bottom - top)
        d.line((left, y, right, y), fill="#E2E5E9", width=2)
        d.text((left - 24, y), f"{yv:.0f}%", font=tick, fill="#444444", anchor="rm")
    max_epoch = 0
    for model, color in zip(models, colors):
        xvals, mean, sd = load_curve(model)
        max_epoch = max(max_epoch, int(xvals[-1]))
        points = []
        upper = []
        lower = []
        for e, m, s in zip(xvals, mean, sd):
            x = left + (e - 1) / max(1, 100 - 1) * (right - left)
            ym = bottom - (100 * m - y_min) / (y_max - y_min) * (bottom - top)
            yu = bottom - (100 * (m + s) - y_min) / (y_max - y_min) * (bottom - top)
            yl = bottom - (100 * (m - s) - y_min) / (y_max - y_min) * (bottom - top)
            points.append((x, ym)); upper.append((x, yu)); lower.append((x, yl))
        d.polygon(upper + list(reversed(lower)), fill=color + "35")
        d.line(points, fill=color, width=7)
    for ev in [1, 20, 40, 60, 80, 100]:
        x = left + (ev - 1) / 99 * (right - left)
        d.text((x, bottom + 25), str(ev), font=tick, fill="#444444", anchor="ma")
    d.text(((left + right) / 2, h - 90), "Epoch", font=label, fill="#333333", anchor="ma")
    d.text((left, 120), "Validation accuracy (%)", font=label, fill="#333333", anchor="ls")
    lx, ly = 240, 1200
    for model, color in zip(models, colors):
        d.line((lx, ly, lx + 75, ly), fill=color, width=7)
        d.text((lx + 95, ly), LABELS[model], font=tick, fill="#333333", anchor="lm")
        lx += 640
    d.text((w / 2, 1310), "Solid line = mean; shaded band = +/- 1 SD. Curves stop at the last epoch shared by all five seeds for each model.", font=tick, fill="#555555", anchor="ms")
    return save_image_pair(image, "representative_validation_accuracy")


def generate_figures(agg):
    core = [
        "vit_baseline", "vit_learnable_position", "vit_row_sinusoidal", "vit_col_sinusoidal",
        "vit_additive_sinusoidal", "vit_additive_sinusoidal_shifted",
        "vit_multiplicative_sinusoidal", "vit_multiplicative_sinusoidal_shifted",
        "vit_squared_multiplicative_sinusoidal", "vit_squared_multiplicative_sinusoidal_shifted",
        "vit_radial_sinusoidal",
    ]
    fusion = [
        "vit_row_col_latent_fusion", "vit_row_col_mean_fusion", "vit_row_col_mean_mlp_fusion",
        "vit_row_col_cross_attention_fusion", "vit_row_col_cross_attention_mlp_head_fusion",
    ]
    return {
        "overview": draw_overview(),
        "core": draw_errorbar_chart(agg, core, "core_pe_comparison", "Core positional encoding comparison", x_min=70.0),
        "order": draw_order_heatmap(agg),
        "fusion": draw_errorbar_chart(agg, fusion, "fusion_comparison", "Row-column fusion comparison", x_min=74.0),
        "curve": draw_learning_curve(),
    }


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_outline_level(style, level):
    p_pr = style.element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def add_multilevel_numbering(doc, styles):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for ilvl, (style, text_fmt) in enumerate(zip(styles, ["%1", "%1.%2", "%1.%2.%3"])):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
        num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "decimal"); lvl.append(num_fmt)
        p_style = OxmlElement("w:pStyle"); p_style.set(qn("w:val"), style.style_id); lvl.append(p_style)
        lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), text_fmt); lvl.append(lvl_text)
        suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "space"); lvl.append(suff)
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId"); abs_id.set(qn("w:val"), str(abstract_id)); num.append(abs_id)
    numbering.append(num)
    for ilvl, style in enumerate(styles):
        p_pr = style.element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl"); ilvl_el.set(qn("w:val"), str(ilvl)); num_pr.append(ilvl_el)
        num_id_el = OxmlElement("w:numId"); num_id_el.set(qn("w:val"), str(num_id)); num_pr.append(num_id_el)
        p_pr.append(num_pr)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.5

    chapter = doc.styles.add_style("Chapter Heading", WD_STYLE_TYPE.PARAGRAPH)
    section = doc.styles.add_style("Section Heading", WD_STYLE_TYPE.PARAGRAPH)
    subsection = doc.styles.add_style("Subsection Heading", WD_STYLE_TYPE.PARAGRAPH)
    for style, size, before, after, level in [
        (chapter, 16, 0, 14, 0), (section, 13, 14, 7, 1), (subsection, 11.5, 10, 5, 2)
    ]:
        style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        if level == 0:
            style.paragraph_format.page_break_before = True
        add_outline_level(style, level)
    add_multilevel_numbering(doc, [chapter, section, subsection])

    front = doc.styles.add_style("Front Matter Heading", WD_STYLE_TYPE.PARAGRAPH)
    front.base_style = normal
    front.font.name = "Times New Roman"; front.font.size = Pt(16); front.font.bold = True
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.space_after = Pt(16)
    add_outline_level(front, 0)

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"; caption.font.size = Pt(9.5); caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4); caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    note = doc.styles.add_style("Drafting Note", WD_STYLE_TYPE.PARAGRAPH)
    note.base_style = normal
    note.font.name = "Times New Roman"; note.font.size = Pt(9.5); note.font.italic = True
    note.font.color.rgb = RGBColor.from_string(MUTED)
    note.paragraph_format.left_indent = Cm(0.5); note.paragraph_format.right_indent = Cm(0.5)
    note.paragraph_format.space_before = Pt(5); note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.15
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), LIGHT); note.element.get_or_add_pPr().append(shd)

    equation = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    equation.base_style = normal
    equation.font.name = "Cambria Math"; equation.font.size = Pt(10.5)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(5); equation.paragraph_format.space_after = Pt(7)


def set_page_number_format(section, fmt, start=1):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    pg_num.set(qn("w:start"), str(start))


def add_field(paragraph, instruction, cached=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = cached
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    return run


def configure_footer(section, short_title=""):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = short_title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, " PAGE ", "1")
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)


def configure_section(section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_body(doc, text, style=None):
    for para in textwrap.dedent(text).strip().split("\n\n"):
        p = doc.add_paragraph(style=style)
        p.add_run(" ".join(line.strip() for line in para.splitlines()))
    return p


def add_rq(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label + " "); set_run_font(r, bold=True)
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)


def add_note(doc, text):
    p = doc.add_paragraph(style="Drafting Note")
    r = p.add_run("Drafting note: "); r.bold = True
    p.add_run(text)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph(style="Equation")
    run = p.add_run(text)
    set_run_font(run, name="Cambria Math", size=10.5)
    return p


def add_picture(doc, path, width=Cm(15.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=width)
    alt_text = Path(path).stem.replace("_", " ")
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return p


def add_caption(doc, label, number, text):
    p = doc.add_paragraph(style="Caption")
    r = p.add_run(f"{label} "); r.bold = True
    seq = add_field(p, f" SEQ {label} \\* ARABIC ", str(number))
    seq.bold = True
    r = p.add_run(". "); r.bold = True
    p.add_run(text)
    return p


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = header
        set_cell_shading(cell, PALE_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, size=font_size, bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    return table


def add_front_heading(doc, text):
    return doc.add_paragraph(text, style="Front Matter Heading")


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(1.8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("UNIVERSITY COLLEGE LONDON")
    set_run_font(r, size=17, bold=True, color=ACCENT)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Cm(2.2)
    r = p.add_run("Investigating Positional Encoding in Vision Transformers")
    set_run_font(r, size=22, bold=True, color=INK)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("An Empirical Study of 2D Designs, Patch Ordering, and Row-Column Fusion")
    set_run_font(r, size=16, italic=True, color=MUTED)
    p.paragraph_format.space_after = Cm(2.0)
    for text, bold in [
        ("Yikai Zhao", True), ("Student ID: 25200353", False),
        ("MSc Scientific and Data Intensive Computing", False),
        ("Scientific Computing Individual Research Project", False),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); set_run_font(r, size=12, bold=bold)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Cm(1.3); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Supervisors: Dr Nikos Nikolaou and Dr Antonis Hadjipittas"); set_run_font(r, size=11)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Cm(2.0); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Physics and Astronomy\n2026"); set_run_font(r, size=11)


def build_document(frame, agg, figures):
    lookup = aggregate_lookup(agg)
    baseline = lookup["vit_baseline"]
    learned = lookup["vit_learnable_position"]
    shifted_mult = lookup["vit_multiplicative_sinusoidal_shifted"]
    hybrid = lookup["vit_normal_col_learnable_multiplicative_sinusoidal"]
    fusion_best = max(
        [lookup[m] for m in ["vit_row_col_latent_fusion", "vit_row_col_mean_fusion", "vit_row_col_mean_mlp_fusion", "vit_row_col_cross_attention_fusion", "vit_row_col_cross_attention_mlp_head_fusion"]],
        key=lambda r: r.mean_acc,
    )

    doc = Document()
    set_styles(doc)
    for section in doc.sections:
        configure_section(section)
    doc.core_properties.title = "Investigating Positional Encoding in Vision Transformers"
    doc.core_properties.subject = "Structured dissertation draft v0.1"
    doc.core_properties.author = "Yikai Zhao"
    doc.core_properties.keywords = "Vision Transformer; positional encoding; patch ordering; CIFAR-10"
    set_update_fields_on_open(doc)

    cover = doc.sections[0]
    cover.header.is_linked_to_previous = False; cover.footer.is_linked_to_previous = False
    cover.header.paragraphs[0].text = ""; cover.footer.paragraphs[0].text = ""
    add_cover(doc)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(front); configure_footer(front, "Dissertation draft v0.1")
    set_page_number_format(front, "lowerRoman", 1)

    add_front_heading(doc, "Declaration")
    add_body(doc, """
        I, Yikai Zhao, confirm that the work presented in this dissertation is my own. Where information has been derived from other sources, this has been indicated in the dissertation. The implementation and experiment materials associated with this project are maintained at https://github.com/RichardKkkkkk/Postgraduate-dissertation.
    """)
    add_note(doc, "Confirm the exact declaration wording required by the current UCL programme handbook before submission, and replace this note with the approved statement.")

    doc.add_page_break()
    add_front_heading(doc, "Abstract")
    abstract = f"""
        Vision Transformers convert an image into a sequence of patch tokens, enabling global self-attention but weakening the direct representation of two-dimensional spatial structure. This dissertation presents a controlled empirical study of positional encoding in a compact Vision Transformer trained on CIFAR-10. The study compares learned absolute embeddings with fixed row-, column-, additive-, multiplicative-, shifted-frequency-, squared-, and radial sinusoidal designs. It further examines four patch traversal orders, a learnable-plus-fixed hybrid encoding, and five row-column feature-fusion strategies. All configurations use the same 32 x 32 input resolution, 4 x 4 patches, 128-dimensional embeddings, four Transformer blocks, four attention heads, a fixed train/validation/test split, and training seeds 42-46. A checkpoint is selected from validation accuracy and evaluated once on the held-out test set.

        The results show that a model without positional encoding performs substantially worse ({acc_text(baseline)}) than the learned absolute baseline ({acc_text(learned)}), confirming that positional information remains important in this setting. Among fixed encodings, shifted multiplicative sinusoidal encoding provides the strongest and most stable result ({acc_text(shifted_mult)}). The hybrid learned-plus-multiplicative model achieves the highest numerical mean ({acc_text(hybrid)}), but its advantage over learned absolute encoding is only {pct(hybrid.mean_acc - learned.mean_acc):.3f} percentage points and is small relative to seed-to-seed variation. Learned embeddings are comparatively insensitive to patch order, whereas fixed encodings show larger order-dependent changes. The best fusion model reaches {acc_text(fusion_best)} but does not exceed the strongest single-encoder or hybrid configurations. These findings support cautious conclusions: positional encoding design affects both accuracy and stability, while more complex spatial fusion is not automatically beneficial under a controlled small-scale ViT protocol.
    """
    add_body(doc, abstract)

    doc.add_page_break()
    add_front_heading(doc, "Acknowledgements")
    add_body(doc, """
        I would like to thank my supervisors, Dr Nikos Nikolaou and Dr Antonis Hadjipittas, for their guidance throughout the project. I am also grateful to the UCL Scientific Computing community for the discussions and computational support that helped shape the experimental work.
    """)
    add_note(doc, "Personalise this section before submission and confirm whether any computing facilities, funding, collaborators, or data providers require formal acknowledgement.")

    doc.add_page_break()
    add_front_heading(doc, "Table of Contents")
    p = doc.add_paragraph(); add_field(p, ' TOC \\o "1-3" \\h \\z \\u ', "Open in Word and update this field (Ctrl+A, F9).")
    doc.add_page_break()
    add_front_heading(doc, "List of Figures")
    add_note(doc, "Experimental figures are intentionally omitted from draft v0.1. Generate this list after the agreed figures are inserted in a later revision.")
    doc.add_page_break()
    add_front_heading(doc, "List of Tables")
    p = doc.add_paragraph(); add_field(p, ' TOC \\h \\z \\c "Table" ', "Open in Word and update this field (Ctrl+A, F9).")

    main = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(main); configure_footer(main, "Investigating Positional Encoding in Vision Transformers")
    set_page_number_format(main, "decimal", 1)

    doc.add_paragraph("Introduction", style="Chapter Heading")
    doc.add_paragraph("Background", style="Section Heading")
    add_body(doc, """
        The Transformer was introduced as an attention-based alternative to recurrent and convolutional sequence models [1]. Its central operation, self-attention, allows each token to aggregate information from every other token in a sequence. This global interaction is attractive for computer vision because image classification depends on both local appearance and relationships between distant regions. The Vision Transformer (ViT) applies this idea by dividing an image into fixed-size patches, projecting each patch to a token embedding, adding a class token and positional information, and processing the sequence with Transformer encoder blocks [2].

        Unlike text, an image is naturally arranged on a two-dimensional grid. Flattening an H x W patch grid into a one-dimensional sequence does not remove the image content, but the content-only self-attention calculation has no intrinsic way to distinguish two token permutations. Positional encoding therefore supplies information about where each patch originated. The original ViT used learned absolute embeddings [2], while the original Transformer used fixed sinusoidal functions [1]. Later work explored relative [5], image-specific relative [6], conditional [7], rotary [8], and hierarchical spatial mechanisms [9]. This variety suggests that positional representation is a design choice rather than a solved implementation detail.
    """)

    doc.add_paragraph("Motivation", style="Section Heading")
    add_body(doc, """
        This project began from a structural question: if an image is converted to a sequence, how should row and column coordinates be represented, combined, and aligned with the token traversal order? Standard learned embeddings can adapt freely to a fixed grid, but their flexibility makes it difficult to infer which spatial assumptions are useful. Fixed encodings expose those assumptions more directly. For example, separate row and column functions preserve axis identity; addition treats the axes symmetrically; element-wise multiplication introduces interactions between them; and different traversal orders alter the relationship between sequence index and image neighbourhood.

        The aim is not to claim a universally optimal positional encoding. The experiments instead isolate a set of concrete design choices under a common architecture, data split, optimiser, and model-selection protocol. This controlled comparison helps distinguish robust effects from changes that are numerically small, unstable across seeds, or confounded by additional model capacity.
    """)
    add_note(doc, "Add an overall study-design diagram here in a later revision, showing the shared CIFAR-10 split and base ViT followed by the positional-encoding, patch-order, hybrid, and fusion experiment branches.")

    doc.add_paragraph("Research Questions and Objectives", style="Section Heading")
    add_rq(doc, "RQ1.", "How do fixed and learned positional encodings affect test accuracy and stability in a compact Vision Transformer trained from scratch on CIFAR-10?")
    add_rq(doc, "RQ2.", "Do explicit two-dimensional combinations of row and column signals provide a consistent benefit over single-axis or learned absolute encodings?")
    add_rq(doc, "RQ3.", "How does the traversal order of an 8 x 8 patch grid interact with positional encoding design?")
    add_rq(doc, "RQ4.", "Can a learned-plus-fixed hybrid or row-column feature fusion improve performance without introducing misleading capacity comparisons?")
    add_body(doc, """
        To answer these questions, the project implements a shared ViT backbone, evaluates 32 configurations across five training seeds, and reports the mean and sample standard deviation of final held-out metrics. The objectives are to build reproducible model variants, maintain an evaluation protocol in which validation data select the checkpoint, and interpret numerical differences in relation to seed variation and architectural complexity.
    """)

    doc.add_paragraph("Contributions", style="Section Heading")
    add_body(doc, """
        The dissertation makes four practical contributions. First, it provides a controlled implementation of multiple one- and two-dimensional positional encodings within the same compact ViT. Second, it tests how row-major, column-major, and two serpentine traversal orders interact with several encoding families. Third, it studies both a learnable-plus-fixed hybrid and multiple row-column fusion mechanisms, including bidirectional cross-attention. Fourth, it evaluates all final configurations with five seeds and a validation-selected, single-use holdout-test protocol, allowing the discussion to focus on reproducible tendencies rather than a single favourable run.
    """)

    doc.add_paragraph("Dissertation Structure", style="Section Heading")
    add_body(doc, """
        Chapter 2 reviews positional information in Transformers and ViTs, with particular attention to two-dimensional structure, patch ordering, and multi-branch fusion. Chapter 3 defines the dataset, backbone, positional encodings, traversal orders, fusion models, training protocol, and statistical summaries. Chapter 4 presents the results in the same order as the research questions. Chapter 5 discusses the interpretation, limitations, and future work. Chapter 6 concludes by answering the research questions and identifying the most reliable findings.
    """)

    doc.add_paragraph("Literature Review", style="Chapter Heading")
    doc.add_paragraph("Transformers and Vision Transformers", style="Section Heading")
    add_body(doc, """
        A Transformer encoder alternates multi-head self-attention and position-wise feed-forward layers, with residual connections and normalisation [1]. Given input tokens X, attention forms query, key, and value projections and computes softmax(QK^T / sqrt(d_k))V. Because the operation depends on token content and pairwise dot products, a consistent permutation of the inputs produces a corresponding permutation of the outputs unless position is injected elsewhere. This property supports flexible set-like processing but conflicts with tasks in which order or geometry carries meaning.

        ViT maps non-overlapping image patches to tokens and uses a learned class token for classification [2]. Its success demonstrated that a convolution-free Transformer can be competitive when trained at scale, while DeiT later showed that careful training and distillation can improve data efficiency [4]. In small-data, from-scratch settings, however, the model cannot rely on large-scale pretraining to learn spatial regularities. Positional design and training protocol may therefore have a more visible influence than in heavily pretrained models.
    """)

    doc.add_paragraph("Absolute and Sinusoidal Position", style="Section Heading")
    add_body(doc, """
        Absolute positional encoding associates each token location with a vector. Learned absolute embeddings allocate one trainable vector per position, allowing the optimiser to adapt the geometry to a fixed patch grid. Fixed sinusoidal encoding instead represents a scalar position with sine and cosine functions at multiple wavelengths [1]. The deterministic construction adds no trainable positional parameters and provides a smooth multi-frequency signal, but a one-dimensional sequence index does not uniquely express the two axes of an image.

        The original ViT reported that several basic position variants were similar at large scale and adopted a learned one-dimensional embedding [2]. That choice does not imply that token order is unimportant. Rather, the learned table can associate each sequence slot with a spatial role when the grid size and traversal convention are fixed. The present study uses this learned embedding as the principal reference because it is both standard and expressive.
    """)

    doc.add_paragraph("Relative, Conditional, and Rotary Encoding", style="Section Heading")
    add_body(doc, """
        Relative position representations modify attention using a relationship between query and key locations instead of, or in addition to, adding an absolute vector to each token [5]. Image-specific relative position encoding can preserve direction and distance on a two-dimensional grid [6]. Conditional positional encoding generates position-dependent information from local token neighbourhoods, which can improve flexibility across image sizes [7]. Rotary position embedding represents relative offsets through rotations of query and key features [8]. These methods broaden the design space, but they also alter the attention computation or introduce local operators. The current experiments deliberately retain the same encoder blocks and manipulate the input positional signal, making the comparison narrower and easier to control.
    """)

    doc.add_paragraph("Two-Dimensional Spatial Structure", style="Section Heading")
    add_body(doc, """
        A patch on an image grid has row coordinate r and column coordinate c. Separate axis encodings preserve this factorisation; a combined encoding must decide how information from the two axes shares the embedding dimensions. Addition is a simple symmetric combination, while concatenation requires dividing or expanding the representation. Element-wise multiplication introduces cross-axis interactions without changing dimensionality. Radial distance collapses the coordinate pair to a scalar and therefore preserves distance from the origin but loses directional identity. No combination is universally preferable: each construction imposes a different spatial inductive bias.

        Hierarchical designs such as Swin Transformer constrain attention to windows and shift those windows across layers, providing spatial locality through the architecture itself [9]. Convolutional tokenisation and conditional encodings provide other routes to position-aware features [7]. These approaches motivate the general importance of spatial structure, but the present work isolates positional signals inside a standard global-attention ViT rather than changing the whole backbone.
    """)

    doc.add_paragraph("Patch Ordering and Flattening", style="Section Heading")
    add_body(doc, """
        A two-dimensional grid must be mapped to a token sequence. Row-major and column-major traversal preserve adjacency along one axis at a time. Serpentine traversal reverses every second row or column, avoiding a large jump between the end of one line and the start of the next. Space-filling curves and learned traversal strategies offer more elaborate locality-preserving alternatives. LOOPE, for example, formulates patch order as a learnable component of positional representation [12].

        Patch order has two distinct roles. It changes the order in which content tokens enter the encoder, and it changes how a sequence-indexed positional table or fixed signal is aligned with the original coordinates. For a globally attending Transformer with a freely learned positional vector at every slot, a consistent reordering can often be represented as a relabelling. A fixed encoding has less freedom, so the interaction between traversal and the encoding formula may be stronger. This distinction motivates the factorial subset of the experiments in Chapter 4.
    """)

    doc.add_paragraph("Row-Column Fusion", style="Section Heading")
    add_body(doc, """
        Multi-branch vision models process complementary views and combine their representations. CrossViT uses cross-attention to exchange information between branches operating at different patch scales [10]. The current project adopts the broader idea of branch interaction but uses row-encoded and column-encoded branches at the same patch scale. Simple averaging tests whether complementary latents can be combined without an expressive fusion module. Concatenation followed by an MLP adds capacity, while bidirectional cross-attention allows every token in one branch to attend to tokens from the other.

        Fusion results must be interpreted carefully because a two-encoder model has substantially more representational capacity and computation than a single-encoder baseline. If such a model fails to improve accuracy, the finding is informative. If it improves, a parameter-matched comparison would still be required before attributing the gain specifically to fusion.
    """)

    doc.add_paragraph("Research Gap and Project Positioning", style="Section Heading")
    add_body(doc, """
        The literature offers many positional mechanisms, but comparisons frequently mix changes in architecture, pretraining, data augmentation, image resolution, or model scale. This dissertation occupies a deliberately limited position: it asks what can be learned from a controlled set of positional interventions in one compact ViT and one dataset. The study does not attempt to displace relative, conditional, or rotary methods. Instead, it provides evidence about axis-aware fixed encodings, their alignment with patch traversal, and the value of hybrid and dual-branch extensions under a consistent protocol.
    """)
    add_note(doc, "For the next literature pass, expand the survey of two-dimensional absolute PE and patch traversal with peer-reviewed sources, then decide whether LOOPE should remain as a recent preprint reference or be replaced by a published traversal study.")

    doc.add_paragraph("Methodology", style="Chapter Heading")
    doc.add_paragraph("Dataset and Data Split", style="Section Heading")
    add_body(doc, """
        CIFAR-10 contains 60,000 colour images of size 32 x 32 across ten balanced object classes [3]. The official 50,000-image training set is divided deterministically into 45,000 training images and 5,000 validation images using split seed 42. The official 10,000-image test set is retained as the final holdout set. All model variants use the same split, so paired comparisons across training seeds are not confounded by different examples.

        The data pipeline uses the transformations implemented in the project repository. The first draft records the confirmed split sizes and image resolution; the exact normalisation constants and any stochastic augmentation should be copied verbatim from the final data-loader implementation during the next methods audit.
    """)

    doc.add_paragraph("Base Vision Transformer", style="Section Heading")
    add_body(doc, """
        Each 32 x 32 image is partitioned by a convolutional patch projection with kernel size and stride 4, producing an 8 x 8 grid of 64 patch tokens. Every token has embedding dimension 128. A learnable class token is prepended, after which positional information is added to the token sequence. The encoder contains four pre-normalisation Transformer blocks, each using four attention heads and a 512-dimensional MLP hidden layer. A final layer normalisation and linear classifier map the class token to ten logits.

        The base configuration sets embedding, attention, projection, and MLP dropout to zero. Holding these choices constant focuses the experiment on positional design. It also means that the absolute accuracy should not be read as a state-of-the-art CIFAR-10 result; the backbone is a controlled experimental instrument rather than a heavily regularised production classifier.
    """)

    doc.add_paragraph("Positional Encoding Variants", style="Section Heading")
    add_body(doc, """
        For scalar position p and embedding channel pair i, the sinusoidal function follows the Transformer construction [1]:
    """)
    add_equation(doc, "S(p, 2i) = sin(p / 10000^(2i/d)),     S(p, 2i+1) = cos(p / 10000^(2i/d)).")
    add_body(doc, """
        Let S_r = S(r) and S_c = S(c) denote row and column encodings. The additive and multiplicative two-dimensional variants are defined as follows:
    """)
    add_equation(doc, "E_add(r,c) = S_r + S_c,     E_mult(r,c) = S_r elementwise-times S_c.")
    add_body(doc, """
        The shifted variants use the even-indexed frequency sequence for rows and an offset frequency sequence for columns before applying addition or multiplication. Squared variants apply an element-wise square after multiplication. The radial variant encodes rho = sqrt(r^2 + c^2) with the same sinusoidal function. The class-token positional vector is zero for fixed encodings. Learned absolute PE instead uses a trainable tensor with one vector for the class token and each patch position.
    """)
    pe_rows = [
        ("None", "No positional vector", "Content-only reference"),
        ("Learned absolute", "Trainable vector per sequence slot", "Flexible fixed-grid reference"),
        ("Row / column", "S(r) or S(c)", "Single-axis structure"),
        ("Additive", "S(r) + S(c)", "Symmetric 2D combination"),
        ("Multiplicative", "S(r) elementwise-times S(c)", "Cross-axis interaction"),
        ("Shifted", "Different row/column frequency indices", "Reduces exact axis-frequency overlap"),
        ("Squared", "E_mult(r,c)^2", "Non-negative interaction magnitude"),
        ("Radial", "S(sqrt(r^2+c^2))", "Distance from grid origin"),
    ]
    add_table(doc, ["Family", "Construction", "Intended structural signal"], pe_rows, [1900, 3200, 3920], 8.8)
    add_caption(doc, "Table", 1, "Positional encoding families implemented in the controlled ViT backbone.")

    doc.add_paragraph("Patch Traversal Orders", style="Section Heading")
    add_body(doc, """
        The patch projection initially produces an 8 x 8 row-major grid. Four deterministic index orders are applied before the class token and positional encoding are added. Normal row order traverses each row left to right. Normal column order traverses each column top to bottom. Proper row order is serpentine: even-indexed rows are left to right and odd-indexed rows are reversed. Proper column order applies the equivalent alternation between columns. The term proper is retained from the implementation; serpentine row and serpentine column are used in the prose for clarity.
    """)

    doc.add_paragraph("Hybrid Positional Encoding", style="Section Heading")
    add_body(doc, """
        The hybrid model combines a trainable absolute embedding L with the fixed multiplicative two-dimensional signal M:
    """)
    add_equation(doc, "E_hybrid = L + alpha M,     alpha initialised to 0 and learned during training.")
    add_body(doc, """
        Initialising alpha to zero makes the model begin as a standard learned absolute PE system and allows optimisation to introduce the fixed component only if it is useful. The implemented hybrid uses column-major traversal. Consequently, its comparison with the row-major learned baseline is informative but not a perfectly isolated ablation; Chapter 5 treats this as a limitation.
    """)

    doc.add_paragraph("Row-Column Fusion Architectures", style="Section Heading")
    add_body(doc, """
        Five dual-encoder models process the same image with separate row-sinusoidal and column-sinusoidal ViTs. The first concatenates the two class-token latents and applies an MLP. The second averages the latents before classification. The third applies an MLP after averaging. The fourth performs bidirectional token-level cross-attention between branches and concatenates the resulting class tokens. The fifth uses the same bidirectional cross-attention with a smoother MLP classification head. All fusion models therefore test interaction strategies, but they are not parameter matched to the single-encoder models.
    """)

    doc.add_paragraph("Training and Model Selection", style="Section Heading")
    protocol_rows = [
        ("Training seeds", "42, 43, 44, 45, 46"), ("Split seed", "42"),
        ("Maximum epochs", "100"), ("Batch size", "128"),
        ("Optimiser", "AdamW [11]"), ("Initial learning rate", "3 x 10^-4"),
        ("Weight decay", "0.05"), ("LR schedule", "ReduceLROnPlateau: patience 5, factor 0.5, minimum 10^-6"),
        ("Early stopping", "Validation accuracy; patience 10; minimum change 0.001"),
        ("Final test", "One evaluation of the validation-selected checkpoint"),
    ]
    add_table(doc, ["Item", "Setting"], protocol_rows, [2700, 6320], 9.0)
    add_caption(doc, "Table", 2, "Training, model-selection, and holdout-test protocol shared by all 32 final configurations.")
    add_body(doc, """
        AdamW decouples weight decay from the adaptive gradient update [11]. Training is permitted to continue for at most 100 epochs. Validation accuracy controls both early stopping and checkpoint selection. ReduceLROnPlateau reduces the learning rate when validation performance stalls. Critically, the test set is not evaluated every epoch and does not select the model. After training, the checkpoint selected by validation accuracy is loaded and evaluated once on the test set. Every final summary file records the protocol as selected_checkpoint_only.
    """)

    doc.add_paragraph("Evaluation and Statistical Summary", style="Section Heading")
    add_body(doc, """
        The primary metric is test accuracy at the selected checkpoint. Macro-averaged precision, recall, and F1 are retained as secondary summaries. For each configuration, the dissertation reports the arithmetic mean and sample standard deviation across the five training seeds. Where two configurations share the same seeds and data split, paired per-seed differences are also inspected. With only five seeds, the analysis avoids treating a small numerical advantage as evidence of statistical significance. The emphasis is on effect size, direction consistency, and variation relative to the observed mean difference.
    """)

    doc.add_paragraph("Reproducibility", style="Section Heading")
    add_body(doc, """
        The project repository contains model definitions, the experiment registry, data loaders, seed-sweep tooling, configuration JSON files, per-epoch training and validation metrics, selected-checkpoint summaries, and confusion matrices. The final analysis in this draft reads 160 summary files: 32 models multiplied by five training seeds. The generated aggregate CSV files and thesis figures are stored under the experiment-specific thesis_core report directory, leaving the original run outputs unchanged.
    """)

    doc.add_paragraph("Experiments and Results", style="Chapter Heading")
    doc.add_paragraph("Reporting Policy and Overall Ranking", style="Section Heading")
    add_body(doc, f"""
        All results in this chapter refer to the held-out test set evaluated at the validation-selected checkpoint. The complete ranking is provided in Appendix A. The highest numerical mean is obtained by the hybrid learned-plus-multiplicative model at {acc_text(hybrid)}, followed closely by several learned absolute PE variants. The no-PE reference is last at {acc_text(baseline)}. Because many top means differ by less than one standard deviation, the ranking is used as a descriptive index rather than as proof that neighbouring models are meaningfully different.
    """)
    top_rows = []
    for row in agg.head(8).itertuples(index=False):
        top_rows.append((row.rank, row.label, f"{pct(row.mean_acc):.3f}", f"{pct(row.sd_acc):.3f}", f"{pct(row.mean_f1):.3f}"))
    add_table(doc, ["Rank", "Configuration", "Accuracy mean (%)", "SD (pp)", "Macro-F1 mean (%)"], top_rows, [650, 3870, 1500, 1200, 1800], 8.5)
    add_caption(doc, "Table", 3, "Top eight configurations by mean selected-checkpoint test accuracy across five training seeds.")

    doc.add_paragraph("Core Positional Encoding Comparison", style="Section Heading")
    add_note(doc, "Insert the core positional-encoding comparison figure after the final visual style is agreed. Plot mean selected-checkpoint test accuracy with sample-SD error bars across seeds 42-46.")
    add_body(doc, f"""
        Removing positional encoding reduces mean accuracy from {pct(learned.mean_acc):.3f}% for learned absolute PE to {pct(baseline.mean_acc):.3f}%, a difference of {pct(learned.mean_acc - baseline.mean_acc):.3f} percentage points. This is the largest and clearest effect in the experiment. The result indicates that global self-attention and patch content alone do not recover all of the spatial information required by this compact model.

        Among fixed encodings, shifted multiplicative PE is strongest at {acc_text(shifted_mult)}. Its standard deviation is only {pct(shifted_mult.sd_acc):.3f} percentage points, making it unusually stable relative to most alternatives. However, it remains {pct(learned.mean_acc - shifted_mult.mean_acc):.3f} points below learned absolute PE on the mean. Additive, single-axis, radial, and squared variants do not establish a general advantage. The fixed results therefore support a narrower conclusion: how row and column signals interact matters, and a shifted multiplicative construction is a promising fixed design in this setting, but learned absolute PE remains a strong reference.
    """)

    doc.add_paragraph("Extended Fixed Designs", style="Section Heading")
    add_body(doc, """
        Squaring the multiplicative interaction changes both sign structure and magnitude. The squared variants do not improve on the unsquared shifted multiplicative design, suggesting that preserving signed channel interactions may be useful. Radial encoding also underperforms the strongest axis-aware designs. A likely explanation is information loss: patches at different directions but equal distance from the upper-left origin receive the same scalar radial coordinate before sinusoidal projection. This interpretation is plausible but not directly measured by the classification experiment.
    """)

    doc.add_paragraph("Patch Traversal Order", style="Section Heading")
    add_note(doc, "Insert a patch-order comparison after the final visual style is agreed. A heatmap can show mean held-out accuracy by traversal order and positional-encoding family, accompanied by a compact patch-index diagram.")
    learned_orders = [lookup[order_model_name(o, "learnable_position")] for o in ORDER_PREFIXES]
    learned_range = pct(max(r.mean_acc for r in learned_orders) - min(r.mean_acc for r in learned_orders))
    add_body(doc, f"""
        Learned absolute PE is highly consistent across traversal conventions: its four order means span only {learned_range:.3f} percentage points. Column-major and both serpentine learned variants are all close to the row-major reference. This is compatible with a relabelling interpretation. If the model attends globally and each sequence slot has an independently learned vector, a fixed permutation can be absorbed by learning a corresponding positional table.

        Fixed encodings show a wider spread. In these models, the sequence order is paired with a deterministic coordinate signal, so the model cannot freely redefine every slot. Some serpentine combinations perform worse even though serpentine traversal reduces geometric jumps between consecutive tokens. This indicates that local sequence continuity alone is not sufficient; alignment between the traversal rule and the formula used to assign row and column signals is also important.
    """)
    add_note(doc, "Add a compact patch-index illustration for all four traversal orders in the next revision. The current heatmap communicates the outcome but not the index geometry itself.")

    doc.add_paragraph("Hybrid Learned and Fixed Encoding", style="Section Heading")
    delta = pct(hybrid.mean_acc - learned.mean_acc)
    hybrid_deltas = frame[frame.model == "vit_normal_col_learnable_multiplicative_sinusoidal"].merge(
        frame[frame.model == "vit_learnable_position"][["seed", "test_acc"]], on="seed", suffixes=("_hybrid", "_learned")
    )
    delta_values = 100 * (hybrid_deltas.test_acc_hybrid - hybrid_deltas.test_acc_learned)
    add_body(doc, f"""
        The hybrid model achieves {acc_text(hybrid)}, compared with {acc_text(learned)} for the learned row-major reference. The mean difference is {delta:.3f} percentage points. Across matched seeds, the differences range from {delta_values.min():.3f} to {delta_values.max():.3f} percentage points. This is not evidence of a reliable optimisation gain: the mean difference is small relative to the variability of both models, and the hybrid also changes traversal from row-major to column-major.

        The useful finding is therefore architectural rather than competitive. Starting from alpha = 0, the model can incorporate a deterministic multiplicative spatial component without degrading the mean result. Future work should log the learned alpha value for every seed and compare hybrid and learned models under the same traversal order before claiming that the fixed component contributes independently.
    """)

    doc.add_paragraph("Row-Column Fusion", style="Section Heading")
    add_note(doc, "Insert the five-model row-column fusion comparison after the final visual style is agreed, explicitly noting that the dual-encoder models are not parameter matched to the single-encoder baselines.")
    add_body(doc, f"""
        The best fusion mean is produced by {fusion_best.label.lower()} at {acc_text(fusion_best)}. Bidirectional cross-attention variants outperform the simpler mean and concatenation alternatives, which suggests that token-level interaction is more effective than combining only final class-token latents. Nevertheless, even the best fusion model remains {pct(hybrid.mean_acc - fusion_best.mean_acc):.3f} percentage points below the hybrid model and {pct(learned.mean_acc - fusion_best.mean_acc):.3f} points below learned absolute PE.

        This outcome is notable because fusion models use two encoders. Increased capacity and computation are not sufficient to guarantee better generalisation. It is possible that separate row and column streams learn redundant information, that optimisation is harder, or that late interaction cannot recover the advantage of a single flexible positional table. Parameter counts and compute should be reported explicitly in a later revision before drawing stronger efficiency conclusions.
    """)

    doc.add_paragraph("Training Behaviour", style="Section Heading")
    add_note(doc, "Insert representative validation-accuracy curves in a later revision. Use five-seed means and sample-SD bands for learned absolute PE, shifted multiplicative PE, and the hybrid model, truncated to the last epoch shared by all seeds of each model.")
    add_body(doc, """
        The representative validation trajectories show that the three strong configurations improve rapidly during early training and then approach similar plateaus. The shifted multiplicative model has a narrow cross-seed band, consistent with its small final test standard deviation. The curves are descriptive rather than a second model-selection mechanism. Early stopping occurs at different epochs across seeds, so each aggregate curve is truncated at the last epoch shared by all five runs of that model.
    """)

    doc.add_paragraph("Summary of Findings", style="Section Heading")
    add_body(doc, """
        The results answer the experimental questions at three levels of confidence. First, positional information is clearly valuable: the no-PE baseline is substantially worse than all leading positional models. Second, learned absolute PE is a strong and order-robust reference, while shifted multiplicative encoding is the most promising fixed construction. Third, hybrid and fusion extensions produce informative numerical differences but do not support broad optimisation claims. The hybrid is only marginally above learned PE on the mean, and the more expensive fusion models remain below the leading single-encoder configurations.
    """)

    doc.add_paragraph("Analysis and Discussion", style="Chapter Heading")
    doc.add_paragraph("Why Positional Information Matters", style="Section Heading")
    add_body(doc, """
        The roughly seven-point gap between learned PE and the no-PE baseline is consistent with the permutation-equivariant nature of content-only self-attention. Patch embeddings contain local appearance, and boundaries or textures may correlate with location, but those cues do not uniquely identify where a patch belongs. Positional encoding gives the classifier a stable frame in which relationships such as above, below, left, and right can be represented. The magnitude of the gap is specific to this architecture and training setup; it should not be generalised to all ViTs or all datasets.
    """)

    doc.add_paragraph("Fixed versus Learned Encoding", style="Section Heading")
    add_body(doc, """
        Learned absolute embeddings perform strongly because they can adapt every slot to the data and fixed grid. This flexibility is also a limitation: they offer no explicit guarantee of smoothness, relative geometry, or transfer to unseen resolutions. Fixed sinusoidal designs have the opposite profile. They add no trainable positional table and make their structural assumptions inspectable, but a poor formula can constrain the model. The shifted multiplicative result suggests that separating axis frequency patterns before element-wise interaction is more effective than using identical row and column frequency sequences.

        The experiment does not reveal the mechanism directly. Visualising pairwise distances between position vectors, measuring attention displacement, or evaluating resolution transfer would help determine whether shifted multiplicative PE preserves more useful directional diversity. These analyses are suitable follow-up work but are not required to report the current classification evidence honestly.
    """)

    doc.add_paragraph("Patch Order and Permutation Reasoning", style="Section Heading")
    add_body(doc, """
        The stability of learned PE across patch orders is theoretically intuitive. Reordering content tokens and assigning an independently learned vector to each reordered slot does not change the expressive capacity of a global self-attention encoder; it changes the labels attached to the positions. Optimisation noise can still produce small differences, but no large systematic effect is expected.

        Fixed encodings break this equivalence because a deterministic function connects sequence slots to coordinate values. A traversal that looks geometrically smooth can still interact poorly with a row-only, column-only, or multiplicative signal. The heatmap therefore argues against treating patch order as an isolated preprocessing decision. It is part of the positional representation and should be reported alongside the encoding formula.
    """)

    doc.add_paragraph("Interpreting the Hybrid Result", style="Section Heading")
    add_body(doc, """
        The hybrid model is the numerical leader, but the correct interpretation is that it matches learned absolute PE while retaining access to a fixed two-dimensional signal. Its 0.06-point mean advantage is much smaller than the cross-seed standard deviations. The result does not justify the word optimised in the dissertation title or abstract. A stronger claim would require an order-matched ablation, logging of alpha, additional datasets, more seeds, and an uncertainty analysis designed before inspecting the outcomes.
    """)

    doc.add_paragraph("Fusion and Capacity Trade-offs", style="Section Heading")
    add_body(doc, """
        Cross-attention is the strongest fusion strategy among those tested, aligning with the general role of cross-attention as an information-exchange mechanism in multi-branch ViTs [10]. However, the fusion study also provides a negative result: duplicating the encoder and adding a richer fusion mechanism does not outperform a single learned positional table. This discourages complexity for its own sake. Any future fusion architecture should be compared at matched parameter count or compute budget and should clarify whether branches receive genuinely complementary information.
    """)

    doc.add_paragraph("Limitations", style="Section Heading")
    add_body(doc, """
        The evidence is limited to CIFAR-10 and a small ViT trained from scratch. The 8 x 8 grid is convenient for controlled experiments but may not reflect behaviour at higher resolution or under large-scale pretraining. Five seeds permit a useful variability estimate but provide limited power for formal hypothesis testing. The study evaluates a selected set of hand-designed encodings rather than an exhaustive search. Some comparisons change more than one factor, most notably traversal in the hybrid model and capacity in the fusion models. Dropout is fixed at zero, and the training recipe is not tuned separately for each architecture.

        The final results are classification outcomes. They do not directly measure positional information retained in attention maps or embeddings. The results also do not demonstrate interpolation to new grid sizes. These boundaries should remain explicit when presenting the findings.
    """)

    doc.add_paragraph("Future Work", style="Section Heading")
    add_body(doc, """
        The first priority is a clean hybrid ablation under identical row-major and column-major orders, including the learned alpha values. The second is a parameter-matched fusion study that reduces branch width or depth to hold total capacity constant. A third direction is evaluation on a larger image dataset or a higher-resolution variant to test whether fixed encodings transfer across grids. Additional analysis could compare attention-distance distributions, positional-vector similarity, robustness to patch permutations, and resolution interpolation.

        Relative, conditional, and rotary encodings provide natural external baselines [5]-[8]. Patch order could also be learned or selected using a locality objective rather than enumerated manually [12]. CADB remains outside the main claims of this draft. If a complete and methodologically comparable CADB experiment is later approved, it should be introduced as a separate appendix study rather than mixed into the CIFAR-10 evidence.
    """)

    doc.add_paragraph("Conclusion", style="Chapter Heading")
    add_body(doc, f"""
        This dissertation investigates how positional encoding, patch traversal, and row-column fusion affect a compact Vision Transformer under a controlled CIFAR-10 protocol. Across 32 configurations and five training seeds, the clearest conclusion is that positional information is essential for this model: the no-PE baseline achieves {pct(baseline.mean_acc):.3f}% mean test accuracy, compared with {pct(learned.mean_acc):.3f}% for learned absolute PE.

        The fixed encoding experiments show that two-dimensional construction matters. Shifted multiplicative sinusoidal encoding is the strongest fixed variant at {pct(shifted_mult.mean_acc):.3f}% and has the lowest variability among the leading fixed models. Learned absolute PE is comparatively insensitive to patch traversal, while fixed encodings show greater interaction with row-major, column-major, and serpentine orders. The hybrid learned-plus-multiplicative model has the highest numerical mean at {pct(hybrid.mean_acc):.3f}%, but its small advantage over learned PE does not support an optimisation claim. Cross-attention is the strongest of the tested fusion mechanisms, yet no fusion model exceeds the best single-encoder or hybrid result.

        The main contribution is therefore a reproducible empirical map of design choices rather than a declaration of a universally best positional encoding. The evidence favours learned absolute PE as a strong default, identifies shifted multiplicative encoding as a stable fixed alternative, and shows that traversal and fusion must be evaluated together with their representational and capacity assumptions.
    """)

    doc.add_paragraph("References", style="Chapter Heading")
    for idx, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(f"[{idx}] "); set_run_font(r, size=9.5)
        r = p.add_run(ref); set_run_font(r, size=9.5)

    doc.add_page_break()
    doc.add_paragraph("Appendix A: Complete Model Results", style="Front Matter Heading")
    add_body(doc, """
        This appendix lists every final configuration. Accuracy and macro-F1 are computed at the validation-selected checkpoint and summarised across training seeds 42-46 using the sample standard deviation.
    """)
    full_rows = []
    for row in agg.itertuples(index=False):
        full_rows.append((row.rank, row.label, f"{pct(row.mean_acc):.3f}", f"{pct(row.sd_acc):.3f}", f"{pct(row.mean_f1):.3f}"))
    # Keep the long appendix table's title with the table instead of allowing a
    # trailing caption to spill onto an otherwise empty page.
    add_caption(doc, "Table", 4, "Complete ranking of all 32 final model configurations.")
    add_table(doc, ["Rank", "Configuration", "Accuracy mean (%)", "SD (pp)", "Macro-F1 mean (%)"], full_rows, [650, 3870, 1500, 1200, 1800], 8.2)

    doc.add_page_break()
    doc.add_paragraph("Appendix B: Supplementary Diagnostics", style="Front Matter Heading")
    add_body(doc, """
        The experiment directory contains per-seed learning curves, confusion matrices, and per-class precision, recall, F1, and accuracy figures. These diagnostics should be selected for the final appendix only when they support a specific error-analysis argument; duplicating every generated image would obscure the main findings.
    """)
    add_note(doc, "Choose one representative confusion matrix and one per-class comparison after the discussion chapter is revised. Record the selection rule before choosing the seed to avoid presenting an unusually favourable run.")

    doc.add_page_break()
    doc.add_paragraph("Appendix C: Reserved CADB Study", style="Front Matter Heading")
    add_body(doc, """
        CADB is intentionally excluded from the main dissertation claims in draft v0.1. This appendix is reserved for a later, self-contained study if the dataset, label protocol, evaluation metrics, and experimental completeness are confirmed. Until then, no CADB result should be cited as evidence for the CIFAR-10 conclusions.
    """)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def write_manifest(agg):
    top = agg.iloc[0]
    content = {
        "source_summary_count": 160,
        "model_count": 32,
        "seeds": [42, 43, 44, 45, 46],
        "protocol": "selected_checkpoint_only",
        "top_model": top.model,
        "top_mean_test_acc": float(top.mean_acc),
        "top_sample_sd": float(top.sd_acc),
        "generated_figures": sorted(p.name for p in FIGURES.glob("*.png")),
        "document": str(OUTPUT.relative_to(ROOT)),
    }
    (REPORT / "thesis_core_manifest.json").write_text(json.dumps(content, indent=2), encoding="utf-8")


def main():
    REPORT.mkdir(parents=True, exist_ok=True)
    FIGURES.mkdir(parents=True, exist_ok=True)
    frame, agg = load_results()
    figures = {}
    path = build_document(frame, agg, figures)
    write_manifest(agg)
    print(path)


if __name__ == "__main__":
    main()
