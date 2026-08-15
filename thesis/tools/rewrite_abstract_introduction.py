from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from standardize_math_notation import add_body_text, clear_paragraph_content


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation.docx"
BUILDING = ROOT / "thesis" / ".Yikai_Zhao_MSc_Dissertation.front.tmp.docx"


ABSTRACT = (
    "Vision Transformers (ViTs) represent images as sequences of patch tokens, "
    "making positional information important for recovering the two-dimensional "
    "arrangement that self-attention does not encode by itself. This dissertation "
    "presents a controlled empirical evaluation of positional encoding in a "
    "shared ViT trained from scratch. It compares no positional encoding, "
    "learnable absolute embeddings, fixed row- and column-wise sinusoidal "
    "encodings, additive and element-wise multiplicative combinations, shifted "
    "frequency schedules, and radial and squared extensions. Separate experiments "
    "examine patch-to-position assignment, reduced training data, a CIFAR-100 "
    "extension, and hybrid and dual-branch fusion models. All formal comparisons "
    "use five seeds, validation-selected checkpoints, held-out test accuracy and "
    "loss, 95% t-confidence intervals, same-seed paired differences, and parameter "
    "counts when model capacity changes. On CIFAR-10, learnable absolute positional "
    "encoding achieved the highest core mean test accuracy at 78.60 ± 0.42%, while "
    "shifted multiplicative encoding was the strongest fixed design at "
    "78.08 ± 0.09%; the no-position baseline reached 71.29 ± 0.83%. Shifted "
    "multiplicative encoding exceeded learnable encoding by 2.54 percentage points "
    "in the 1,000-example setting, but the mean ordering reversed with 5,000, "
    "10,000, and full training data. For the four pre-selected models, CIFAR-100 "
    "retained the CIFAR-10 ranking, with learnable encoding first. The best dual-branch "
    "fusion model reached 77.74 ± 0.44% on CIFAR-10 despite using approximately "
    "2.5 times as many trainable parameters as the learnable single-branch model. "
    "These results show that positional information is valuable in the evaluated "
    "ViT, fixed spatial priors can be competitive in a highly data-limited regime, "
    "and added architectural complexity does not necessarily improve classification "
    "performance."
)


INTRO_1_A = (
    "Vision Transformers (ViTs) have become an important alternative to "
    "convolutional neural networks for image classification. A ViT divides an "
    "image into patches and processes their embeddings as a token sequence, "
    "allowing self-attention to model interactions across the full image. This "
    "formulation, however, does not by itself preserve token order or the "
    "two-dimensional location of each patch. Positional information is therefore "
    "required if the model is to distinguish patches by where they originated in "
    "the image grid "
)

INTRO_2_A = (
    "The original ViT addresses this requirement with a learnable absolute "
    "positional embedding for every token position "
)

INTRO_2_B = (
    ". Such embeddings can adapt to the training task, but they represent location "
    "through a table indexed by sequence position rather than an explicit "
    "row-column construction. Fixed two-dimensional encodings offer a complementary "
    "trade-off: they calculate positional vectors directly from patch coordinates "
    "and add no learned positional parameters, but their spatial representation "
    "cannot adapt during training. The relative value of these choices may therefore "
    "depend on the amount of training data, the patch-to-position mapping, and the "
    "difficulty of the classification task."
)

INTRO_3 = (
    "This dissertation examines positional encoding as a controlled design choice "
    "rather than proposing a new general-purpose Transformer architecture. Using "
    "one shared ViT backbone, it compares no positional encoding and learnable "
    "absolute positional encoding with fixed two-dimensional designs that vary the "
    "combination of row and column information and their sinusoidal frequency "
    "schedules. Further experiments distinguish patch-to-position assignment from "
    "token ordering and evaluate whether hybrid or dual-branch extensions provide "
    "a measurable benefit. The data split, optimisation, checkpoint selection, and "
    "random seeds are kept fixed within each comparison so that the effect of each "
    "design change can be examined directly."
)

INTRO_4 = (
    "The main evaluation uses CIFAR-10 with five training seeds and "
    "validation-selected checkpoints. Reduced-data experiments use 1,000, 5,000, "
    "and 10,000 training examples to examine whether a fixed spatial prior becomes "
    "more useful when labelled data are limited. CIFAR-100 provides a more "
    "fine-grained extension at the same input resolution, allowing the relative "
    "ordering of selected methods to be compared under a harder classification "
    "task. Test accuracy and loss are reported with 95% t-confidence intervals and "
    "same-seed paired differences, while parameter counts accompany hybrid and "
    "fusion comparisons."
)

INTRO_5 = (
    "The results establish three main patterns. In the core CIFAR-10 comparison, "
    "learnable absolute positional encoding achieved the highest mean accuracy "
    "(78.60 ± 0.42%), shifted multiplicative encoding was the strongest fixed "
    "method (78.08 ± 0.09%), and the no-position baseline was lower "
    "(71.29 ± 0.83%). Shifted multiplicative encoding exceeded learnable encoding "
    "only in the 1,000-example setting; the mean ordering reversed as more training "
    "data were used. The four selected methods retained the same ranking on "
    "CIFAR-100, while the best fusion model remained below learnable positional "
    "encoding despite more than doubling the parameter count. These findings "
    "support a cautious conclusion: a fixed two-dimensional encoding can be a useful "
    "deterministic alternative, but learnable positional encoding performs best in "
    "most of the evaluated settings, and additional complexity alone is not a "
    "reliable route to better performance."
)


def field_counts(document: Document) -> dict[str, int]:
    codes = [
        (element.text or "").strip()
        for element in document.element.body.xpath(".//w:instrText")
    ]
    return {
        "zotero": sum(code.startswith("ADDIN ZOTERO") for code in codes),
        "seq": sum(code.startswith("SEQ ") for code in codes),
        "ref": sum(code.startswith("REF ") for code in codes),
    }


def paragraphs_between(document: Document, start_heading: str, end_heading: str):
    paragraphs = document.paragraphs
    starts = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == start_heading
    ]
    ends = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip() == end_heading
    ]
    if len(starts) != 1 or len(ends) != 1 or starts[0] >= ends[0]:
        raise RuntimeError(f"Could not locate section {start_heading!r}")
    return paragraphs[starts[0] + 1 : ends[0]]


def complex_field_blocks(paragraph, instruction_prefix: str):
    blocks = []
    current = []
    depth = 0

    for child in list(paragraph._p):
        field_chars = child.xpath(".//*[local-name()='fldChar']")
        begins = sum(
            element.get(qn("w:fldCharType")) == "begin" for element in field_chars
        )
        ends = sum(
            element.get(qn("w:fldCharType")) == "end" for element in field_chars
        )

        if depth == 0 and begins:
            current = []
        if begins or depth:
            current.append(deepcopy(child))
            depth += begins
            depth -= ends
            if depth == 0 and current:
                instructions = "".join(
                    (element.text or "")
                    for item in current
                    for element in item.xpath(".//*[local-name()='instrText']")
                ).strip()
                if instructions.startswith(instruction_prefix):
                    blocks.append(current)
                current = []
    return blocks


def append_field_block(paragraph, block) -> None:
    for element in block:
        paragraph._p.append(deepcopy(element))


def remove_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is None:
        raise RuntimeError("Paragraph is already detached")
    parent.remove(paragraph._p)


def rewrite(document: Document) -> dict[str, int]:
    expected_fields = field_counts(document)
    abstract_paragraphs = paragraphs_between(document, "Abstract", "1 Introduction")
    introduction_paragraphs = paragraphs_between(
        document, "1 Introduction", "2 Literature Review"
    )

    if len(abstract_paragraphs) not in (1, 2):
        raise RuntimeError(
            f"Expected one abstract paragraph with an optional drafting note, found "
            f"{len(abstract_paragraphs)} paragraphs"
        )
    if len(introduction_paragraphs) not in (5, 6):
        raise RuntimeError(
            f"Expected five introduction paragraphs with an optional placeholder, found "
            f"{len(introduction_paragraphs)} paragraphs"
        )

    citation_blocks = []
    for paragraph in introduction_paragraphs[:2]:
        citation_blocks.extend(complex_field_blocks(paragraph, "ADDIN ZOTERO"))
    if len(citation_blocks) != 3:
        raise RuntimeError(
            f"Expected three reusable Introduction citations, found {len(citation_blocks)}"
        )

    clear_paragraph_content(abstract_paragraphs[0])
    add_body_text(abstract_paragraphs[0], ABSTRACT)
    if len(abstract_paragraphs) == 2:
        remove_paragraph(abstract_paragraphs[1])

    first, second, third, fourth, fifth = introduction_paragraphs[:5]

    clear_paragraph_content(first)
    add_body_text(first, INTRO_1_A)
    append_field_block(first, citation_blocks[0])
    add_body_text(first, ", ")
    append_field_block(first, citation_blocks[1])
    add_body_text(first, ".")

    clear_paragraph_content(second)
    add_body_text(second, INTRO_2_A)
    append_field_block(second, citation_blocks[2])
    add_body_text(second, INTRO_2_B)

    for paragraph, text in (
        (third, INTRO_3),
        (fourth, INTRO_4),
        (fifth, INTRO_5),
    ):
        clear_paragraph_content(paragraph)
        add_body_text(paragraph, text)

    if len(introduction_paragraphs) == 6:
        remove_paragraph(introduction_paragraphs[5])
    return expected_fields


def validate(document: Document, expected_fields: dict[str, int]) -> None:
    abstract_paragraphs = paragraphs_between(document, "Abstract", "1 Introduction")
    introduction_paragraphs = paragraphs_between(
        document, "1 Introduction", "2 Literature Review"
    )

    if len(abstract_paragraphs) != 1:
        raise RuntimeError("Abstract must contain exactly one body paragraph")
    if len(introduction_paragraphs) != 5:
        raise RuntimeError("Introduction must contain exactly five body paragraphs")

    abstract_text = abstract_paragraphs[0].text
    introduction_text = "\n".join(p.text for p in introduction_paragraphs)
    abstract_words = re.findall(r"\b[\w-]+\b", abstract_text)
    introduction_words = re.findall(r"\b[\w-]+\b", introduction_text)
    if not (220 <= len(abstract_words) <= 300):
        raise RuntimeError(f"Unexpected abstract length: {len(abstract_words)} words")
    if not (420 <= len(introduction_words) <= 650):
        raise RuntimeError(
            f"Unexpected introduction length: {len(introduction_words)} words"
        )

    required_abstract = (
        "78.60 ± 0.42%",
        "78.08 ± 0.09%",
        "71.29 ± 0.83%",
        "2.54 percentage points",
        "77.74 ± 0.44%",
    )
    required_intro = (
        "78.60 ± 0.42%",
        "78.08 ± 0.09%",
        "71.29 ± 0.83%",
        "1,000-example setting",
        "same ranking on CIFAR-100",
    )
    if any(value not in abstract_text for value in required_abstract):
        raise RuntimeError("A required Abstract result is missing")
    if any(value not in introduction_text for value in required_intro):
        raise RuntimeError("A required Introduction result is missing")

    combined = f"{abstract_text}\n{introduction_text}".lower()
    for forbidden in (
        "drafting note",
        "补充result",
        "significantly improves",
        "optimises",
        "universally superior",
    ):
        if forbidden in combined:
            raise RuntimeError(f"Forbidden or placeholder wording remains: {forbidden}")

    abstract_fields = [
        (element.text or "").strip()
        for paragraph in abstract_paragraphs
        for element in paragraph._p.xpath(".//w:instrText")
    ]
    intro_fields = [
        (element.text or "").strip()
        for paragraph in introduction_paragraphs
        for element in paragraph._p.xpath(".//w:instrText")
        if (element.text or "").strip().startswith("ADDIN ZOTERO")
    ]
    if abstract_fields:
        raise RuntimeError("Abstract must not contain citations")
    if len(intro_fields) != 3:
        raise RuntimeError("Introduction citation count changed")

    if field_counts(document) != expected_fields:
        raise RuntimeError("Document field counts changed")
    if len(document.inline_shapes) != 10 or len(document.tables) != 11:
        raise RuntimeError("Figure or table counts changed")

    bookmarks = {
        element.get(qn("w:name"))
        for element in document.element.body.xpath(".//w:bookmarkStart")
    }
    if not {f"eq{number}" for number in range(1, 15)}.issubset(bookmarks):
        raise RuntimeError("An equation bookmark was lost")

    print(f"Abstract words: {len(abstract_words)}")
    print(f"Introduction words: {len(introduction_words)}")


def main() -> None:
    document = Document(str(DOCX))
    expected_fields = rewrite(document)
    validate(document, expected_fields)

    document.save(str(BUILDING))
    reopened = Document(str(BUILDING))
    validate(reopened, expected_fields)
    BUILDING.replace(DOCX)

    with ZipFile(DOCX) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("The saved DOCX package failed its ZIP integrity test")

    print(f"Updated: {DOCX}")
    print(f"Fields preserved: {expected_fields}")


if __name__ == "__main__":
    main()
