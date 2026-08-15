from pathlib import Path
from docx import Document
from docx.shared import RGBColor


DOCX_PATH = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
OLD_START = "The experiments show that positional information is important for the ViT used in this project."
NEW_TEXT = (
    "The results show that positional information is important for the ViT used in this project. "
    "On the full CIFAR-10 training set, learnable absolute PE achieved a mean test accuracy of 78.60%, "
    "compared with 78.08% for shifted multiplicative PE and 71.29% without PE. With only 1,000 training "
    "examples, shifted multiplicative PE reached 40.84%, while learnable PE achieved 37.47%. Learnable PE "
    "returned to the higher mean at 10,000 examples and with the full training set, and it also ranked first "
    "on CIFAR-100. The dual-branch models added substantial capacity but did not improve on the learned "
    "single-branch model."
)


document = Document(DOCX_PATH)
old_matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(OLD_START)]
new_matches = [paragraph for paragraph in document.paragraphs if paragraph.text == NEW_TEXT]

if len(old_matches) == 1 and not new_matches:
    paragraph = old_matches[0]
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    run = paragraph.add_run(NEW_TEXT)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)

    temporary_path = DOCX_PATH.with_suffix(".tmp.docx")
    document.save(temporary_path)
    temporary_path.replace(DOCX_PATH)
    document = Document(DOCX_PATH)
elif len(new_matches) != 1:
    raise RuntimeError(
        f"Expected one old or new results paragraph, found old={len(old_matches)}, new={len(new_matches)}"
    )

verified_matches = [paragraph for paragraph in document.paragraphs if paragraph.text == NEW_TEXT]
verified_runs = verified_matches[0].runs if verified_matches else []
font_names = sorted({run.font.name or "(inherited)" for run in verified_runs})
font_colours = sorted(
    {str(run.font.color.rgb) if run.font.color.rgb is not None else "(inherited)" for run in verified_runs}
)
equation_count = len(document.element.xpath(".//*[local-name()='oMathPara']"))
print(f"Updated paragraph matches: {len(verified_matches)}")
print(f"Result paragraph fonts: {font_names}")
print(f"Result paragraph colours: {font_colours}")
print(f"Paragraphs: {len(document.paragraphs)}")
print(f"Tables: {len(document.tables)}")
print(f"Inline images: {len(document.inline_shapes)}")
print(f"Displayed equations: {equation_count}")
