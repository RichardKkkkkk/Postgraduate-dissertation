from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
doc = Document(path)

replacements = {
    "The study is organised around three research questions.": (
        "This study examines the role of positional information in the ViT used in this project. "
        "Learned absolute PE is compared with several fixed two-dimensional designs under the same "
        "architecture and training protocol. The experiments also investigate how this comparison "
        "changes when less training data are available and whether the main findings remain consistent "
        "on CIFAR-100."
    ),
    "5.1 RQ1: The Role of Positional Encoding": "5.1 The Role of Positional Encoding",
    "5.2 RQ2: Comparison of Fixed Designs": "5.2 Comparison of Fixed Designs",
    "5.3 RQ3: Data Availability and Cross-Dataset Consistency": (
        "5.3 Data Availability and Cross-Dataset Consistency"
    ),
    "Chapter 4 reported the experimental outcomes.": (
        "Chapter 4 reported the experimental outcomes. This chapter draws together the main findings, "
        "considers what they mean for the choice of positional encoding, and discusses the limits of the "
        "evidence."
    ),
    "The shifted comparisons require a narrower conclusion.": (
        "The shifted comparisons require a narrower conclusion. Shifted additive PE changed mean accuracy "
        "by only 0.17 points, and shifted multiplicative PE changed it by 0.55 points. Both paired confidence "
        "intervals included zero. The full shifted multiplicative configuration can therefore be described "
        "as the strongest fixed "
        "design observed in this study, but the results do not isolate the frequency shift as the cause of "
        "that ranking. The comparison therefore supports a conclusion about the complete fixed designs, "
        "not a general claim that shifting the frequency schedule provides a reliable improvement."
    ),
    "The clearest result is that positional information matters": (
        "The clearest result is that positional information matters when the full training set is used. On "
        "CIFAR-10, learnable PE improved mean test accuracy over no PE by 7.31 percentage points. The paired "
        "95% confidence interval ranged from 6.51 to 8.12 points, and the difference was positive for all five "
        "seeds. A similar gap appeared on CIFAR-100, where learnable PE achieved 50.40% mean test accuracy "
        "compared with 43.08% for no PE. Removing positional information therefore caused a clear loss of "
        "classification performance under full-data training in both datasets."
    ),
    "Training-set size changed the ordering between the two leading methods.": (
        "Training-set size changed the ordering between the two leading methods. Shifted multiplicative PE "
        "led learnable PE at 1,000 and 5,000 examples, while the learned method had the higher mean at "
        "10,000 examples and with the full training set. The four pre-selected methods also retained the "
        "same accuracy ranking on CIFAR-100. Together, these results show that the relative performance of "
        "learned and fixed PE depends on the amount of training data in the evaluated settings. They do not "
        "identify a universal data threshold or establish broad cross-domain generalisation."
    ),
    "The experiments answer the three research questions within this setting.": (
        "Within this setting, the experiments provide three main findings. Positional information was "
        "important in full-data training, and learnable absolute PE achieved the highest core mean test "
        "accuracy. Shifted multiplicative PE was the strongest fixed design. It led learnable PE at 1,000 "
        "and 5,000 training examples, but the learned method had the higher mean at 10,000 examples and with "
        "the full training set. The selected methods kept the same test-accuracy order on CIFAR-100, while "
        "the dual-branch models did not improve on the learned single-branch reference."
    ),
    "Chapter 2 reviews previous work on positional information": (
        "Chapter 2 reviews previous work on positional information in Transformers and explains the research "
        "gap addressed by this project. Chapter 3 describes the datasets, model architecture, positional "
        "encoding methods, fusion extensions, and evaluation process. Chapter 4 presents the experiments and "
        "results. Chapter 5 discusses the findings and their limitations. Chapter 6 summarises the main "
        "conclusions and suggests possible directions for future work."
    ),
}


def replace(paragraph, text):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Arial")


seen = set()
for paragraph in doc.paragraphs:
    for prefix, text in replacements.items():
        if paragraph.text.strip().startswith(prefix):
            replace(paragraph, text)
            seen.add(prefix)
            break

missing = set(replacements) - seen
if missing:
    raise SystemExit(f"Targets not found: {sorted(missing)}")

tmp = path.with_suffix(".rq.tmp.docx")
doc.save(tmp)
tmp.replace(path)
print(path)
