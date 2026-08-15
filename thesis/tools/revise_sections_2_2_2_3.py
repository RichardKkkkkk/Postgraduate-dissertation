from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor


DOCX_PATH = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")


def find_paragraph(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}, found {len(matches)}")
    return matches[0]


def citation_elements(paragraph):
    selected = []
    active = False
    for child in paragraph._p:
        field_chars = list(child.iter(qn("w:fldChar")))
        if any(node.get(qn("w:fldCharType")) == "begin" for node in field_chars):
            active = True
        if active:
            selected.append(deepcopy(child))
        if active and any(node.get(qn("w:fldCharType")) == "end" for node in field_chars):
            break
    if not selected:
        raise RuntimeError(f"No citation field found in paragraph: {paragraph.text[:80]!r}")
    return selected


def clear_content(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_run(paragraph, text, *, italic=False, subscript=False, superscript=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.italic = italic
    run.font.subscript = subscript
    run.font.superscript = superscript
    return run


def replace_plain(paragraph, text):
    clear_content(paragraph)
    add_run(paragraph, text)


def replace_with_citation(paragraph, before, citation_source, after="."):
    field = citation_elements(citation_source)
    clear_content(paragraph)
    add_run(paragraph, before)
    for element in field:
        paragraph._p.append(element)
    add_run(paragraph, after)


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


document = Document(DOCX_PATH)
before_counts = (
    len(document.paragraphs),
    len(document.tables),
    len(document.inline_shapes),
    len(document.element.xpath(".//*[local-name()='oMathPara']")),
)

p22_intro = find_paragraph(document, "The Vision Transformer (ViT) applies")
p221_detail = find_paragraph(document, "Each patch has dimensions")
p222_heading = find_paragraph(document, "2.2.2 Patch Embedding and Class Token")
p222_formula = find_paragraph(document, "Each patch is flattened into a vector")
p222_class = find_paragraph(document, "The projected patch embeddings form the input token sequence")
p223_heading = find_paragraph(document, "2.2.3 Learnable Positional Embeddings in the ViT Baseline")
p223_input = find_paragraph(document, "Although a patch embedding represents local visual content")
p223_rows = find_paragraph(document, "Each row of Epos corresponds")
p223_encoder = find_paragraph(document, "The resulting token sequence is processed")
p23_intro = find_paragraph(document, "Spatial representation in Vision Transformers is shaped")
p231_first = find_paragraph(document, "Learnable absolute positional embeddings associate")
p231_second = find_paragraph(document, "Because the positional vectors are learned")
p231_third = find_paragraph(document, "However, this representation does not explicitly encode")
p232_third = find_paragraph(document, "The main advantage of this approach")
p233_first = find_paragraph(document, "Patch tokenisation determines how image pixels are grouped")
p233_second = find_paragraph(document, "Yuan et al. proposed the Tokens-to-Token")
p233_third = find_paragraph(document, "Patch size creates a related design trade-off")

replace_with_citation(
    p22_intro,
    "The Vision Transformer (ViT) adapts the Transformer encoder for image classification by representing an image as a sequence of patch tokens. Figure 3 summarises the process. The image is divided into fixed-size patches, each patch is projected into an embedding, and the sequence is processed by a Transformer encoder. Sections 2.2.1–2.2.3 describe patch formation, patch projection, and the construction of the input sequence ",
    p22_intro,
)

clear_content(p221_detail)
add_run(p221_detail, "Each patch has dimensions ")
add_run(p221_detail, "P", italic=True)
add_run(p221_detail, " × ")
add_run(p221_detail, "P", italic=True)
add_run(p221_detail, " × ")
add_run(p221_detail, "C", italic=True)
add_run(p221_detail, " and contains ")
add_run(p221_detail, "P", italic=True)
add_run(p221_detail, "2", superscript=True)
add_run(p221_detail, "C", italic=True)
add_run(p221_detail, " scalar values. Flattening a patch retains its pixel values but removes its explicit row and column coordinates. The next stage therefore represents patch content and position separately.")

replace_plain(p222_heading, "2.2.2 Patch Embedding")
replace_with_citation(
    p222_class,
    "The same projection is applied to every patch. It gives all patch embeddings the same dimension, which allows them to be processed as one token sequence ",
    p222_class,
)

replace_plain(p223_heading, "2.2.3 Classification Token and Baseline Input")
field = citation_elements(p223_input)
clear_content(p223_input)
add_run(p223_input, "ViT places a learnable classification token at the start of the projected patch sequence. This token is processed with the patch tokens, and its final representation is passed to the classification head. Before the sequence enters the encoder, learnable positional embeddings are added to the classification token and the patch tokens. Equation (5) defines this initial sequence, where ")
add_run(p223_input, "x", italic=True)
add_run(p223_input, "class", italic=True, subscript=True)
add_run(p223_input, " denotes the classification token and ")
add_run(p223_input, "E", italic=True)
add_run(p223_input, "pos", italic=True, subscript=True)
add_run(p223_input, " ∈ ℝ")
add_run(p223_input, "(N+1) × D", superscript=True)
add_run(p223_input, " denotes the positional embedding matrix ")
for element in field:
    p223_input._p.append(element)
add_run(p223_input, ".")
delete_paragraph(p223_rows)
replace_plain(
    p223_encoder,
    "The complete sequence is processed by a stack of Transformer encoder layers. Each layer contains multi-head self-attention, a position-wise multilayer perceptron, layer normalisation, and residual connections. All models in this dissertation keep the patch projection, classification token, encoder, and prediction head unchanged. The positional representation can therefore be compared within the same ViT structure.",
)

replace_plain(
    p23_intro,
    "Once the ViT input has been defined, the main remaining question is how position should be represented. This section compares learnable absolute PE with fixed two-dimensional signals. It also reviews alternative tokenisation methods because they change the local information available before self-attention, even though they are separate from PE.",
)
replace_with_citation(
    p231_first,
    "The baseline in Equation (5) uses learnable absolute PE. Its positional vectors are updated jointly with the model weights, so they can adapt to the training task rather than follow a predefined coordinate function ",
    p231_first,
)
replace_plain(
    p231_second,
    "The term absolute means that each vector identifies one token index. It does not directly encode the distance or direction between two patches. The learned table is also tied to a particular token grid, so changing the input resolution usually requires the positional embeddings to be interpolated.",
)
replace_with_citation(
    p231_third,
    "Dosovitskiy et al. found that learned embeddings could still recover properties of the two-dimensional grid. Nearby patches and patches in the same row or column tended to have similar embeddings ",
    p231_third,
    ". This shows that spatial structure can emerge during training, although it is not built directly into the learned table. Learnable absolute PE is therefore a flexible reference for comparison with fixed coordinate-based signals.",
)
replace_with_citation(
    p232_third,
    "A fixed encoding supplies spatial structure before training, but its coordinate function cannot adapt to the task. In the original ViT study, hand-crafted two-dimensional variants did not improve on the learned embeddings ",
    p232_third,
    ". Fixed and learned encodings therefore offer different trade-offs. One provides an explicit spatial prior, while the other learns the representation from data.",
)
replace_with_citation(
    p233_first,
    "The baseline in Section 2.2 projects each patch independently. This front end is simple and efficient, but it does not combine information from neighbouring patches before the Transformer encoder ",
    p233_first,
    ". Other tokenisation methods alter this stage to introduce local interactions earlier. They should therefore be treated as changes to token formation rather than changes to positional encoding.",
)
replace_with_citation(
    p233_second,
    "Yuan et al. proposed the Tokens-to-Token Vision Transformer (T2T-ViT), which repeatedly combines neighbouring tokens before the main Transformer encoder. This process introduces local structure while gradually reducing the sequence length ",
    p233_second,
    ". It also changes the model front end and adds computation, so it is not directly comparable with a change to input-level PE.",
)
replace_plain(
    p233_third,
    "Patch size creates a related design trade-off. Smaller patches retain finer spatial detail but produce longer token sequences, increasing the memory and computational cost of self-attention. Larger patches shorten the sequence but combine a wider image region within each token, which may remove fine detail before attention is applied. Recent work on patchification scaling reports that smaller patches can improve predictive performance, although they also increase sequence length and computation [4]. Patch size is therefore a separate architectural choice that should be controlled when PE methods are compared.",
)

temporary_path = DOCX_PATH.with_suffix(".tmp.docx")
document.save(temporary_path)
temporary_path.replace(DOCX_PATH)

verified = Document(DOCX_PATH)
after_counts = (
    len(verified.paragraphs),
    len(verified.tables),
    len(verified.inline_shapes),
    len(verified.element.xpath(".//*[local-name()='oMathPara']")),
)
print("Before:", before_counts)
print("After:", after_counts)
print("Updated headings:")
for paragraph in verified.paragraphs:
    if paragraph.text.startswith(("2.2.2 ", "2.2.3 ")):
        print(paragraph.text)
