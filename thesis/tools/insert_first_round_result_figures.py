from __future__ import annotations

from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation.docx"
BUILDING = ROOT / "thesis" / ".Yikai_Zhao_MSc_Dissertation.figures.tmp.docx"

MAIN_FIGURES = (
    ROOT
    / "results"
    / "cifar10_final_vit_models_5seeds"
    / "reports"
    / "thesis_comparison_figures_v2"
    / "figures"
)
ROBUSTNESS_FIGURES = (
    ROOT / "results" / "reports" / "thesis_robustness_figures_v2" / "figures"
)
WORD_FIGURES = ROOT / "results" / "reports" / "thesis_word_figures_v2"

TABLE_CAPTIONS = {
    1: "Dataset configuration and experimental role.",
    2: "Shared ViT configuration.",
    3: "Positional-encoding configurations evaluated in Section 3.4.",
    4: "Hardware and software environment.",
    5: "Experimental reporting conventions.",
    6: "Core PE comparison on CIFAR-10.",
    7: "Base and shifted PE comparison on CIFAR-10.",
    8: "Patch-to-position assignment results on CIFAR-10.",
    9: "Reduced-data PE comparison on CIFAR-10.",
    10: "Selected PE configurations on CIFAR-10 and CIFAR-100.",
    11: "Hybrid, fusion and fixed-PE extensions.",
}


def set_run_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Arial")


def stack_images(top_path: Path, bottom_path: Path, output_path: Path) -> None:
    """Stack two thesis plots vertically without resampling either source."""
    with Image.open(top_path) as top_source, Image.open(bottom_path) as bottom_source:
        top = top_source.convert("RGB")
        bottom = bottom_source.convert("RGB")
        width = max(top.width, bottom.width)
        gap = max(36, round(width * 0.012))
        canvas = Image.new("RGB", (width, top.height + gap + bottom.height), "white")
        canvas.paste(top, ((width - top.width) // 2, 0))
        canvas.paste(bottom, ((width - bottom.width) // 2, top.height + gap))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        canvas.save(output_path, format="PNG", dpi=(300, 300), optimize=True)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def previous_paragraph(paragraph: Paragraph) -> Paragraph | None:
    element = paragraph._p.getprevious()
    while element is not None:
        if element.tag == qn("w:p"):
            return Paragraph(element, paragraph._parent)
        element = element.getprevious()
    return None


def remove_existing_result_figures(document: Document) -> None:
    captions = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if any(text.startswith(f"Figure {number}:") for number in range(4, 11)):
            captions.append(paragraph)

    for caption in captions:
        image_paragraph = previous_paragraph(caption)
        if image_paragraph is not None and image_paragraph._p.xpath(
            ".//w:drawing|.//w:pict"
        ):
            remove_paragraph(image_paragraph)
        remove_paragraph(caption)


def clear_paragraph_content(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def add_text_run(paragraph: Paragraph, text: str, *, caption: bool = False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, 10 if caption else 11, bold=caption, italic=caption)


def add_field_run(
    paragraph: Paragraph,
    instruction: str,
    display_text: str,
    *,
    caption: bool = False,
) -> None:
    begin_run = paragraph.add_run()
    set_run_font(begin_run, 10 if caption else 11, bold=caption, italic=caption)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    set_run_font(
        instruction_run, 10 if caption else 11, bold=caption, italic=caption
    )
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    instruction_run._r.append(instruction_element)

    separate_run = paragraph.add_run()
    set_run_font(separate_run, 10 if caption else 11, bold=caption, italic=caption)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(display_text)
    set_run_font(result_run, 10 if caption else 11, bold=caption, italic=caption)

    end_run = paragraph.add_run()
    set_run_font(end_run, 10 if caption else 11, bold=caption, italic=caption)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_ref(paragraph: Paragraph, bookmark: str, display_number: int) -> None:
    add_field_run(paragraph, f"REF {bookmark} \\h", str(display_number))


def replace_result_paragraph(paragraph: Paragraph, parts: list[object]) -> None:
    clear_paragraph_content(paragraph)
    for part in parts:
        if isinstance(part, str):
            add_text_run(paragraph, part)
        else:
            bookmark, number = part
            add_ref(paragraph, bookmark, number)


def find_paragraph(document: Document, starts_with: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(starts_with)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph starting with {starts_with!r}, found {len(matches)}"
        )
    return matches[0]


def shorten_table_captions(document: Document) -> None:
    for number, title in TABLE_CAPTIONS.items():
        paragraph = find_paragraph(document, f"Table {number}.")
        runs = paragraph.runs
        try:
            suffix_index = next(
                index for index, run in enumerate(runs) if run.text.startswith(". ")
            )
        except StopIteration as error:
            raise RuntimeError(
                f"Could not locate the title run for Table {number}"
            ) from error

        runs[suffix_index].text = f". {title}"
        set_run_font(runs[suffix_index], 10, bold=True, italic=True)
        for run in runs[suffix_index + 1 :]:
            if run.text:
                run.text = ""


def update_reporting_convention(document: Document) -> None:
    paragraph = find_paragraph(document, "All values in this chapter are generated")
    replace_result_paragraph(
        paragraph,
        [
            "All values in this chapter are generated from the validation-selected five-seed summary files. The tables report selected-checkpoint test performance using the frozen protocol described in Chapter 3; values following ± and bracketed ranges are 95% t confidence intervals. Validation curves show the mean across seeds 42-46, with shaded bands representing one sample standard deviation. For each configuration, a curve ends at the last epoch shared by all five runs."
        ],
    )


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    inserted = Paragraph(new_element, paragraph._parent)
    if style:
        inserted.style = style
    return inserted


def next_bookmark_id(document: Document) -> int:
    ids = []
    for element in document.element.body.xpath(".//w:bookmarkStart"):
        value = element.get(qn("w:id"))
        if value is not None and value.isdigit():
            ids.append(int(value))
    return max(ids, default=-1) + 1


def add_caption(
    paragraph: Paragraph,
    figure_number: int,
    caption_text: str,
    bookmark_id: int,
) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = False

    add_text_run(paragraph, "Figure ", caption=True)
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), f"fig{figure_number}")
    paragraph._p.append(bookmark_start)

    instruction = (
        f"SEQ Figure \\r {figure_number}"
        if figure_number == 4
        else "SEQ Figure"
    )
    add_field_run(paragraph, instruction, str(figure_number), caption=True)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(bookmark_end)
    add_text_run(paragraph, f": {caption_text}", caption=True)


def insert_figure(
    anchor: Paragraph,
    image_path: Path,
    figure_number: int,
    caption_text: str,
    alt_text: str,
    bookmark_id: int,
) -> Paragraph:
    figure_paragraph = insert_paragraph_after(anchor, "Normal")
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.space_before = Pt(6)
    figure_paragraph.paragraph_format.space_after = Pt(3)
    figure_paragraph.paragraph_format.keep_with_next = True
    figure_paragraph.paragraph_format.keep_together = True

    run = figure_paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.15))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", f"Figure {figure_number}")

    caption = insert_paragraph_after(figure_paragraph, "Caption")
    add_caption(caption, figure_number, caption_text, bookmark_id)
    return caption


def set_update_fields_on_open(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_existing_figure_alt_text(document: Document) -> None:
    descriptions = [
        "Transformer encoder-decoder architecture showing token embeddings, positional encoding, multi-head attention, feed-forward blocks and the output projection.",
        "Self-attention diagram showing query, key and value projections, scaled dot-product attention and the weighted output representation.",
        "Vision Transformer architecture showing image patchification, patch embeddings, the classification token, positional embeddings, Transformer encoder blocks and the classification head.",
    ]
    for number, (shape, description) in enumerate(
        zip(list(document.inline_shapes)[:3], descriptions), start=1
    ):
        shape._inline.docPr.set("descr", description)
        shape._inline.docPr.set("title", f"Figure {number}")


def validate(document: Document) -> None:
    if len(document.inline_shapes) != 10:
        raise RuntimeError(f"Expected 10 inline figures, found {len(document.inline_shapes)}")
    if len(document.tables) != 11:
        raise RuntimeError(f"Expected 11 tables, found {len(document.tables)}")

    bookmarks = {
        element.get(qn("w:name"))
        for element in document.element.body.xpath(".//w:bookmarkStart")
    }
    expected_bookmarks = {f"fig{number}" for number in range(4, 11)}
    if not expected_bookmarks.issubset(bookmarks):
        missing = sorted(expected_bookmarks - bookmarks)
        raise RuntimeError(f"Missing figure bookmarks: {missing}")

    captions = [
        paragraph
        for paragraph in document.paragraphs
        if any(
            paragraph.text.strip().startswith(f"Figure {number}:")
            for number in range(4, 11)
        )
    ]
    if len(captions) != 7:
        raise RuntimeError(f"Expected seven result captions, found {len(captions)}")

    for shape in document.inline_shapes:
        if not shape._inline.docPr.get("descr"):
            raise RuntimeError("A figure is missing alternative text")

    field_codes = [
        (element.text or "").strip()
        for element in document.element.body.xpath(".//w:instrText")
    ]
    figure_sequences = [code for code in field_codes if code.startswith("SEQ Figure")]
    if len(figure_sequences) != 7:
        raise RuntimeError(
            f"Expected seven Figure SEQ fields, found {len(figure_sequences)}"
        )
    for number in range(4, 11):
        if not any(code.startswith(f"REF fig{number}") for code in field_codes):
            raise RuntimeError(f"Missing cross-reference field for fig{number}")

    zotero_fields = [code for code in field_codes if code.startswith("ADDIN ZOTERO")]
    if len(zotero_fields) != 31:
        raise RuntimeError(f"Expected 31 Zotero fields, found {len(zotero_fields)}")


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    stack_images(
        ROBUSTNESS_FIGURES / "low_data_validation_accuracy_epoch.png",
        ROBUSTNESS_FIGURES / "low_data_validation_loss_epoch.png",
        WORD_FIGURES / "low_data_validation_accuracy_loss.png",
    )
    stack_images(
        ROBUSTNESS_FIGURES / "cifar100_validation_accuracy_epoch.png",
        ROBUSTNESS_FIGURES / "cifar100_validation_loss_epoch.png",
        WORD_FIGURES / "cifar100_validation_accuracy_loss.png",
    )

    document = Document(str(DOCX))
    remove_existing_result_figures(document)
    add_existing_figure_alt_text(document)
    shorten_table_captions(document)
    update_reporting_convention(document)

    core = find_paragraph(document, "Learnable absolute PE achieved")
    shifted = find_paragraph(document, "The shifted additive variant changed")
    assignment = find_paragraph(document, "The deterministic mapping test passed")
    low_data = find_paragraph(document, "Under the reduced-data protocol")
    cifar100 = find_paragraph(document, "The ordering of the four pre-selected models")
    fusion = find_paragraph(document, "The hybrid differed from its order-matched")

    replace_result_paragraph(
        core,
        [
            "Learnable absolute PE achieved the highest mean test accuracy in the core comparison (78.60 ± 0.42%), while shifted multiplicative PE was the strongest fixed design (78.08 ± 0.09%). The no-position control reached 71.29 ± 0.83%, below every positional-encoding variant in Table 6. Among the fixed designs, interval widths varied substantially; the two shifted variants had the narrowest intervals in this experiment. ",
            "Figure ",
            ("fig4", 4),
            " shows validation accuracy and loss for the six basic configurations. The shifted variants are examined separately in Figure ",
            ("fig5", 5),
            ", while the radial extension remains reported in Table 6.",
        ],
    )
    replace_result_paragraph(
        shifted,
        [
            "The shifted additive variant changed test accuracy by +0.17 percentage points (95% CI [−0.57, +0.91]) and improved over its base variant in three of five paired seeds. The shifted multiplicative variant changed accuracy by +0.55 percentage points (95% CI [−0.23, +1.32]) and was higher in all five paired seeds. Because both paired intervals include zero, shifting did not produce a clear or consistently substantial improvement under this five-seed protocol. ",
            "Figure ",
            ("fig5", 5),
            " shows the validation trajectories for the same two paired comparisons.",
        ],
    )
    replace_result_paragraph(
        assignment,
        [
            "The deterministic mapping test passed for all four assignment conventions before these results were aggregated. Mean accuracy varied by only 0.22 percentage points across the no-PE controls and by 0.06 points across the learnable controls. The corresponding ranges were wider for row PE (0.70 points), column PE (2.08 points) and multiplicative PE (2.70 points), recording an assignment interaction for the fixed positional mappings without attributing it to token adjacency. ",
            "Figures ",
            ("fig6", 6),
            " and ",
            ("fig7", 7),
            " show the corresponding validation accuracy and loss trajectories.",
        ],
    )
    replace_result_paragraph(
        low_data,
        [
            "Under the reduced-data protocol, shifted multiplicative PE exceeded learnable absolute PE by 2.54 percentage points at 1,000 examples (95% CI [1.19, 3.89]). The ordering reversed at 5,000 examples, where the paired difference was −1.81 points (95% CI [−2.51, −1.11]), and at 10,000 examples, where the difference was −0.95 points (95% CI [−3.06, 1.15]). The separately trained Full reference favoured learnable PE by 0.52 points (95% CI [0.15, 0.89]). The reversal at 1,000 examples is therefore specific to the most data-limited condition rather than a consistent advantage across subset sizes. ",
            "Figure ",
            ("fig8", 8),
            " shows validation accuracy and loss across epochs for the four reduced-data configurations. The Full reference is not joined to these curves because it was trained under the separately frozen learning-rate protocol.",
        ],
    )
    replace_result_paragraph(
        cifar100,
        [
            "The ordering of the four pre-selected models was unchanged between CIFAR-10 and CIFAR-100. Learnable absolute PE ranked first on CIFAR-100 at 47.66 ± 0.21%, followed by shifted multiplicative PE at 45.77 ± 1.27%, shifted additive PE at 45.24 ± 1.35%, and the no-position control at 40.16 ± 0.69%. This preserves the within-dataset ranking but does not by itself establish that the same gaps will hold for other architectures or image domains. ",
            "Figure ",
            ("fig9", 9),
            " shows the corresponding CIFAR-100 validation accuracy and loss trajectories.",
        ],
    )
    replace_result_paragraph(
        fusion,
        [
            "The hybrid differed from its order-matched learned control by +0.09 pp ([-0.24, +0.41] pp); the selected-checkpoint fixed-position scale had mean 0.0060 and ranged from -0.0403 to 0.0492. The best dual-branch result was the cross-attention model with an MLP head at 77.74 ± 0.44%, but its paired difference from learnable absolute PE was -0.86 pp ([-1.45, -0.27] pp) while its parameter count increased from 809,354 to 2,031,242. Neither squared variant nor radial PE exceeded the strongest core single-branch results. ",
            "Figure ",
            ("fig10", 10),
            " shows the validation dynamics of the five fusion variants alongside learnable absolute PE.",
        ],
    )

    figure_specs = [
        (
            core,
            MAIN_FIGURES / "basic_pe_validation_dynamics.png",
            4,
            "Validation dynamics for six basic PE configurations on CIFAR-10.",
            "Line charts of CIFAR-10 validation accuracy and validation loss over epochs for the core positional-encoding models, averaged over seeds 42 to 46 with sample-standard-deviation bands.",
        ),
        (
            shifted,
            MAIN_FIGURES / "shift_validation_dynamics.png",
            5,
            "Validation dynamics for base and shifted additive and multiplicative PE on CIFAR-10.",
            "Line charts comparing validation accuracy and validation loss for base and shifted additive and multiplicative positional encodings over five CIFAR-10 seeds.",
        ),
        (
            assignment,
            MAIN_FIGURES / "patch_assignment_val_acc_epoch.png",
            6,
            "Patch-assignment validation accuracy on CIFAR-10.",
            "Five-panel line chart of CIFAR-10 validation accuracy for no PE, learnable PE, row PE, column PE and multiplicative PE under four patch-to-position assignments.",
        ),
        (
            assignment,
            MAIN_FIGURES / "patch_assignment_val_loss_epoch.png",
            7,
            "Patch-assignment validation loss on CIFAR-10.",
            "Five-panel line chart of CIFAR-10 validation loss for no PE, learnable PE, row PE, column PE and multiplicative PE under four patch-to-position assignments.",
        ),
        (
            low_data,
            WORD_FIGURES / "low_data_validation_accuracy_loss.png",
            8,
            "Validation dynamics under reduced CIFAR-10 training sets.",
            "Stacked validation accuracy and loss charts for four positional-encoding configurations using 1,000, 5,000 and 10,000 CIFAR-10 training examples across five seeds.",
        ),
        (
            cifar100,
            WORD_FIGURES / "cifar100_validation_accuracy_loss.png",
            9,
            "Validation dynamics for four PE configurations on CIFAR-100.",
            "Stacked CIFAR-100 validation accuracy and loss charts comparing no PE, learnable PE, shifted additive PE and shifted multiplicative PE across five seeds.",
        ),
        (
            fusion,
            MAIN_FIGURES / "fusion_validation_dynamics.png",
            10,
            "Validation dynamics for learnable PE and five fusion variants on CIFAR-10.",
            "Line charts of CIFAR-10 validation accuracy and validation loss for five row-column fusion variants and learnable absolute positional encoding across five seeds.",
        ),
    ]

    for _, image_path, _, _, _ in figure_specs:
        if not image_path.exists():
            raise FileNotFoundError(image_path)

    bookmark_id = next_bookmark_id(document)
    last_insertion_by_anchor: dict[int, Paragraph] = {}
    for anchor, image_path, number, caption, alt_text in figure_specs:
        anchor_key = id(anchor._p)
        insertion_anchor = last_insertion_by_anchor.get(anchor_key, anchor)
        inserted_caption = insert_figure(
            insertion_anchor,
            image_path,
            number,
            caption,
            alt_text,
            bookmark_id,
        )
        last_insertion_by_anchor[anchor_key] = inserted_caption
        bookmark_id += 1

    set_update_fields_on_open(document)
    validate(document)
    document.save(str(BUILDING))

    reopened = Document(str(BUILDING))
    validate(reopened)
    BUILDING.replace(DOCX)

    print(f"Updated: {DOCX}")
    print(f"Figures: {len(reopened.inline_shapes)}")
    print(f"Tables: {len(reopened.tables)}")
    print(f"Composite assets: {WORD_FIGURES}")


if __name__ == "__main__":
    main()
