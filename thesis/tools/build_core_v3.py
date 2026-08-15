from __future__ import annotations

import hashlib
import os
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation_Aligned_v2.docx"
OUTPUT = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation_Core_Draft_v3.docx"
EXPECTED_SOURCE_SHA256 = "40D3680B0C20A240F11EDE118139EEB9E855BDC71D0F855101A0DD28EE4F5B41"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def find_paragraph(document: Document, starts_with: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(starts_with)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {starts_with!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(document: Document, starts_with: str, text: str, style: str | None = None):
    paragraph = find_paragraph(document, starts_with)
    paragraph.clear()
    paragraph.add_run(text)
    if style:
        paragraph.style = style
    return paragraph


def delete_paragraph(paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def remove_body_range(document: Document, start_text: str, end_text: str, include_end: bool = False) -> None:
    start = find_paragraph(document, start_text)._p
    end = find_paragraph(document, end_text)._p
    body = document._element.body
    children = list(body)
    start_index = children.index(start)
    end_index = children.index(end)
    if end_index <= start_index:
        raise RuntimeError(f"Invalid removal range: {start_text!r} -> {end_text!r}")
    stop = end_index + 1 if include_end else end_index
    for element in children[start_index:stop]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def remove_from_paragraph_to_end(document: Document, start_text: str) -> None:
    start = find_paragraph(document, start_text)._p
    body = document._element.body
    children = list(body)
    start_index = children.index(start)
    for element in children[start_index:]:
        if element.tag == qn("w:sectPr"):
            continue
        body.remove(element)


def patch_structure(document: Document) -> None:
    # Clean the title page while preserving the inherited cover design.
    replace_paragraph(document, "MASTER'S DISSERTATION FRAMEWORK", "MSc Dissertation")
    replace_paragraph(document, "Document status:", "Draft status: Core structure v3 — 5 August 2026")
    delete_paragraph(find_paragraph(document, "UCL MSc dissertation planning document"))

    # Cover -> Abstract. Remove all framework-only guidance before the actual abstract.
    remove_body_range(document, "How to Use This Framework", "Proposed Front Matter")
    remove_body_range(document, "Proposed Front Matter", "Abstract")
    abstract_heading = find_paragraph(document, "Abstract")
    abstract_heading.style = "Heading 1"
    delete_paragraph(find_paragraph(document, "DRAFTING NOTE  Target roughly"))
    abstract = replace_paragraph(
        document,
        "Provisional abstract:",
        "Vision Transformers process images as patch-token sequences, yet content-only self-attention does not identify the two-dimensional origin of each patch. This dissertation presents a controlled empirical evaluation of positional encoding in a compact Vision Transformer trained from scratch on CIFAR-10. It compares no positional encoding, learnable absolute embeddings, fixed row- and column-based sinusoidal constructions, additive and multiplicative two-dimensional variants, coordinate-shifted variants, patch-to-position assignments, and exploratory hybrid and dual-branch fusion models. All locally consolidated configurations are evaluated across five training seeds under a common architecture, fixed data split and validation-selected checkpoint protocol. The no-position baseline is clearly weaker than the leading positional models. Learnable absolute encoding remains a strong reference, while shifted multiplicative encoding is the strongest tested fixed positional encoding. Learnable models vary little across the tested patch assignments, whereas several fixed encodings are more sensitive to the mapping between image patches and positional vectors. A hybrid model attains the highest numerical mean, but its small margin and traversal confound do not support a general optimisation claim; similarly, no dual-branch fusion model exceeds the best single-branch or hybrid result. The study therefore contributes a reproducible empirical map of positional-design choices and their limitations rather than a universally superior architecture. Generalisation beyond CIFAR-10 remains to be established.",
        "Normal",
    )
    abstract.paragraph_format.first_line_indent = None
    delete_paragraph(find_paragraph(document, "CITATION POINT  Abstracts normally minimise citations"))

    acknowledgements = find_paragraph(document, "Acknowledgements")
    acknowledgements.style = "Heading 1"
    replace_paragraph(
        document,
        "Optional. Keep brief and personal",
        "DRAFTING NOTE  Add concise acknowledgements after the technical chapters are stable.",
        "Callout",
    )
    contents = replace_paragraph(document, "Contents, Lists and Abbreviations", "Contents", "Heading 1")
    contents.style = "Heading 1"

    # Remove the argument map and internal word-budget planning. Introduction is now the first main chapter.
    remove_body_range(document, "Dissertation Argument at a Glance", "Introduction")

    # Combine experiment design and results into one top-level chapter.
    replace_paragraph(document, "Experiments", "Experiments and Results", "Heading 1")
    results_heading = find_paragraph(document, "Results")
    delete_paragraph(results_heading)
    delete_paragraph(find_paragraph(document, "CHAPTER JOB  Present observations in the order"))

    # Fold limitations into Discussion so the chapter sequence remains compact.
    limitations = replace_paragraph(
        document,
        "Limitations and Threats to Validity",
        "Limitations and Threats to Validity",
        "Heading 2",
    )
    demote = False
    for paragraph in document.paragraphs:
        if paragraph._p is limitations._p:
            demote = True
            continue
        if demote and paragraph.text.strip().startswith("Conclusion and Future Work"):
            break
        if demote and paragraph.style.name == "Heading 2":
            paragraph.style = "Heading 3"

    # Conclusion is the final top-level section; everything after it is removed.
    replace_paragraph(document, "Conclusion and Future Work", "Conclusion", "Heading 1")
    conclusion_subheadings = [
        p for p in document.paragraphs if p.text.strip() == "Conclusion" and p.style.name == "Heading 2"
    ]
    if len(conclusion_subheadings) != 1:
        raise RuntimeError(f"Expected one redundant Conclusion subheading; found {len(conclusion_subheadings)}")
    delete_paragraph(conclusion_subheadings[0])
    replace_paragraph(
        document,
        "Provisional conclusion starter:",
        "This dissertation examined positional encoding as a controlled design choice in compact Vision Transformers. Across 32 CIFAR-10 configurations and five training seeds per configuration, positional information consistently improved the leading models over the no-position baseline. Learnable absolute PE remained the strongest core reference, while shifted multiplicative PE was the strongest tested fixed encoding. Alternative patch-to-position assignments had little effect on learnable PE but interacted more strongly with several fixed encodings. Hybrid and fusion studies provide useful boundary evidence, yet their confounds prevent them from supporting headline optimisation claims. The conclusions remain specific to the evaluated architecture, dataset, split and training protocol until the planned generalisation experiments are complete.",
        "Normal",
    )
    remove_from_paragraph_to_end(document, "Appendix Plan")

    for section in document.sections:
        for header in (section.header, section.first_page_header):
            for paragraph in header.paragraphs:
                if "MSc Dissertation Framework" in paragraph.text:
                    paragraph.text = "MSc Dissertation Draft | Positional Encoding in Compact ViTs"

    document.core_properties.title = "Positional Encoding in Compact Vision Transformers: A Controlled Evaluation"
    document.core_properties.subject = "Core MSc dissertation draft structure"
    document.core_properties.comments = "Core structure v3: cover to Abstract; Conclusion is the final section; no figures embedded."


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    if sha256(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("Aligned v2 changed after the v3 structure was planned; refresh the source audit before editing.")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    os.chmod(OUTPUT, 0o666)
    document = Document(OUTPUT)
    patch_structure(document)
    document.save(OUTPUT)
    if sha256(SOURCE) != EXPECTED_SOURCE_SHA256:
        raise RuntimeError("The retained v2 source changed during generation.")
    print(OUTPUT)


if __name__ == "__main__":
    main()
