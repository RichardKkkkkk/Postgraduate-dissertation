from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
doc = Document(path)

errors = []

# Hybrid belongs only to Future Work.
future_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "6.1 Future Work")
for i, paragraph in enumerate(doc.paragraphs):
    if "hybrid" in paragraph.text.lower() and i <= future_index:
        errors.append(f"Hybrid remains before Future Work at paragraph {i}: {paragraph.text}")

# Equation labels and prose references must now stop at 13.
equation_labels = []
for paragraph in doc.paragraphs:
    text = "".join(paragraph._p.xpath(".//m:t/text()"))
    if "#" in text:
        equation_labels.append(text.rsplit("#", 1)[-1])
if equation_labels != [str(i) for i in range(1, 14)]:
    errors.append(f"Unexpected equation labels: {equation_labels}")
for paragraph in doc.paragraphs:
    if "Equation (14)" in paragraph.text:
        errors.append(f"Stale Equation (14) reference: {paragraph.text}")

# Verify headline structure and table sizes.
expected_headings = {
    "3.6 Dual-Branch Fusion Extensions",
    "3.6.1 Latent Fusion",
    "3.6.2 Bidirectional Cross-Attention Fusion",
    "4.7 Fusion and Other Extensions",
    "6.1 Future Work",
}
headings = {p.text.strip() for p in doc.paragraphs if p.style.name.startswith("Heading")}
missing = expected_headings - headings
if missing:
    errors.append(f"Missing headings: {sorted(missing)}")

if len(doc.tables) != 12:
    errors.append(f"Expected 12 tables, found {len(doc.tables)}")
if len(doc.tables[8].rows) != 13:
    errors.append(f"Low-data table should have 12 data rows, found {len(doc.tables[8].rows)-1}")
if len(doc.tables[10].rows) != 10:
    errors.append(f"Extension table should have 9 data rows, found {len(doc.tables[10].rows)-1}")
if len(doc.tables[11].rows) != 5:
    errors.append(f"Future hybrid table should have 4 data rows, found {len(doc.tables[11].rows)-1}")

# Walk body order to ensure each figure drawing is immediately followed by a caption.
body_items = []
for child in doc.element.body.iterchildren():
    if child.tag == qn("w:p"):
        item = Paragraph(child, doc)
        body_items.append(("p", item))
    elif child.tag == qn("w:tbl"):
        body_items.append(("tbl", Table(child, doc)))

for i, (kind, item) in enumerate(body_items):
    if kind == "p" and item._p.xpath(".//w:drawing"):
        if i + 1 >= len(body_items) or body_items[i + 1][0] != "p" or body_items[i + 1][1].style.name != "Caption":
            errors.append(f"Figure at body item {i} is not followed by a caption")

# Every table should be preceded by a caption, apart from front-matter tables if any.
for i, (kind, item) in enumerate(body_items):
    if kind == "tbl" and i > 0:
        previous = body_items[i - 1]
        if previous[0] != "p" or previous[1].style.name != "Caption":
            errors.append(f"Table at body item {i} is not preceded by a caption")

# New/modified non-equation text should use black Arial. Report all text-run exceptions.
font_exceptions = []
colour_exceptions = []
for p_idx, paragraph in enumerate(doc.paragraphs):
    for run in paragraph.runs:
        if not run.text.strip():
            continue
        rpr = run._element.rPr
        fonts = rpr.rFonts if rpr is not None else None
        ascii_font = fonts.get(qn("w:ascii")) if fonts is not None else None
        if ascii_font not in (None, "Arial"):
            font_exceptions.append((p_idx, run.text[:30], ascii_font))
        color = run.font.color.rgb
        if color is not None and str(color) != "000000":
            colour_exceptions.append((p_idx, run.text[:30], str(color)))

print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Inline images: {len(doc.inline_shapes)}")
print(f"Equation labels: {equation_labels}")
print(f"Font exceptions: {font_exceptions[:10]}")
print(f"Colour exceptions: {colour_exceptions[:10]}")
if errors:
    print("AUDIT FAILED")
    for error in errors:
        print(f"- {error}")
    raise SystemExit(1)
print("AUDIT PASSED")
