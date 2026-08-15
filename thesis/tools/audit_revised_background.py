from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
document = Document(path)

section = []
inside = False
for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text.startswith("2.2 "):
        inside = True
    if inside:
        section.append(text)
    if text.startswith("2.4 "):
        break

joined = "\n".join(section)
print("paragraphs", len(document.paragraphs))
print("tables", len(document.tables))
print("images", len(document.inline_shapes))
print("displayed_equations", len(document.element.xpath(".//*[local-name()='oMathPara']")))
print("class_token_mentions_2_2_to_2_3", joined.lower().count("classification token"))
print("equation_5_mentions_2_2_to_2_3", joined.count("Equation (5)"))
print("positional_matrix_definition_mentions", joined.count("positional embedding matrix"))
print("broken_todo_markers", sum(joined.count(marker) for marker in ("TODO", "[CITATION NEEDED]", "[?]")))
print("file_size", path.stat().st_size)
with ZipFile(path) as archive:
    document_xml = archive.read("word/document.xml")
    media_files = [name for name in archive.namelist() if name.startswith("word/media/") and not name.endswith("/")]
print("zotero_fields", document_xml.count(b"ADDIN ZOTERO_ITEM"))
print("media_files", len(media_files))
print("new_heading_xml", document_xml.count(b"Classification Token and Baseline Input"))
print("old_heading_xml", document_xml.count(b"Patch Embedding and Class Token"))
