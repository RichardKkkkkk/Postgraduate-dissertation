from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation.docx"
BUILDING = ROOT / "thesis" / ".Yikai_Zhao_MSc_Dissertation.math.tmp.docx"


def set_body_run_font(run) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), "Arial")


def add_body_text(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    set_body_run_font(run)


def add_ref_field(paragraph, bookmark: str, display_text: str) -> None:
    begin_run = paragraph.add_run()
    set_body_run_font(begin_run)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    set_body_run_font(instruction_run)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" REF {bookmark} \\h "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    set_body_run_font(separate_run)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(display_text)
    set_body_run_font(result_run)

    end_run = paragraph.add_run()
    set_body_run_font(end_run)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def clear_paragraph_content(paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def math_run(text: str) -> OxmlElement:
    run = OxmlElement("m:r")
    math_text = OxmlElement("m:t")
    math_text.text = text
    run.append(math_text)

    word_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Cambria Math")
    fonts.set(qn("w:hAnsi"), "Cambria Math")
    word_properties.append(fonts)
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "000000")
    word_properties.append(colour)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "21")
    word_properties.append(size)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "21")
    word_properties.append(size_cs)
    run.append(word_properties)
    return run


def fraction(numerator: str, denominator: str) -> OxmlElement:
    fraction_element = OxmlElement("m:f")
    properties = OxmlElement("m:fPr")
    fraction_type = OxmlElement("m:type")
    fraction_type.set(qn("m:val"), "bar")
    properties.append(fraction_type)
    fraction_element.append(properties)

    numerator_element = OxmlElement("m:num")
    numerator_element.append(math_run(numerator))
    fraction_element.append(numerator_element)

    denominator_element = OxmlElement("m:den")
    denominator_element.append(math_run(denominator))
    fraction_element.append(denominator_element)
    return fraction_element


def subscript(base: str, sub: str) -> OxmlElement:
    element = OxmlElement("m:sSub")
    base_element = OxmlElement("m:e")
    base_element.append(math_run(base))
    element.append(base_element)
    sub_element = OxmlElement("m:sub")
    sub_element.append(math_run(sub))
    element.append(sub_element)
    return element


def append_inline_math(paragraph, *children: OxmlElement) -> None:
    math = OxmlElement("m:oMath")
    for child in children:
        math.append(child)
    paragraph._p.append(math)


def find_paragraph(document: Document, starts_with: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(starts_with)]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph starting with {starts_with!r}, found {len(matches)}"
        )
    return matches[0]


def paragraph_with_bookmark(document: Document, bookmark: str):
    matches = [
        p
        for p in document.paragraphs
        if any(
            element.get(qn("w:name")) == bookmark
            for element in p._p.xpath(".//w:bookmarkStart")
        )
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph containing bookmark {bookmark!r}, found {len(matches)}"
        )
    return matches[0]


def replace_math_token(paragraph, old: str, new: str) -> None:
    matches = [element for element in paragraph._p.xpath(".//m:t") if element.text == old]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one math token {old!r} in paragraph, found {len(matches)}"
        )
    matches[0].text = new


def rewrite_shift_explanation(document: Document) -> None:
    paragraph = find_paragraph(document, "The shifted variants keep")
    clear_paragraph_content(paragraph)
    add_body_text(
        paragraph,
        "The shifted variants keep the same combination rules but use different sinusoidal frequency schedules for the two axes. For the sine–cosine pair indexed by ",
    )
    append_inline_math(paragraph, math_run("i"))
    add_body_text(paragraph, ", the row frequency uses exponent ")
    append_inline_math(paragraph, math_run("−"), fraction("2i", "d"))
    add_body_text(paragraph, ", and the column frequency uses ")
    append_inline_math(paragraph, math_run("−"), fraction("2i+1", "d"))
    add_body_text(
        paragraph,
        ", as defined in Equation (",
    )
    add_ref_field(paragraph, "eq8", "8")
    add_body_text(
        paragraph,
        "). The sine and cosine channels in each pair share the corresponding axis-specific frequency. The resulting vectors ",
    )
    append_inline_math(paragraph, subscript("S", "r"), math_run("(r)"))
    add_body_text(paragraph, " and ")
    append_inline_math(paragraph, subscript("S", "c"), math_run("(c)"))
    add_body_text(paragraph, " are combined in Equation (")
    add_ref_field(paragraph, "eq9", "9")
    add_body_text(paragraph, ").")


def rewrite_hadamard_definition(document: Document) -> None:
    paragraph = find_paragraph(document, "To represent both axes")
    clear_paragraph_content(paragraph)
    add_body_text(
        paragraph,
        "To represent both axes, the row and column vectors are combined in two ways. The additive encoding sums the vectors, and the multiplicative encoding takes their element-wise product. Both operations preserve the 128-dimensional embedding size and are defined in Equation (",
    )
    add_ref_field(paragraph, "eq7", "7")
    add_body_text(paragraph, "). The symbol ")
    append_inline_math(paragraph, math_run("⊙"))
    add_body_text(paragraph, " denotes element-wise multiplication.")


def math_text(paragraph) -> str:
    return "".join((element.text or "") for element in paragraph._p.xpath(".//m:t"))


def validate(document: Document) -> None:
    all_math_tokens = [
        element.text or "" for element in document.element.body.xpath(".//m:t")
    ]
    forbidden = [token for token in all_math_tokens if any(x in token for x in "/∕÷*·")]
    if forbidden:
        raise RuntimeError(f"Non-standard math operators remain: {forbidden}")

    expected_fraction_counts = {
        "eq1": 4,
        "eq2": 1,
        "eq3": 1,
        "eq8": 2,
        "eq11": 1,
        "eq13": 1,
        "eq14": 1,
    }
    for bookmark, minimum in expected_fraction_counts.items():
        paragraph = paragraph_with_bookmark(document, bookmark)
        count = len(paragraph._p.xpath(".//m:f"))
        if count < minimum:
            raise RuntimeError(
                f"{bookmark} should contain at least {minimum} stacked fraction(s), found {count}"
            )

    shift_paragraph = find_paragraph(document, "The shifted variants retain")
    if len(shift_paragraph._p.xpath(".//m:f")) != 2 or "/" in shift_paragraph.text:
        raise RuntimeError("The shifted-frequency exponents are not two stacked fractions")

    for bookmark in ("eq7", "eq9", "eq10"):
        equation_text = math_text(paragraph_with_bookmark(document, bookmark))
        if equation_text.count("⊙") != 1:
            raise RuntimeError(f"{bookmark} must contain exactly one Hadamard product symbol")
        if "×" in equation_text:
            raise RuntimeError(f"{bookmark} incorrectly uses ordinary multiplication")

    squared = paragraph_with_bookmark(document, "eq10")
    grouped_square = [
        element
        for element in squared._p.xpath(".//m:sSup")
        if element.xpath(
            "./*[local-name()='e']/*[local-name()='d']"
            "//*[local-name()='t'][contains(text(), '⊙')]"
        )
    ]
    if len(grouped_square) != 1:
        raise RuntimeError("Equation (10) does not square the grouped Hadamard product")

    if "H×W" not in math_text(paragraph_with_bookmark(document, "eq3")):
        raise RuntimeError("Equation (3) does not use ordinary multiplication for H×W")
    if "α·" in math_text(paragraph_with_bookmark(document, "eq12")):
        raise RuntimeError("Equation (12) still uses a centred dot for scalar multiplication")

    ci_paragraph = find_paragraph(document, "For each five-seed model")
    ci_text = math_text(ci_paragraph)
    if "2.776×" not in ci_text or len(ci_paragraph._p.xpath(".//m:f")) < 1:
        raise RuntimeError("The confidence-interval half-width is not an explicit product and fraction")

    bookmarks = {
        element.get(qn("w:name"))
        for element in document.element.body.xpath(".//w:bookmarkStart")
    }
    expected_equations = {f"eq{number}" for number in range(1, 15)}
    if not expected_equations.issubset(bookmarks):
        raise RuntimeError("An equation bookmark was lost")

    if len(document.inline_shapes) != 10 or len(document.tables) != 11:
        raise RuntimeError("Figure or table count changed during the math edit")

    field_codes = [
        (element.text or "").strip()
        for element in document.element.body.xpath(".//w:instrText")
    ]
    if sum(code.startswith("ADDIN ZOTERO") for code in field_codes) != 31:
        raise RuntimeError("Zotero fields changed during the math edit")
    if sum(code.startswith("SEQ ") for code in field_codes) != 18:
        raise RuntimeError("SEQ fields changed during the math edit")
    if sum(code.startswith("REF ") for code in field_codes) != 17:
        raise RuntimeError("REF fields changed during the math edit")


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    document = Document(str(DOCX))
    replace_math_token(paragraph_with_bookmark(document, "eq3"), "H·W", "H×W")
    replace_math_token(
        paragraph_with_bookmark(document, "eq12"), "]+L+α·", "]+L+α"
    )
    confidence_interval = find_paragraph(document, "For each five-seed model")
    replace_math_token(confidence_interval, "2.776·", "2.776×")
    rewrite_hadamard_definition(document)
    rewrite_shift_explanation(document)

    validate(document)
    document.save(str(BUILDING))
    reopened = Document(str(BUILDING))
    validate(reopened)
    BUILDING.replace(DOCX)

    with ZipFile(DOCX) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("The saved DOCX package failed its ZIP integrity test")

    print(f"Updated: {DOCX}")
    print("Equation bookmarks: eq1-eq14")
    print("Math convention: stacked fractions, ordinary multiplication, Hadamard ⊙")


if __name__ == "__main__":
    main()
