from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation.docx"
OUT = ROOT / "thesis" / "_extract_review.txt"
CIFAR100 = ROOT / "results" / "cifar100_4models_5seeds" / "metrics"


def main() -> None:
    print("exists", DOCX.exists(), "size", DOCX.stat().st_size if DOCX.exists() else None)
    doc = Document(str(DOCX))
    lines: list[str] = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if text:
            lines.append(f"[{style}] {text}")
    for ti, table in enumerate(doc.tables, 1):
        lines.append("")
        lines.append(f"=== TABLE {ti} ===")
        for row in table.rows:
            cells = " | ".join(c.text.strip().replace("\n", " ") for c in row.cells)
            lines.append(cells)
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print("lines", len(lines), "chars", OUT.stat().st_size, "tables", len(doc.tables))
    if CIFAR100.exists():
        summaries = sorted(CIFAR100.rglob("*_summary.json"))
        print("cifar100_summaries", len(summaries))
        for path in summaries:
            print(path.relative_to(CIFAR100).as_posix())
    else:
        print("cifar100_summaries", 0)


if __name__ == "__main__":
    main()
