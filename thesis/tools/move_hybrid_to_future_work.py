from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


DOCX = Path(r"D:\code\Postgraduate-dissertation\thesis\Yikai_Zhao_MSc_Dissertation.docx")
LOW_DATA_FIGURE = Path(r"D:\code\Postgraduate-dissertation\thesis\assets\low_data_validation_accuracy_epoch_valid.png")
FUSION_FIGURE = Path(r"D:\code\Postgraduate-dissertation\thesis\assets\fusion_validation_accuracy_epoch_valid.png")


def set_run_font(run, size=11, bold=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Arial")


def replace_paragraph(paragraph, text):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def paragraph_with_prefix(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"Paragraph not found: {prefix}")


def equation_paragraph_containing(doc, text):
    for paragraph in doc.paragraphs:
        equation_text = "".join(paragraph._p.xpath(".//m:t/text()"))
        if text in equation_text:
            return paragraph
    raise ValueError(f"Equation paragraph not found: {text}")


def replace_prefix(doc, prefix, text):
    replace_paragraph(paragraph_with_prefix(doc, prefix), text)


def set_equation_number(paragraph, old_number, new_number):
    nodes = paragraph._p.xpath(".//m:t")
    for node in reversed(nodes):
        if node.text == str(old_number):
            node.text = str(new_number)
            return
    raise ValueError(f"Equation number {old_number} not found")


def remove_table_rows(table, predicate):
    for row in list(table.rows):
        values = [cell.text.strip() for cell in row.cells]
        if predicate(values):
            table._tbl.remove(row._tr)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                shade_cell(cell, "D9E2F3")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=9, bold=(row_idx == 0))


def insert_after(anchor, element):
    anchor._p.addnext(element)


def replace_figure(paragraph, path, width_inches=6.45):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))


def paragraph_before(paragraph):
    previous = paragraph._p.getprevious()
    if previous is None or previous.tag != qn("w:p"):
        raise ValueError(f"No paragraph immediately before: {paragraph.text}")
    return Paragraph(previous, paragraph._parent)


doc = Document(DOCX)

# Abstract and Introduction: keep the main story centred on learned versus fixed PE.
replace_prefix(
    doc,
    "Vision Transformers (ViTs) are widely used",
    "Vision Transformers (ViTs) are widely used for image classification, but self-attention alone does not represent the spatial arrangement of image patches. Positional encoding (PE) supplies this missing information. This dissertation compares learned and fixed PE in a small ViT trained from scratch. CIFAR-10 is used for the main experiments, and each comparison is repeated with five training seeds. Positional encoding improved full-data performance over a model without PE. Learnable absolute PE achieved the highest mean accuracy in the core comparison, while shifted multiplicative PE was the strongest fixed method. Under reduced training data, shifted multiplicative PE led at 1,000 and 5,000 examples. Learnable PE had a slightly higher mean at 10,000 examples and performed best with the full training set. The selected methods kept the same test-accuracy ranking on CIFAR-100. More complex dual-branch models did not improve on the learned single-branch reference. These results show that the value of learned and fixed PE depends on the amount of training data, while additional architectural complexity does not necessarily improve classification.",
)
replace_prefix(
    doc,
    "Many previous studies compare positional encoding methods",
    "Many previous studies compare positional encoding methods using different model architectures or training settings [5], [10]–[12]. This makes it difficult to know whether a change in performance comes from the positional encoding itself or from other differences in the model. This project therefore keeps the patch size, ViT backbone, and training process the same. It studies whether positional encoding improves the model, which fixed PE methods perform best, and whether the results change when the patch-to-position mapping or amount of training data is changed. Dual-branch models are tested separately because they change the architecture and number of parameters.",
)
replace_prefix(
    doc,
    "The experiments show that positional information is important",
    "The experiments show that positional information is important for the ViT used in this project. Learnable absolute PE has the highest mean accuracy in the full-data core comparison, while shifted multiplicative PE is the strongest fixed method. The fixed method leads at 1,000 and 5,000 training examples, but not at 10,000 examples or with the full training set. The same four-model test-accuracy ranking is observed on CIFAR-100. Dual-branch fusion adds substantial capacity without improving on the learned single-branch model.",
)

# Literature and Methodology: remove the hybrid from the formal method pipeline.
replace_prefix(
    doc,
    "The practical gap is therefore one of attribution",
    "The practical gap is therefore one of attribution, not a lack of proposed encodings. Results from different studies may reflect changes in architecture, scale, data, augmentation or optimisation as well as the position method. This project instead uses one ViT and a common protocol to separate four factors: PE presence and form, patch-coordinate mapping, limited training data, and the cost of multi-branch complexity. This controlled comparison does not claim a universally superior encoding or a new general-purpose Transformer.",
)
replace_prefix(
    doc,
    "This chapter explains how the experiments were constructed",
    "This chapter explains how the experiments were constructed so that each reported comparison can be reproduced and interpreted. Every image passes through the same broad pipeline: dataset-specific preprocessing, division into patches, assignment of positional information, and processing by a shared Transformer encoder. Validation performance selects the checkpoint. The official test set is used only after selection. The following sections define the controlled design, data, shared architecture, positional encodings, patch mappings, fusion extensions, training protocol and statistical analysis.",
)
replace_prefix(
    doc,
    "The study changes one main factor at a time",
    "The study changes one main factor at a time rather than introducing a new general-purpose Transformer. Within each comparison, the ViT backbone, data split, optimiser and checkpoint-selection rule remain fixed. The factors that do change are the positional encoding, the assignment of positional vectors to physical patches, and, in the architecture extension, the use of dual-branch fusion components. This design reduces the number of competing explanations for an observed accuracy difference.",
)
replace_prefix(
    doc,
    "The main evaluation uses CIFAR-10 and five training seeds",
    "The main evaluation uses CIFAR-10 and five training seeds. No PE provides a genuine no-position baseline, while learnable absolute PE provides the standard learned reference. Fixed row-column encodings then test different ways to represent the two-dimensional grid without trainable positional parameters. Patch-assignment and fusion models are analysed separately because they change more than the core positional vector. CIFAR-100 tests whether the observed ranking is retained on a finer-grained classification problem, and reduced-data CIFAR-10 tests whether the relationship changes when fewer labelled images are available.",
)
replace_prefix(
    doc,
    "The reduced-data study draws subsets",
    "The reduced-data study draws subsets of 1,000, 5,000 and 10,000 images from the fixed 45,000-image CIFAR-10 training pool. It compares no PE, learnable absolute PE and shifted multiplicative PE using the common learning rate in Section 3.7. The full-data results use the same optimisation settings and are included as a fourth data-size condition. Validation and test sets do not change. Within each seed, all three models receive the same subset. The subset changes across seeds, so the reported variation includes both stochastic training and the sampled training images.",
)
replace_prefix(
    doc,
    "The core comparison is deliberately limited",
    "The core comparison is deliberately limited to methods that produce a 128-component positional term and add it at the same point in the shared ViT. Other approaches reviewed in Sections 2.3 and 2.6, including relative-position, conditional and rotary encodings [10]–[12], either modify the attention calculation or introduce learned position-generating components. Including them would change more than the positional vector itself. The dual-branch models also alter parameter count and architecture, so they are evaluated separately in Section 3.6.",
)

replace_paragraph(paragraph_with_prefix(doc, "3.6 Hybrid and Dual-Branch Extensions"), "3.6 Dual-Branch Fusion Extensions")
replace_prefix(
    doc,
    "The extension models ask whether extra positional",
    "The fusion extensions ask whether separate row and column branches provide useful information beyond the single-branch models. Each design processes the same image through two ViT encoders and combines their representations before classification. These models change both the architecture and parameter count, so their results are reported separately and interpreted alongside model capacity.",
)
delete_paragraph(paragraph_with_prefix(doc, "3.6.1 Learnable-Fixed Hybrid"))
delete_paragraph(paragraph_with_prefix(doc, "The hybrid uses normal column-major ordering"))
hybrid_eq = equation_paragraph_containing(doc, "Z0=[")
delete_paragraph(hybrid_eq)
delete_paragraph(paragraph_with_prefix(doc, "Starting at zero makes the initial input"))
replace_paragraph(paragraph_with_prefix(doc, "3.6.2 Latent Fusion"), "3.6.1 Latent Fusion")
replace_paragraph(paragraph_with_prefix(doc, "3.6.3 Bidirectional Cross-Attention Fusion"), "3.6.2 Bidirectional Cross-Attention Fusion")

latent_text = paragraph_with_prefix(doc, "Latent fusion uses two ViT encoders")
replace_paragraph(
    latent_text,
    "Latent fusion uses two ViT encoders on the same image: one with row PE and one with column PE. Their classification-token vectors are h_r and h_c. One variant averages them, one applies an MLP after averaging, and one concatenates them before a layer-normalised MLP. Equation (12) summarises the basic rules.",
)
latent_eq = equation_paragraph_containing(doc, "hmean=")
set_equation_number(latent_eq, 13, 12)
replace_prefix(
    doc,
    "In Equation (13),",
    "In Equation (12), g(·) denotes the learned fusion MLP and square brackets denote vector concatenation. A linear head maps the fused representation to class scores. On CIFAR-10, mean, mean-MLP and concatenation-MLP fusion contain 1,600,778, 1,732,746 and 1,798,538 trainable parameters, compared with 801,034 for one fixed-PE encoder. Most of the increase comes from duplicating the encoder. The MLP variants add further parameters in the fusion stage.",
)
cross_text = paragraph_with_prefix(doc, "Cross-attention allows one branch")
replace_paragraph(
    cross_text,
    "Cross-attention allows one branch to request information from the other before classification. The models therefore retain the complete row and column token sequences rather than only their final classification tokens. For one attention head, Q is the sequence being updated and C is the other branch used as context. Equation (13) gives the operation.",
)
cross_eq = equation_paragraph_containing(doc, "CA(Q,C)=")
set_equation_number(cross_eq, 14, 13)
replace_prefix(
    doc,
    "In Equation (14),",
    "In Equation (13), W_Q, W_K and W_V are learned query, key and value projections, and d_h = D/4 = 32 is the width of each attention head. Four heads perform this exchange in parallel. Their outputs are concatenated and projected back to the model dimension. One block updates row tokens from column context and another updates column tokens from row context. Both use pre-normalisation, residual connections and an MLP. The updated classification tokens are concatenated and passed to either a linear or MLP head. These variants contain 1,999,114 and 2,031,242 parameters, so accuracy is interpreted alongside model size.",
)

replace_prefix(
    doc,
    "Test accuracy and cross-entropy loss are the primary outcomes",
    "Test accuracy and cross-entropy loss are the primary outcomes. Accuracy is the proportion of images assigned the correct class, so higher is better. Cross-entropy also reflects the confidence of the predicted probabilities, and lower is better. Macro-precision, macro-recall and macro-F1 are calculated from the confusion matrix as secondary measures. Each class contributes equally, which is useful for the 100-class extension. Per-class scores and confusion matrices support diagnosis, while trainable parameter count is reported for the fusion models because they change capacity.",
)

# Results: keep low-data as a three-way core comparison and retain fusion as the architecture extension.
replace_prefix(
    doc,
    "This chapter presents the experiments in the order needed",
    "This chapter presents the experiments in the order needed to answer the study questions. It begins with the core positional-encoding comparison, then isolates the shifted schedule, changes patch-to-position assignment, reduces the available training data, moves selected models to CIFAR-100, and finally examines fusion and fixed-PE extensions. Each section states what changes and what remains controlled before reporting the selected-checkpoint test results. Mechanistic explanations and practical recommendations are reserved for Chapter 5 so that observation and interpretation remain distinct.",
)
replace_prefix(
    doc,
    "The reduced-data experiment asks whether a fixed spatial prior",
    "The reduced-data experiment asks whether a fixed spatial prior becomes more useful when fewer labelled images are available. Three models are compared at 1,000, 5,000, 10,000 and the full 45,000-example training split: no PE, learnable PE and shifted multiplicative PE. Within each data size, all three models use the same sampled images within a seed and share the architecture, learning rate, optimisation, checkpoint rule and seed set. The paired comparisons therefore isolate the PE choice as far as this protocol allows.",
)
replace_prefix(
    doc,
    "The test ranking changes with the amount of training data",
    "The test ranking changes with the amount of training data. Shifted multiplicative PE exceeded learnable PE by 3.37 percentage points at 1,000 examples with a paired 95% CI from 2.55 to 4.19. It also led by 1.12 points at 5,000 examples with an interval from 0.46 to 1.78. At 10,000 examples, the difference was −0.65 points with an interval from −1.31 to 0.02. The full-data comparison favoured learnable PE by 0.52 points with an interval from 0.15 to 0.89. Figure 7 shows the validation trajectories under the four data sizes. The fixed-versus-learnable separation is most visible at 1,000 and 5,000 examples.",
)
replace_paragraph(
    paragraph_with_prefix(doc, "Figure 7. Mean CIFAR-10 validation accuracy"),
    "Figure 7. Mean CIFAR-10 validation accuracy for three PE settings and four training-set sizes. Shading shows pointwise 95% t confidence intervals across five seeds.",
)

replace_paragraph(paragraph_with_prefix(doc, "4.7 Hybrid, Fusion and Other Extensions"), "4.7 Fusion and Other Extensions")
replace_prefix(
    doc,
    "The final group asks whether extra complexity earns better performance",
    "The final group asks whether dual-branch complexity earns better performance than the single-branch models. Fusion models are shown with parameter count because they duplicate the encoder or add learned fusion layers. Squared and radial PE remain in the table as targeted fixed-encoding tests, even though they were not among the strongest core methods.",
)
replace_paragraph(paragraph_with_prefix(doc, "Table 11. Hybrid, fusion and fixed-PE extensions."), "Table 11. Fusion and fixed-PE extensions.")
replace_prefix(
    doc,
    "The hybrid differed from its order-matched learned reference",
    "The best dual-branch result came from cross-attention with an MLP head at 77.74 ± 0.44%, but it remained 0.86 points below learnable PE in the paired test comparison while increasing the parameter count from 809,354 to 2,031,242. Figure 9 separates aggregation-based and cross-attention fusion. The validation trajectories show similar convergence, but Table 11 confirms that none of the larger fusion models improves on the learned single-branch reference at the selected checkpoint.",
)
replace_prefix(
    doc,
    "Across the experiments, positional information improved full-data performance",
    "Across the experiments, positional information improved full-data performance over no PE. Learnable PE had the highest core mean on CIFAR-10, while shifted multiplicative PE was the strongest fixed design. The fixed method led at 1,000 and 5,000 examples, then fell slightly behind learnable PE at 10,000 and full data. Patch assignment had little effect on learnable PE but larger effects on several fixed encodings. The four selected methods retained the same test-accuracy order on CIFAR-100. Dual-branch fusion did not improve on the learned single-branch reference.",
)
replace_prefix(
    doc,
    "The evidence now covers two datasets",
    "The evidence now covers two datasets, four training-set sizes, several patch assignments, and both single-branch and dual-branch models. It remains bounded by one small ViT architecture, one split-construction rule, and five seeds. Chapter 5 discusses what these results support and where further evidence is still needed.",
)

# Remove hybrid rows from the low-data and extension tables.
low_data_table = doc.tables[8]
remove_table_rows(low_data_table, lambda values: len(values) > 1 and "Learnable + multiplicative PE" in values[1])
extension_table = doc.tables[10]
remove_table_rows(extension_table, lambda values: values and ("Order-matched learned PE" in values[0] or "Hybrid learned + fixed PE" in values[0]))
format_table(low_data_table, [1700, 2350, 1350, 1300, 2660])
format_table(extension_table, [2600, 1600, 1400, 1550, 1550])

# Replace low-data and fusion figures with the versions that exclude the hybrid.
replace_figure(paragraph_before(paragraph_with_prefix(doc, "Figure 7. Mean CIFAR-10 validation accuracy")), LOW_DATA_FIGURE)
replace_figure(paragraph_before(paragraph_with_prefix(doc, "Figure 9.")), FUSION_FIGURE)

# Discussion and Conclusion: remove hybrid from headline findings.
replace_prefix(
    doc,
    "The hybrid and fusion results also show",
    "The fusion results show that extra architectural complexity is not enough on its own. The best fusion model remained below learnable PE even though it used about 2.5 times as many trainable parameters. Because the fusion models were not parameter matched, the experiment cannot isolate the effect of the fusion operation from the change in capacity. It can still support the practical conclusion that the added computation did not produce a better result under the tested design.",
)
replace_prefix(
    doc,
    "This dissertation examined positional encoding as a controlled design choice",
    "This dissertation examined positional encoding as a controlled design choice in a small Vision Transformer trained from scratch. The patch representation, encoder, optimisation and checkpoint rule were held constant so that the main comparisons could be linked to PE design. The evaluation covered CIFAR-10, CIFAR-100, four CIFAR-10 training-set sizes, four patch-to-position mappings, and several dual-branch extensions.",
)
replace_prefix(
    doc,
    "The experiments answer the three research questions within this setting",
    "The experiments answer the three research questions within this setting. Positional information was important in full-data training, and learnable absolute PE achieved the highest core mean test accuracy. Shifted multiplicative PE was the strongest fixed design. It led learnable PE at 1,000 and 5,000 training examples, but the learned method had the higher mean at 10,000 examples and full data. Fixed encodings were also more sensitive to patch-to-position assignment. The selected methods kept the same test-accuracy order on CIFAR-100, while the dual-branch models did not improve on the learned single-branch reference.",
)
replace_prefix(
    doc,
    "The main contribution is therefore empirical rather than a new general-purpose ViT architecture",
    "The main contribution is therefore empirical rather than a new general-purpose ViT architecture. The study shows that PE choices can be compared under a shared protocol, that the amount of training data can change the ordering between learned and fixed PE, and that additional fusion complexity does not automatically improve classification. Learned PE is the stronger full-data default for this model, while shifted multiplicative PE is a credible option when labelled data are limited. These conclusions remain bounded by the low-resolution datasets, model scale, fixed data split and five-seed protocol.",
)

# Future Work: retain the completed hybrid only as preliminary evidence and state concrete next steps.
future_heading = paragraph_with_prefix(doc, "6.1 Future Work")
hybrid_intro = doc.add_paragraph(style="Normal")
replace_paragraph(
    hybrid_intro,
    "The reduced-data results motivate a more complete study of learned and fixed PE combinations. A preliminary model added an unshifted multiplicative encoding to the learnable table through one trainable global scale. Table 12 reports this initial attempt alongside the two components that motivated it. The hybrid remained close to learnable PE at every training size and did not provide a clear improvement. It is included here as evidence that the idea was explored, rather than as a main method or conclusion of this dissertation.",
)
future_heading._p.addnext(hybrid_intro._p)

caption = doc.add_paragraph(style="Caption")
replace_paragraph(caption, "Table 12. Preliminary learnable and multiplicative PE combination.")
hybrid_intro._p.addnext(caption._p)

future_table = doc.add_table(rows=1, cols=4)
headers = ["Training examples", "Learnable PE", "Shifted multiplicative PE", "Preliminary hybrid"]
for idx, value in enumerate(headers):
    future_table.rows[0].cells[idx].text = value
future_rows = [
    ("1,000", "37.47 ± 0.44%", "40.84 ± 0.85%", "37.45 ± 0.41%"),
    ("5,000", "55.18 ± 0.90%", "56.30 ± 0.51%", "55.06 ± 0.87%"),
    ("10,000", "63.13 ± 1.11%", "62.48 ± 1.04%", "63.31 ± 1.12%"),
    ("Full (45k)", "78.60 ± 0.42%", "78.08 ± 0.09%", "78.66 ± 0.51%"),
]
for values in future_rows:
    cells = future_table.add_row().cells
    for idx, value in enumerate(values):
        cells[idx].text = value
format_table(future_table, [1700, 2200, 2700, 2200])
caption._p.addnext(future_table._tbl)

hybrid_next = doc.add_paragraph(style="Normal")
replace_paragraph(
    hybrid_next,
    "The preliminary model does not test the most promising combination directly. Its fixed component was the unshifted multiplicative PE, while the reduced-data advantage was observed for shifted multiplicative PE. Future work should therefore test learnable PE with shifted multiplicative PE, compare different initial values for the scale, and replace the single global scale with layer-wise or dimension-wise gates. The study should also record the scale throughout training and verify that every fixed positional vector remains aligned with the physical patch coordinate under alternative unfolding patterns. These variants require a new, pre-defined experiment rather than further interpretation of the preliminary result.",
)
future_table._tbl.addnext(hybrid_next._p)

replace_prefix(
    doc,
    "The assignment and architecture studies can also be made more precise",
    "The assignment and architecture studies can also be made more precise. A factorial assignment experiment should compare reordering tokens alone, reordering positional vectors alone, and applying the same permutation to both. Fusion should be revisited only with parameter-matched single- and dual-branch models so that the fusion operation can be separated from model capacity.",
)

# Apply Arial and black text to new and modified content without disturbing equations.
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        set_run_font(run, size=run.font.size.pt if run.font.size else 11)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=run.font.size.pt if run.font.size else 9)

tmp = DOCX.with_suffix(".tmp.docx")
doc.save(tmp)
tmp.replace(DOCX)
print(DOCX)
