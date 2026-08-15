from __future__ import annotations

import csv
import hashlib
import os
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[2]
REFERENCE = Path(
    r"E:\xwechat_files\wxid_18ldxieih0qz22_77d6\msg\file\2026-08\Yikai_Zhao_MSc_Dissertation_Framework_v1.docx"
)
OUTPUT = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation_Aligned_v2.docx"
AGGREGATE = (
    ROOT
    / "results"
    / "cifar10_final_vit_models_5seeds"
    / "reports"
    / "thesis_core"
    / "aggregate_results.csv"
)

EXPECTED_REFERENCE_SHA256 = "74B2DFE5D920714A600FD3C8BF05D895D5378E25BD9999CD6F73FD40673A307C"

MODEL_LABELS = {
    "vit_baseline": "No positional encoding",
    "vit_learnable_position": "Learnable absolute PE",
    "vit_row_sinusoidal": "Row-only sinusoidal PE",
    "vit_col_sinusoidal": "Column-only sinusoidal PE",
    "vit_additive_sinusoidal": "Additive 2D sinusoidal PE",
    "vit_additive_sinusoidal_shifted": "Shifted additive 2D PE",
    "vit_multiplicative_sinusoidal": "Multiplicative 2D sinusoidal PE",
    "vit_multiplicative_sinusoidal_shifted": "Shifted multiplicative 2D PE",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest().upper()


def find_paragraph(document: Document, starts_with: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(starts_with)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {starts_with!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(document: Document, starts_with: str, text: str, label: str | None = None) -> Paragraph:
    paragraph = find_paragraph(document, starts_with)
    paragraph.clear()
    if label:
        run = paragraph.add_run(label)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)
    return paragraph


def insert_paragraph_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def copy_numbering(source: Paragraph, target: Paragraph) -> None:
    if source._p.pPr is None or source._p.pPr.numPr is None:
        return
    target_ppr = target._p.get_or_add_pPr()
    existing = target_ppr.find(qn("w:numPr"))
    if existing is not None:
        target_ppr.remove(existing)
    target_ppr.append(deepcopy(source._p.pPr.numPr))


def add_update_fields_on_open(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_twips: tuple[int, ...], indent_twips: int = 120) -> None:
    total = sum(widths_twips)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_twips))

    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths_twips):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_twips):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def insert_results_table(document: Document, after: Paragraph, rows: list[tuple[str, str, str]]) -> None:
    caption = insert_paragraph_after(
        after,
        "Table 2. Current verified Tier 1 CIFAR-10 evidence. Accuracy is the mean selected-checkpoint test accuracy across seeds 42–46; SD is the sample standard deviation in percentage points.",
        "Caption",
    )
    table = document.add_table(rows=1, cols=3)
    table.style = document.tables[0].style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (Cm(10.2), Cm(2.7), Cm(2.7))
    widths_twips = (5900, 1563, 1563)
    headers = ("Configuration", "Mean accuracy (%)", "Sample SD (pp)")
    for index, (cell, value, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = value
        set_cell_shading(cell, "2F5597")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8.5)
    set_repeat_table_header(table.rows[0])

    for label, mean, sd in rows:
        cells = table.add_row().cells
        values = (label, mean, sd)
        for index, (cell, value, width) in enumerate(zip(cells, values, widths)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
    set_table_geometry(table, widths_twips)
    caption._p.addnext(table._tbl)


def load_core_results() -> list[tuple[str, str, str]]:
    with AGGREGATE.open(newline="", encoding="utf-8") as handle:
        rows = {row["model"]: row for row in csv.DictReader(handle)}
    missing = set(MODEL_LABELS) - set(rows)
    if missing:
        raise RuntimeError(f"Missing aggregate rows: {sorted(missing)}")
    return [
        (
            label,
            f"{100 * float(rows[model]['mean_acc']):.3f}",
            f"{100 * float(rows[model]['sd_acc']):.3f}",
        )
        for model, label in MODEL_LABELS.items()
    ]


def patch_content(document: Document) -> None:
    replacements: list[tuple[str, str]] = [
        (
            "Evaluating Positional Encoding and Patch-to-Position Assignment",
            "Positional Encoding in Compact Vision Transformers: A Controlled Evaluation",
        ),
        (
            "A controlled image-classification study",
            "Fixed and learnable 2D encodings, patch-to-position assignment, and exploratory extensions on CIFAR-10",
        ),
        ("Document status:", "Document status: Aligned framework v2 — 5 August 2026"),
        (
            "Generic planning template",
            "UCL MSc dissertation planning document — confirm only programme-specific submission details with the module handbook and supervisor.",
        ),
        (
            "This is an evidence-led planning document",
            "This is an evidence-led MSc dissertation framework rather than a submission-ready thesis. It uses Framework v1 as the structural and visual authority, incorporates the strongest reusable prose from the earlier structured draft, and aligns all numerical statements with the locally verified final CIFAR-10 bundle. No figures are embedded; future figure locations are described explicitly. The suggested 10,000–12,000-word range is an internal planning aid only, not a claimed UCL requirement.",
        ),
        (
            "RESULT SYNC REQUIRED  The repository currently contains useful CIFAR-10 evidence",
            "LOCAL EVIDENCE VERIFIED  The local bundle contains 32 configurations and 160 selected-checkpoint summaries: five training seeds (42–46) for every model, with split seed 42 fixed throughout. Synchronisation is required only if another computer contains a newer bundle that supersedes these files.",
        ),
        (
            "RESULT SYNC REQUIRED — marks a dependency",
            "EVIDENCE STATUS — distinguishes locally verified results from experiments or external bundles that still require checking.",
        ),
        ("Synchronise the final experiment outputs", "Methodology — write directly from the model registry, PE implementations, data split code, run configurations and checkpoint-selection protocol."),
        ("Write Methodology and Experiments", "Experiments and Results — freeze the evidence tiers, then report the verified multi-seed outcomes without mechanism claims."),
        ("Write Results without interpretation", "Discussion — answer the three research questions, distinguish observations from explanations, and state boundary conditions."),
        ("Rewrite Background and Related Work", "Literature Review — retain only literature that defines the design space or supports interpretation of the tested factors."),
        ("Write Introduction, Abstract and Conclusion last", "Introduction — write the motivation, gap, contributions and chapter map after the evidence-led chapters stabilise."),
        (
            "Starter idea: Vision Transformers require positional information",
            "Provisional abstract: Vision Transformers process images as patch-token sequences, yet content-only self-attention does not identify the two-dimensional origin of each patch. This dissertation presents a controlled empirical evaluation of positional encoding in a compact Vision Transformer trained from scratch on CIFAR-10. It compares no positional encoding, learnable absolute embeddings, fixed row- and column-based sinusoidal constructions, additive and multiplicative two-dimensional variants, coordinate-shifted variants, patch-to-position assignments, and exploratory hybrid and dual-branch fusion models. All locally consolidated configurations are evaluated across five training seeds under a common architecture, fixed data split and validation-selected checkpoint protocol. The no-position baseline is clearly weaker than the leading positional models. Learnable absolute encoding remains a strong reference, while shifted multiplicative encoding is the strongest tested fixed positional encoding. Learnable models vary little across the tested patch assignments, whereas several fixed encodings are more sensitive to the mapping between image patches and positional vectors. A hybrid model attains the highest numerical mean, but its small margin and traversal confound do not support a general optimisation claim; similarly, no dual-branch fusion model exceeds the best single-branch or hybrid result. The study therefore contributes a reproducible empirical map of positional-design choices and their limitations rather than a universally superior architecture. Generalisation beyond CIFAR-10 remains to be established.",
        ),
        (
            "This dissertation provides a controlled empirical study",
            "This dissertation provides a controlled empirical study of how absent, learnable and fixed two-dimensional positional encodings—and their interaction with patch-to-position assignment—affect compact Vision Transformer image classification under a common protocol.",
        ),
        (
            "What: compare positional encoding families",
            "What: compare positional encoding families within the same compact ViT, CIFAR-10 split, optimisation settings and selected-checkpoint evaluation protocol, then test representative methods under alternative patch-to-position assignments.",
        ),
        (
            "So what: the experiments identify which fixed 2D constructions",
            "So what: the experiments quantify which fixed 2D constructions recover most of the performance associated with learnable absolute embeddings, how stable the differences are across seeds, and where patch-to-position assignment changes the result.",
        ),
        (
            "RQ1. Under a controlled CIFAR-10 protocol",
            "RQ1. Under a controlled protocol, how do absent, learnable absolute and fixed two-dimensional positional encodings affect classification accuracy and run-to-run stability?",
        ),
        (
            "RQ2. Which properties of the tested fixed encodings",
            "RQ2. How do row- and column-based constructions—including axis-specific, additive, multiplicative and wavelength-shifted variants—differ under the shared compact-ViT setting?",
        ),
        (
            "RQ3. How does patch-to-position assignment",
            "RQ3. How do row-major, column-major and alternative unfolding conventions change patch-to-position assignment, how do they interact with fixed positional encoding, and do the principal trends generalise across datasets or split conditions?",
        ),
        (
            "Do claim that positional encoding improves",
            "Do claim that positional information is important for this compact CIFAR-10 ViT: every leading positional model substantially exceeds the no-position baseline in the verified five-seed comparison.",
        ),
        (
            "Do not claim that a custom fixed encoding beats",
            "Do not claim that a custom fixed encoding generally beats learnable absolute PE. Shifted multiplicative encoding is the strongest tested fixed design, but learned absolute PE has the higher mean in the matched core comparison.",
        ),
        (
            "Treat single-seed squared, radial, low-data, fusion and hybrid results",
            "Treat squared, radial, hybrid and fusion variants as secondary or exploratory evidence. Five-seed coverage is now available locally, but the hybrid changes traversal and the fusion models change capacity, so neither isolates a single causal factor.",
        ),
        (
            "A practical generic allocation",
            "The allocation below is an internal planning aid for a substantial MSc dissertation. It is not presented as a UCL word-count requirement; the definitive module rules and supervisor guidance take precedence.",
        ),
        ("Background and Related Work", "Literature Review"),
        (
            "Document image size, patch size, number of patches",
            "The verified CIFAR-10 runs use 32 × 32 RGB inputs, 4 × 4 non-overlapping patches (an 8 × 8 grid), embedding dimension 128, four Transformer blocks, four attention heads, MLP hidden dimension 512 and zero embedding, attention, projection and MLP dropout. Record the resulting parameter count for each model family before final submission.",
        ),
        (
            "For CIFAR-10 and each additional dataset",
            "For the current CIFAR-10 study, report 45,000 training images, 5,000 validation images and the official 10,000-image test set; split seed 42 is held fixed while training seeds vary from 42 to 46. Document normalisation and augmentation directly from the data pipeline. For every additional dataset, report its provenance, preprocessing, split construction and leakage controls separately.",
        ),
        (
            "Specify optimiser, learning rate, weight decay",
            "The verified runs use AdamW, learning rate 3 × 10⁻⁴, weight decay 0.05, batch size 128 and at most 100 epochs. ReduceLROnPlateau monitors validation evidence; early stopping monitors validation accuracy with patience 10 and minimum delta 0.001. The selected validation checkpoint is loaded and evaluated once on the test set, recorded as test_evaluation_protocol=selected_checkpoint_only.",
        ),
        (
            "EVIDENCE REQUIRED  Some existing CIFAR-10 reports contain per-epoch test metrics",
            "LOCAL PROTOCOL VERIFIED  All 160 consolidated summaries record test_evaluation_protocol=selected_checkpoint_only. Preserve this rule for any new dataset or robustness run; do not use the test set for checkpoint selection or experiment pruning.",
        ),
        (
            "Primary confirmatory experiment: compare no PE",
            "Tier 1 confirmatory experiment: compare no PE, learnable absolute PE, row-only, column-only, additive, shifted additive, multiplicative and shifted multiplicative PE under the same CIFAR-10 architecture, split seed 42, training settings and seeds 42–46. The local bundle is complete for all eight conditions.",
        ),
        (
            "RESULT SYNC REQUIRED  Replace all preliminary summaries",
            "LOCAL EVIDENCE VERIFIED  Regenerate thesis tables from results/cifar10_final_vit_models_5seeds/reports/thesis_core rather than copying values manually. If a newer external bundle is synchronised, archive the current manifest and compare checksums before replacing it.",
        ),
        (
            "Test squared, radial, rotary or other completed extensions",
            "Tier 3 extensions: squared multiplicative and radial variants have five-seed local results, while any rotary or newly added variant must be checked separately. Keep these extensions outside the headline Tier 1 table because they broaden the search after the primary comparison.",
        ),
        (
            "Compare row-major, column-major and the final split/unfolding strategies",
            "Tier 2 assignment experiment: compare normal_row, normal_col, proper_row and proper_col within no-PE, learnable, row-only, column-only and multiplicative families. Five-seed results are available locally. Retain no-PE and learnable controls so that sensitivity can be attributed cautiously to the patch-to-fixed-PE mapping rather than to 1D adjacency alone.",
        ),
        (
            "Repeat a compact subset of the strongest baselines",
            "Tier 2 generalisation experiment to complete: pre-select no PE, learnable absolute PE, the strongest additive design and the strongest multiplicative design, then repeat them on at least one additional dataset or independently constructed split using the same reporting protocol.",
        ),
        (
            "Hybrid learnable + fixed PE:",
            "Hybrid learnable + fixed PE: five-seed evidence is available, but an order-matched ablation and learned mixing-coefficient logs are required before attributing the numerical difference to the fixed component.",
        ),
        (
            "Fusion variants: place in an appendix",
            "Fusion variants: five-seed evidence is available, but dual encoders change parameter count and compute. Keep the study exploratory or appendical until capacity and cost are reported or matched.",
        ),
        (
            "Current repository evidence (preliminary, pending synchronisation)",
            "The locally verified Tier 1 means ± sample SD are: learnable absolute PE 78.602 ± 0.339%, shifted multiplicative PE 78.082 ± 0.076%, multiplicative PE 77.536 ± 0.668%, shifted additive PE 77.148 ± 0.077%, additive PE 76.978 ± 0.611%, row-only PE 74.928 ± 0.701%, column-only PE 74.242 ± 0.567%, and no PE 71.288 ± 0.672%. These are selected-checkpoint test accuracies across five training seeds; the percent sign denotes accuracy and the SD values are percentage points.",
        ),
        (
            "EVIDENCE REQUIRED  Verify these values",
            "EVIDENCE STATUS  The values above are verified against the local aggregate and per-seed reports. Before submission, record a checksum or immutable manifest for the final bundle and include paired seed-level differences in the appendix.",
        ),
        (
            "Starter results style: Across five seeds",
            "Starter results style: Across five seeds, all seven tested positional encodings exceeded the no-position baseline in mean test accuracy. Learnable absolute PE produced the strongest mean within the pre-specified core comparison. Shifted multiplicative PE was the strongest fixed design and showed low cross-seed dispersion. These observations establish the ranking only for the evaluated compact ViT, CIFAR-10 split and training protocol; causal explanations belong in the Discussion chapter.",
        ),
        (
            "The current single-seed records include squared multiplicative",
            "The five-seed secondary results are 77.606 ± 0.660% for shifted squared multiplicative PE, 77.128 ± 0.416% for squared multiplicative PE and 75.644 ± 0.486% for radial PE. They do not exceed shifted multiplicative PE and should be presented as extensions rather than folded into the primary eight-condition claim.",
        ),
        (
            "In the current seed-42 unfolding report",
            "Across the four tested assignments, learnable absolute PE means span only 78.576–78.638%, while no-PE means span 71.288–71.512%. Fixed families show larger and method-dependent changes; for example, multiplicative PE ranges from 74.834% to 77.536%. Report the full interaction table and paired seed results, then interpret these differences as patch-to-position assignment effects only after the index mapping is verified.",
        ),
        (
            "EVIDENCE REQUIRED  Repeat the assignment comparison across seeds",
            "EVIDENCE STATUS  The four-assignment comparison is complete across five seeds. The remaining requirement is an explicit unit or synthetic test that documents, for each unfolding mode, the physical patch index, sequence slot and fixed positional vector that are paired.",
        ),
        (
            "For each additional dataset, report multi-seed",
            "[RESULT TO VERIFY] Cross-dataset or independent-split evidence is not present in the local final bundle. Report multi-seed selected-checkpoint accuracy and uncertainty only after the pre-selected representative models have been completed.",
        ),
        (
            "If confirmed, argue that positional information materially improves",
            "The verified CIFAR-10 evidence supports the narrow conclusion that positional information materially improves this compact ViT relative to no PE. Keep this broad result separate from the method ranking: learnable absolute PE remains the strongest pre-specified core reference.",
        ),
        (
            "Discuss the comparatively strong multiplicative row/column coupling",
            "Discuss shifted multiplicative row/column coupling as the strongest tested fixed pattern, not as proof of a universal geometric advantage. Its low sample SD is descriptive evidence of stability under these five seeds; cross-dataset evidence is still required.",
        ),
        (
            "Seed coverage: the main experiment has five seeds",
            "Seed coverage: all 32 models in the current consolidated CIFAR-10 bundle have five training seeds, but five runs remain insufficient for strong inference about very small differences, and cross-dataset/split evidence is absent.",
        ),
        (
            "Starter idea: This dissertation examined positional encoding",
            "Provisional conclusion starter: This dissertation examined positional encoding as a controlled design choice in compact Vision Transformers. Across 32 CIFAR-10 configurations and five training seeds per configuration, positional information consistently improved the leading models over the no-position baseline. Learnable absolute PE remained the strongest core reference, while shifted multiplicative PE was the strongest tested fixed encoding. Alternative patch-to-position assignments had little effect on learnable PE but interacted more strongly with several fixed encodings. Hybrid and fusion studies provide useful boundary evidence, yet their confounds prevent them from supporting headline optimisation claims. The conclusions remain specific to the evaluated architecture, dataset, split and training protocol until the planned generalisation experiments are complete.",
        ),
        (
            "RESULT SYNC REQUIRED  This inventory is a writing map",
            "EVIDENCE STATUS  This inventory is aligned with the local final bundle dated 5 August 2026. Treat it as provisional only if a newer bundle exists on another computer; synchronise by checksum and regenerate all derived summaries rather than merging values manually.",
        ),
        (
            "Headline evidence: results/cifar10_positional_8models_5seeds",
            "Tier 1 headline evidence: results/cifar10_final_vit_models_5seeds/reports/thesis_core/aggregate_results.csv and per_seed_results.csv. Eight core PE conditions, five seeds, fixed split seed 42 and selected-checkpoint-only test evaluation.",
        ),
        ("Squared extensions: single-seed result reports", "Squared and radial extensions: five-seed result reports are available in the final bundle. Keep them secondary because they extend the pre-specified core family."),
        ("Radial PE: seed-42 selected-checkpoint result", "Patch assignment: five-seed results cover normal_row, normal_col, proper_row and proper_col for representative PE families. Add a mapping correctness test before causal wording."),
        ("Unfolding/assignment: seed-42", "Hybrid: five-seed results are available, but the current hybrid changes traversal; add an order-matched ablation and log the learned mixing coefficient."),
        ("Low-data regimes: 1k/5k/10k/full", "Fusion: five-seed results are available for five dual-branch variants. Report parameters and compute; do not compare efficiency until capacity is controlled."),
        ("Hybrid and fusion: single-seed", "Generalisation gap: no completed additional-dataset or independent-split evidence is included in the local final bundle. This is the highest-priority experiment family still missing."),
        ("CADB Elements: exploratory multi-label outputs", "CADB Elements and low-data runs remain exploratory. Keep them outside the main thesis argument unless their labels, metrics, balance and multi-seed protocol are independently resolved."),
        (
            "☐ Final runs from the other computer have been copied",
            "☐ Any external result bundle newer than the local manifest has been copied, checksummed, indexed and reconciled before derived tables are regenerated.",
        ),
    ]

    for starts_with, text in replacements:
        replace_paragraph(document, starts_with, text)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text.startswith("FIGURE PLACEHOLDER"):
            continue
        if "indicates a useful future visual" in text:
            paragraph.text = "[FIGURE HERE: concise description] marks a planned visual. No figures are embedded in this version."
        else:
            description = text.removeprefix("FIGURE PLACEHOLDER").lstrip(" —-")
            paragraph.text = f"[FIGURE HERE: {description}]"

    # Add the final two stages of the user's preferred drafting sequence.
    anchor = find_paragraph(document, "Introduction — write the motivation")
    conclusion_step = insert_paragraph_after(anchor, "Conclusion — answer the three research questions without introducing new evidence.", "Normal")
    copy_numbering(anchor, conclusion_step)
    abstract_step = insert_paragraph_after(conclusion_step, "Abstract — revise the provisional abstract last so every claim matches the completed thesis.", "Normal")
    copy_numbering(anchor, abstract_step)
    anchor = abstract_step

    # Make the alignment choice explicit without turning it into another chapter.
    anchor = find_paragraph(document, "Abstract — revise the provisional abstract")
    anchor = insert_paragraph_after(anchor, "Alignment with the Existing Structured Draft", "Heading 2")
    anchor = insert_paragraph_after(
        anchor,
        "Framework v1 is retained as the stronger planning structure. The earlier structured draft contributes its verified multi-seed results, useful provisional prose and complete-model appendix logic. The aligned version removes the fourth research question, recentres the dissertation on three evidence-backed questions, and downscopes hybrid and fusion work to exploratory evidence.",
        "Normal",
    )
    insert_paragraph_after(
        anchor,
        "DRAFTING NOTE  Expand full prose in the recommended order. Do not treat this framework's compact starter paragraphs as final chapter length.",
        "Callout",
    )

    results_anchor = find_paragraph(document, "Starter results style:")
    insert_results_table(document, results_anchor, load_core_results())

    # Make the core-results table easier to identify in Word's navigation/accessibility tools.
    document.tables[-1].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Preserve the source header/footer and design authority; update only the running title.
    for section in document.sections:
        for header in (section.header, section.first_page_header):
            for paragraph in header.paragraphs:
                if "MSc Dissertation Framework" in paragraph.text:
                    paragraph.text = "MSc Dissertation Framework | Controlled Evaluation of Positional Encoding"

    document.core_properties.title = "Positional Encoding in Compact Vision Transformers: A Controlled Evaluation"
    document.core_properties.subject = "Aligned MSc dissertation framework"
    document.core_properties.author = "Yikai Zhao"
    document.core_properties.comments = "Generated from Framework v1; local five-seed evidence aligned; no figures embedded."
    # The source already contains a PAGE field. Do not force a full field refresh
    # on open: in some desktop Word installations that can block headless export.


def main() -> None:
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    if sha256(REFERENCE) != EXPECTED_REFERENCE_SHA256:
        raise RuntimeError("Reference document changed since the template audit; refusing to build from an unverified source.")
    if not AGGREGATE.exists():
        raise FileNotFoundError(AGGREGATE)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    os.chmod(OUTPUT, 0o666)
    document = Document(OUTPUT)
    patch_content(document)
    document.save(OUTPUT)

    if sha256(REFERENCE) != EXPECTED_REFERENCE_SHA256:
        raise RuntimeError("Reference document was modified during generation.")
    print(OUTPUT)


if __name__ == "__main__":
    main()
