from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
THESIS_DIR = ROOT / "thesis"
OUTPUT = THESIS_DIR / "Yikai_Zhao_MSc_Dissertation.docx"
TEMP_OUTPUT = THESIS_DIR / ".Yikai_Zhao_MSc_Dissertation.building.docx"
REPORT_DIR = ROOT / "results" / "cifar10_final_vit_models_5seeds" / "reports" / "thesis_core"
SUMMARY_CSV = REPORT_DIR / "selected_test_summary_with_ci.csv"
CORE_CSV = REPORT_DIR / "core_pe_with_radial_test_summary.csv"
PAIRED_CSV = REPORT_DIR / "key_paired_test_contrasts.csv"


LABELS = {
    "vit_baseline": "No PE",
    "vit_learnable_position": "Learned absolute PE",
    "vit_row_sinusoidal": "Row sinusoidal",
    "vit_col_sinusoidal": "Column sinusoidal",
    "vit_additive_sinusoidal": "Additive 2D sinusoidal",
    "vit_additive_sinusoidal_shifted": "Shifted additive 2D",
    "vit_multiplicative_sinusoidal": "Multiplicative 2D sinusoidal",
    "vit_multiplicative_sinusoidal_shifted": "Shifted multiplicative 2D",
    "vit_radial_sinusoidal": "Radial sinusoidal",
    "vit_squared_multiplicative_sinusoidal": "Squared multiplicative",
    "vit_squared_multiplicative_sinusoidal_shifted": "Shifted squared multiplicative",
    "vit_normal_col_learnable_multiplicative_sinusoidal": "Hybrid learned + fixed multiplicative",
    "vit_row_col_latent_fusion": "Latent concatenation fusion",
    "vit_row_col_mean_fusion": "Mean fusion",
    "vit_row_col_mean_mlp_fusion": "Mean + MLP fusion",
    "vit_row_col_cross_attention_fusion": "Cross-attention fusion",
    "vit_row_col_cross_attention_mlp_head_fusion": "Cross-attention + MLP head",
}


REFERENCES = [
    "A. Vaswani et al., ‘Attention Is All You Need,’ in Advances in Neural Information Processing Systems 30, 2017.",
    "A. Dosovitskiy et al., ‘An Image Is Worth 16×16 Words: Transformers for Image Recognition at Scale,’ in International Conference on Learning Representations, 2021.",
    "H. Touvron, M. Cord, M. Douze, F. Massa, A. Sablayrolles, and H. Jégou, ‘Training Data-Efficient Image Transformers & Distillation through Attention,’ in Proceedings of the 38th International Conference on Machine Learning, pp. 10347–10357, 2021.",
    "P. Dufter, M. Schmitt, and H. Schütze, ‘Position Information in Transformers: An Overview,’ Computational Linguistics, vol. 48, no. 3, pp. 733–763, 2022, doi: 10.1162/coli_a_00445.",
    "P. Shaw, J. Uszkoreit, and A. Vaswani, ‘Self-Attention with Relative Position Representations,’ in Proceedings of NAACL-HLT, vol. 2, pp. 464–468, 2018, doi: 10.18653/v1/N18-2074.",
    "K. Wu, H. Peng, M. Chen, J. Fu, and H. Chao, ‘Rethinking and Improving Relative Position Encoding for Vision Transformer,’ in Proceedings of the IEEE/CVF International Conference on Computer Vision, pp. 10033–10041, 2021.",
    "X. Chu, Z. Tian, B. Zhang, X. Wang, and C. Shen, ‘Conditional Positional Encodings for Vision Transformers,’ in International Conference on Learning Representations, 2023.",
    "Y. Li, S. Si, G. Li, C.-J. Hsieh, and S. Bengio, ‘Learnable Fourier Features for Multi-Dimensional Spatial Positional Encoding,’ in Advances in Neural Information Processing Systems 34, 2021.",
    "B. Heo, S. Park, D. Han, and S. Yun, ‘Rotary Position Embedding for Vision Transformer,’ in Computer Vision – ECCV 2024, LNCS 15068, pp. 289–305, 2024, doi: 10.1007/978-3-031-72684-2_17.",
    "M. A. M. Chowdhury, M. R. U. Rahman, and A. A. Taki, ‘LOOPE: Learnable Optimal Patch Order in Positional Embeddings for Vision Transformers,’ arXiv:2504.14386, 2025, preprint.",
    "D. Kutscher, D. M. Chan, Y. Bai, T. Darrell, and R. Gupta, ‘REOrdering Patches Improves Vision Models,’ arXiv:2505.23751, 2025, preprint.",
    "C.-F. R. Chen, Q. Fan, and R. Panda, ‘CrossViT: Cross-Attention Multi-Scale Vision Transformer for Image Classification,’ in Proceedings of the IEEE/CVF International Conference on Computer Vision, pp. 357–366, 2021.",
    "A. Krizhevsky, ‘Learning Multiple Layers of Features from Tiny Images,’ University of Toronto, Technical Report, 2009.",
    "J. Demšar, ‘Statistical Comparisons of Classifiers over Multiple Data Sets,’ Journal of Machine Learning Research, vol. 7, pp. 1–30, 2006.",
]


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def black_arial(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_body(doc: Document, text: str) -> None:
    for block in [part.strip() for part in text.strip().split("\n\n") if part.strip()]:
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.add_run(" ".join(line.strip() for line in block.splitlines()))


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered_heading(doc: Document, number: str, title: str, level: int = 1) -> None:
    if level == 1 and doc.paragraphs:
        doc.add_page_break()
    doc.add_paragraph(f"{number} {title}", style=f"Heading {level}")


def add_caption(doc: Document, number: int, title: str) -> None:
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.add_run(f"Table {number}. {title}")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, "D9E2F3")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        black_arial(run, size=8.5, bold=True)
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            black_arial(run, size=8.3)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    heading_sizes = {1: 17, 2: 13, 3: 11}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if level > 1 else 0)
        style.paragraph_format.space_after = Pt(8)

    caption = styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(9)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(3)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.different_first_page_header_footer = True

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    black_arial(run, size=9)


def add_title_page(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(40)
    run = paragraph.add_run("UNIVERSITY COLLEGE LONDON")
    black_arial(run, 14, True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(50)
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run("Positional Encoding in Compact Vision Transformers:\nA Controlled Evaluation")
    black_arial(run, 22, True)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(55)
    run = paragraph.add_run("MSc Dissertation")
    black_arial(run, 14)

    for line in (
        "MSc Scientific and Data Intensive Computing",
        "Yikai Zhao",
        "Student Number: 25200353",
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        black_arial(run, 12)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(70)
    run = paragraph.add_run("Draft prepared August 2026")
    black_arial(run, 11)
    doc.add_page_break()


def format_acc(row: dict[str, str]) -> str:
    return f"{float(row['mean_test_acc_pct']):.3f} ± {float(row['ci95_half_width_test_acc_pp']):.3f}"


def format_loss(row: dict[str, str]) -> str:
    return f"{float(row['mean_test_loss']):.4f} ± {float(row['ci95_half_width_test_loss']):.4f}"


def build_document() -> Document:
    summary_rows = read_csv(SUMMARY_CSV)
    core_rows = read_csv(CORE_CSV)
    paired_rows = read_csv(PAIRED_CSV)
    summary = {row["model"]: row for row in summary_rows}
    paired = {row["contrast"]: row for row in paired_rows}

    baseline = summary["vit_baseline"]
    learned = summary["vit_learnable_position"]
    shifted_mult = summary["vit_multiplicative_sinusoidal_shifted"]
    hybrid = summary["vit_normal_col_learnable_multiplicative_sinusoidal"]
    fusion = summary["vit_row_col_cross_attention_mlp_head_fusion"]

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_page(doc)

    doc.add_paragraph("Abstract", style="Heading 1")
    add_body(doc, f"""
        Vision Transformers convert two-dimensional images into patch-token sequences, but content-only self-attention does not identify the spatial origin of each patch. This dissertation presents a controlled empirical evaluation of positional encoding in a compact Vision Transformer trained from scratch. The study compares the absence of positional encoding, learned absolute embeddings, fixed row and column sinusoidal signals, additive, multiplicative, shifted, radial and squared two-dimensional constructions, four patch-to-position assignments, an exploratory learned–fixed hybrid, and five dual-branch fusion models. The consolidated CIFAR-10 experiment contains 32 configurations evaluated with five training seeds, a common data split and architecture, validation-based checkpoint selection, and a single holdout-test evaluation per run. Learned absolute positional encoding achieves {float(learned['mean_test_acc_pct']):.3f}% mean test accuracy, compared with {float(baseline['mean_test_acc_pct']):.3f}% without positional encoding. Shifted multiplicative encoding is the strongest tested fixed design at {float(shifted_mult['mean_test_acc_pct']):.3f}%. Order-matched learned and hybrid models differ by only {float(paired['order_matched_learnable_to_hybrid']['mean_test_acc_delta_pp']):.3f} percentage points on average, while the best fusion model remains below the learned single-branch reference despite its larger parameter count. These results show that positional information is important in this compact setting, that fixed two-dimensional constructions interact with patch-to-position assignment, and that additional hybrid or fusion complexity does not automatically translate into a commensurate accuracy gain. The conclusions are limited to the evaluated architecture and CIFAR-10 evidence; low-data and cross-dataset evaluations are specified as the next tests of generality.
    """)

    add_numbered_heading(doc, "1", "Introduction")
    doc.add_paragraph("1.1 Background and Motivation", style="Heading 2")
    add_body(doc, """
        The Transformer replaced recurrent sequence processing with attention mechanisms that allow every token to interact directly with every other token [1]. Vision Transformer (ViT) adapted this architecture to images by dividing an image into patches, mapping the patches to embeddings, prepending a classification token, and processing the resulting sequence with Transformer encoder blocks [2]. This formulation is conceptually simple, but it creates an immediate representational question: an image is two-dimensional, whereas the encoder receives a one-dimensional list of tokens.

        Self-attention without a positional signal is permutation-equivariant. If the input patch tokens are permuted consistently, the corresponding token outputs are permuted in the same way. The operation can model similarities among patch contents, but it cannot determine whether a patch originated above, below, left or right of another patch. ViT implementations therefore inject position through learned tables, deterministic functions, relative attention terms or content-conditioned mechanisms [2], [4]–[9]. The number of alternatives shows that position is not a minor formatting detail; it is part of the model’s inductive bias.

        Published positional-encoding results are difficult to compare directly because architecture size, pretraining data, augmentation, optimisation and evaluation protocol often change together. This dissertation narrows the question to a compact ViT trained from scratch under one controlled pipeline. The aim is not to introduce a universally superior Transformer architecture. It is to determine which empirical differences survive when the backbone, data split, optimiser, checkpoint rule and test protocol are held constant.
    """)
    doc.add_paragraph("1.2 Research Questions", style="Heading 2")
    add_bullets(doc, [
        "RQ1: Is positional encoding necessary for the evaluated compact ViT, and how does learned absolute encoding compare with fixed alternatives?",
        "RQ2: How do row and column signals, their additive or multiplicative combination, wavelength shifts, and radial or squared transformations affect performance?",
        "RQ3: How does the mapping from physical image patches to sequence slots and positional vectors interact with fixed and learned encodings?",
        "RQ4: Do an exploratory learned–fixed hybrid or dual-branch row–column fusion provide gains commensurate with their additional complexity?",
        "RQ5: Do the principal CIFAR-10 findings persist when labelled training data are reduced and when the dataset changes?",
    ])
    doc.add_paragraph("1.3 Objectives and Contributions", style="Heading 2")
    add_body(doc, """
        The practical objectives are to implement positional variants behind a shared model interface, verify patch-to-position mappings deterministically, run multi-seed comparisons, and report selected-checkpoint test accuracy and loss with uncertainty. The current completed evidence contributes a 160-run CIFAR-10 comparison covering 32 configurations and five seeds per configuration; a mapping test that explicitly records physical patch coordinate, sequence slot and assigned positional coordinate; and capacity-aware analyses of hybrid and dual-branch extensions. The low-data and second-dataset protocols are defined but are not presented as completed contributions until their five-seed result sets pass the same aggregation checks.
    """)
    doc.add_paragraph("1.4 Scope and Dissertation Structure", style="Heading 2")
    add_body(doc, """
        The scope is image classification with a compact ViT and absolute input-level positional signals, rather than a comprehensive benchmark of every relative or rotary attention mechanism. Chapter 2 reviews the literature needed to motivate the comparison. Chapter 3 defines the dataset, model, positional formulas, mapping conditions and statistical protocol. Chapter 4 presents the experiments in predefined suites. Chapter 5 answers the research questions, separates observations from possible mechanisms, and states threats to validity. Chapter 6 concludes with the most reliable claims and their boundaries.
    """)

    add_numbered_heading(doc, "2", "Literature Review")
    doc.add_paragraph("2.1 Self-Attention and Position Information", style="Heading 2")
    add_body(doc, """
        For an input matrix X, scaled dot-product attention forms query, key and value projections and computes Attention(Q,K,V) = softmax(QKᵀ/√dₖ)V [1]. Multi-head attention applies this operation in parallel representation subspaces, followed by residual connections, normalisation and a position-wise feed-forward network. If no positional signal is included, applying the same permutation to the rows of X permutes the attention outputs correspondingly. The relevant property is permutation equivariance, not simple invariance, because token-level outputs move with their inputs. Dufter et al. survey the ways Transformer models inject absolute, relative and contextual position information [4].

        The original Transformer added fixed sine and cosine functions at multiple wavelengths [1]. A deterministic encoding provides a smoothly varying representation without allocating one trainable vector to each position. Shaw et al. instead introduced relative position representations inside attention, making pairwise offsets part of the compatibility and aggregation calculations [5]. These constructions establish the two main perspectives used later in vision: encode each token’s absolute location or modify token interactions according to relative geometry.
    """)
    doc.add_paragraph("2.2 Vision Transformers and Data Scale", style="Heading 2")
    add_body(doc, """
        ViT maps non-overlapping image patches to tokens and uses a learned absolute embedding for each sequence position [2]. Its original large-scale results relied on substantial pretraining, while DeiT demonstrated that a carefully designed training recipe and distillation can make image Transformers more data efficient [3]. These studies motivate two aspects of the present work. First, learned absolute PE is a strong conventional reference. Second, results from large pretrained models cannot be assumed to predict behaviour for a compact model trained from scratch on CIFAR-10.

        CIFAR-10 and CIFAR-100 contain 32×32 colour images and share the same broad acquisition source, while differing in class granularity [13]. Their common size makes CIFAR-100 a controlled second-dataset option: the patch grid and model can remain unchanged while the classification problem becomes more fine-grained. This comparability is preferable to mixing an unrelated multi-label task into the main single-label evidence.
    """)
    doc.add_paragraph("2.3 Absolute, Relative, Conditional and Two-Dimensional Encoding", style="Heading 2")
    add_body(doc, """
        Vision introduces a two-dimensional coordinate (r,c) for each patch. A one-dimensional learned table can associate every flattened slot with a spatial role, but it does not explicitly factor row and column. Fixed two-dimensional schemes make that factorisation visible. An encoding may use only r or c, add separate axis signals, multiply them element-wise, concatenate subspaces, or transform the coordinate pair before applying a sinusoidal basis. Each choice discards or preserves different information, so no universal ordering follows from the formula alone.

        Image RPE adapts relative-position modelling to directional two-dimensional distances [6]. CPVT generates positional information conditionally from local token neighbourhoods [7]. Learnable Fourier features map multidimensional coordinates through trainable Fourier projections [8]. RoPE-ViT applies rotary position embedding to two-dimensional vision data and studies resolution extrapolation [9]. These methods demonstrate a broad design space, but they also alter different parts of the architecture. The present comparison is deliberately narrower: it keeps the encoder blocks fixed and changes the positional vector added to patch tokens.
    """)
    doc.add_paragraph("2.4 Patch Ordering and Coordinate Assignment", style="Heading 2")
    add_body(doc, """
        Converting a patch grid to a sequence requires an ordering convention, commonly row-major raster order. Recent work has treated patch order itself as a learnable or searchable design variable. LOOPE proposes learning an ordering for positional embeddings [10], while REOrder studies alternative and learned patch orders in architectures whose long-sequence approximations introduce order sensitivity [11]. Both are 2025 preprints and are used here as recent context rather than settled evidence.

        A full-attention ViT requires a more precise distinction. Reordering both patch tokens and their matched positional vectors is a joint permutation and should preserve the functional problem up to optimisation noise. By contrast, keeping a sequence order while changing which coordinate vector is assigned to each physical patch changes the positional content presented to the model. The present study therefore records three objects separately: physical patch coordinate, sequence slot and assigned PE coordinate. This implementation-level distinction is central to interpreting the order experiments.
    """)
    doc.add_paragraph("2.5 Dual-Branch Fusion", style="Heading 2")
    add_body(doc, """
        CrossViT provides a prominent example of a dual-branch vision Transformer in which tokens at different patch scales are processed separately and exchanged through cross-attention [12]. It establishes a precedent for branch-specific encoders and attention-based fusion, but the present fusion models are not claimed as a reproduction of CrossViT or as an original general architecture. They are exploratory tests of whether separate row- and column-oriented streams contain complementary information. Because they duplicate substantial encoder capacity, accuracy must be interpreted jointly with parameter count.
    """)
    doc.add_paragraph("2.6 Positioning of This Dissertation", style="Heading 2")
    add_body(doc, """
        Prior work proposes increasingly expressive positional mechanisms, whereas this dissertation asks a controlled diagnostic question. It compares transparent fixed formulas, a learned table and a small set of extensions inside one compact architecture. The research contribution is therefore empirical and implementation-aware: it tests necessity, two-dimensional combination, patch-to-position alignment, data regime and capacity trade-offs without presenting a new general-purpose Transformer family.
    """)

    add_numbered_heading(doc, "3", "Methodology")
    doc.add_paragraph("3.1 Controlled Empirical Design", style="Heading 2")
    add_body(doc, """
        The independent variables are positional encoding, patch extraction and assignment convention, and the presence of hybrid or fusion components. All core CIFAR-10 comparisons share the same train/validation split, image transformations, compact ViT dimensions, optimiser settings, maximum epoch budget and model-selection rule. Seeds 42–46 change training stochasticity while the split seed remains 42. This design supports paired seed contrasts because each model is evaluated under matched seed identifiers and the same data partition.
    """)
    add_caption(doc, 1, "Shared CIFAR-10 experimental protocol.")
    add_table(doc, ["Component", "Setting"], [
        ["Dataset split", "45,000 train; 5,000 validation; 10,000 test; split seed 42"],
        ["Input", "32×32 RGB; random crop with padding 4 and horizontal flip for training"],
        ["Normalisation", "CIFAR-10 channel mean (0.4914, 0.4822, 0.4465) and SD (0.2470, 0.2435, 0.2616)"],
        ["Backbone", "Patch 4×4; 64 patch tokens; embedding 128; 4 blocks; 4 heads; MLP width 512"],
        ["Optimiser", "AdamW; learning rate 3×10⁻⁴; weight decay 0.05; batch size 128"],
        ["Schedule", "ReduceLROnPlateau, patience 5, factor 0.5, minimum 10⁻⁶"],
        ["Stopping", "At most 100 epochs; validation-accuracy early stopping, patience 10, minimum delta 0.001"],
        ["Evaluation", "Validation-selected checkpoint evaluated once on the holdout test set"],
        ["Seeds", "Training seeds 42, 43, 44, 45 and 46; fixed split seed 42"],
    ], [1900, 7000])

    doc.add_paragraph("3.2 Dataset and Preprocessing", style="Heading 2")
    add_body(doc, """
        CIFAR-10 contains 50,000 training images and 10,000 test images in ten balanced classes [13]. Ten per cent of the original training partition is reserved for validation using a deterministic split. Training images receive random 32×32 crops after four-pixel padding and random horizontal flips. Validation and test images receive only tensor conversion and normalisation. The test set is excluded from checkpoint selection and is evaluated only after the selected validation checkpoint has been restored.

        The planned second-dataset protocol uses CIFAR-100 unless a fully validated CADB Scene pipeline is selected at the predefined gate. CIFAR-100 retains the 32×32 input and single-label classification interface but changes the classifier to 100 outputs. Dataset-specific normalisation is applied. Partial seed sets are not mixed into the main results: the cross-dataset table is populated only after all four prespecified models complete seeds 42–46.
    """)
    doc.add_paragraph("3.3 Compact Vision Transformer", style="Heading 2")
    add_body(doc, """
        Each 32×32 image is partitioned into an 8×8 grid of non-overlapping 4×4 patches. A convolution with kernel and stride four projects each patch to 128 dimensions. A learned classification token is prepended, after which positional information is added to the 64 patch tokens. Four pre-normalisation Transformer blocks use four attention heads and a feed-forward hidden width of 512. The classification head maps the final class-token representation to ten logits. The no-PE and fixed-PE models contain 801,034 parameters, while learned absolute PE adds one vector per token and contains 809,354 parameters.
    """)
    doc.add_paragraph("3.4 Positional Encoding Families", style="Heading 2")
    add_body(doc, """
        Let S(p) be the standard d-dimensional sinusoidal encoding of scalar position p, with alternating sine and cosine channels and wavelengths determined by 10000^(2i/d). For a patch at row r and column c, the tested fixed encodings are S(r), S(c), S(r)+S(c), S(r)⊙S(c), and S(√(r²+c²)), where ⊙ denotes element-wise multiplication. Shifted variants compute the column component with an offset wavelength index before addition or multiplication. Squared variants apply an element-wise square to the multiplicative vector. These are hand-designed absolute signals, not relative attention mechanisms.
    """)
    add_caption(doc, 2, "Positional-encoding constructions and their roles.")
    add_table(doc, ["Family", "Construction", "Trainable positional parameters", "Role"], [
        ["No PE", "0", "None", "Necessity baseline"],
        ["Learned absolute", "One learned vector per sequence position", "Yes", "Flexible reference"],
        ["Row / column", "S(r) or S(c)", "None", "Single-axis ablation"],
        ["Additive", "S(r) + S(c)", "None", "Symmetric axis combination"],
        ["Multiplicative", "S(r) ⊙ S(c)", "None", "Element-wise axis interaction"],
        ["Shifted variants", "Offset column frequency indices", "None", "Reduce identical axis frequency alignment"],
        ["Radial", "S(√(r²+c²))", "None", "Distance-from-origin extension"],
        ["Squared", "[S(r) ⊙ S(c)]²", "None", "Magnitude/sign extension"],
    ], [1500, 2700, 1900, 2800])

    doc.add_paragraph("3.5 Patch Sequence and PE Coordinate Assignment", style="Heading 2")
    add_body(doc, """
        Four deterministic conventions are implemented: normal_row, normal_col, proper_row and proper_col. The first two traverse rows or columns monotonically. The proper variants use serpentine traversal, reversing alternate rows or columns. A mapping test constructs a labelled patch grid, passes it through patch extraction, and records physical coordinate → sequence slot → assigned PE coordinate/vector. This test is required before the Results chapter interprets a model difference as an assignment effect.

        The distinction matters because sequence storage and PE assignment may be changed jointly or separately. A joint permutation preserves which positional vector belongs to each physical patch. A reassignment changes that relation. Consequently, statements such as ‘token order destroys neighbourhoods’ are too broad for full self-attention; the measured effect must be attributed to the exact token–position mapping verified by the deterministic test.
    """)
    doc.add_paragraph("3.6 Hybrid and Fusion Extensions", style="Heading 2")
    add_body(doc, """
        The exploratory hybrid adds a fixed multiplicative vector to a learned absolute embedding through a learned scalar: x′ = x + Elearned + αEfixed. The scalar is initialised at zero, allowing training to start from the learned-only representation. The comparison is order-matched against normal_col learned PE. The selected checkpoints record the parameter under the implementation name fixed_pos_scale.

        Five fusion configurations process row- and column-oriented inputs in separate branches and combine their class-token representations by concatenation, averaging, an MLP, bidirectional cross-attention, or cross-attention followed by an MLP head. These models are exploratory because they contain considerably more parameters than a single branch. Their results answer whether added capacity produces a commensurate gain, not whether fusion is intrinsically inferior under a matched budget.
    """)
    doc.add_paragraph("3.7 Training, Selection and Reproducibility", style="Heading 2")
    add_body(doc, """
        Cross-entropy loss is optimised with AdamW. ReduceLROnPlateau lowers the learning rate when validation performance stalls. Early stopping and checkpoint selection use validation accuracy only. The test loader is evaluated once after training for the selected checkpoint, and every accepted summary must contain test_evaluation_protocol = selected_checkpoint_only. Configurations, seeds, split seed, selected epoch, metrics and checkpoint paths are stored in machine-readable summaries. Aggregate tables are regenerated from these summaries rather than edited manually.
    """)
    doc.add_paragraph("3.8 Metrics and Statistical Reporting", style="Heading 2")
    add_body(doc, """
        For each model, test accuracy and test loss are reported as the five-seed mean with a two-sided 95% Student-t confidence interval. With n=5 and four degrees of freedom, the half-width is 2.776 × SD/√5, where SD is the sample standard deviation. Key comparisons use seed-matched differences and the same confidence-interval formula. These intervals describe uncertainty across the five training runs; they do not turn the seeds into independent datasets.

        Formal claims are deliberately limited. An exact two-sided Wilcoxon signed-rank test with five non-zero pairs cannot attain p<0.05; its minimum p-value is 0.0625. Critical-difference diagrams and the multiple-dataset procedures reviewed by Demšar [14] are therefore not used to claim significance across five seeds of one dataset. The discussion focuses on effect magnitude, paired direction, interval width and consistency, and does not interpret failure to reject a null hypothesis as equivalence.
    """)
    doc.add_paragraph("3.9 Low-Data and Cross-Dataset Protocols", style="Heading 2")
    add_body(doc, """
        The low-data experiment compares learned absolute PE with unshifted multiplicative PE at 1,000, 5,000, 10,000 and full training-set sizes. The reduced conditions are rerun with seeds 42–46 under the same split and selection protocol; old single-seed exploratory results are not mixed with them. The principal output is the paired accuracy difference at each data size, accompanied by loss and 95% intervals.

        The second-dataset experiment is limited to no PE, learned absolute PE, shifted additive PE and shifted multiplicative PE. It retains the compact backbone and five-seed protocol. Until every prespecified run is complete, this dissertation treats cross-dataset evidence as pending and does not place partial seeds beside the final CIFAR-10 table.
    """)

    add_numbered_heading(doc, "4", "Experiments and Results")
    doc.add_paragraph("4.1 Reporting Convention", style="Heading 2")
    add_body(doc, """
        Each suite is presented as a question, controlled comparison, numerical result and factual observation. Mechanistic interpretation is deferred to Chapter 5. Accuracy is expressed in percentage points and loss is unitless cross-entropy. The ± value is the 95% confidence-interval half-width, not the sample standard deviation.
    """)
    doc.add_paragraph("4.2 Core Positional-Encoding Comparison", style="Heading 2")
    add_body(doc, """
        This suite asks whether positional information is necessary and how basic fixed two-dimensional encodings compare with a learned absolute table. All nine models share the same row-major patch extraction and compact backbone. Radial PE is included as a prespecified fixed extension, while the confirmatory core remains the no-PE, learned, row, column, additive, shifted additive, multiplicative and shifted multiplicative models.
    """)
    add_caption(doc, 3, "Core CIFAR-10 selected-checkpoint test results across five seeds.")
    core_table = []
    for row in core_rows:
        core_table.append([
            LABELS[row["model"]],
            format_acc(row),
            format_loss(row),
            f"{int(row['parameter_count']):,}",
        ])
    add_table(doc, ["Model", "Accuracy (%) ± 95% CI", "Loss ± 95% CI", "Parameters"], core_table, [3300, 2100, 1900, 1500])
    add_body(doc, f"""
        No PE obtains {format_acc(baseline)}% accuracy, whereas learned absolute PE obtains {format_acc(learned)}%. The seed-matched mean difference is {float(paired['no_pe_to_learnable']['mean_test_acc_delta_pp']):.3f} percentage points, with a 95% interval from {float(paired['no_pe_to_learnable']['ci95_lower_test_acc_delta_pp']):.3f} to {float(paired['no_pe_to_learnable']['ci95_upper_test_acc_delta_pp']):.3f}. Shifted multiplicative PE is the strongest fixed design at {format_acc(shifted_mult)}%. Its paired difference relative to learned PE is {float(paired['learnable_to_shifted_multiplicative']['mean_test_acc_delta_pp']):.3f} points when defined as shifted multiplicative minus learned.

        Among the remaining fixed designs, unshifted multiplicative and the additive variants rank above the single-axis and radial encodings. Radial PE does not alter the main ordering: it exceeds the no-PE baseline but remains below the two-axis additive and multiplicative constructions.
    """)
    doc.add_paragraph("4.3 Shifted Variants", style="Heading 2")
    add_body(doc, f"""
        The shifted designs ask whether offsetting the column-axis wavelength indices changes the additive and multiplicative constructions. Shifted additive changes mean accuracy by {float(paired['additive_to_shifted_additive']['mean_test_acc_delta_pp']):.3f} points relative to additive, with a paired 95% interval from {float(paired['additive_to_shifted_additive']['ci95_lower_test_acc_delta_pp']):.3f} to {float(paired['additive_to_shifted_additive']['ci95_upper_test_acc_delta_pp']):.3f}. Shifted multiplicative changes the mean by {float(paired['multiplicative_to_shifted_multiplicative']['mean_test_acc_delta_pp']):.3f} points, with an interval from {float(paired['multiplicative_to_shifted_multiplicative']['ci95_lower_test_acc_delta_pp']):.3f} to {float(paired['multiplicative_to_shifted_multiplicative']['ci95_upper_test_acc_delta_pp']):.3f}. The factual result is that both means increase, but neither paired interval establishes a clear and consistently substantial improvement across five seeds.
    """)
    doc.add_paragraph("4.4 Patch-to-Position Assignment", style="Heading 2")
    order_models = [
        ("Row-major", "vit_learnable_position", "vit_multiplicative_sinusoidal"),
        ("Column-major", "vit_normal_col_learnable_position", "vit_normal_col_multiplicative_sinusoidal"),
        ("Serpentine rows", "vit_proper_row_learnable_position", "vit_proper_row_multiplicative_sinusoidal"),
        ("Serpentine columns", "vit_proper_col_learnable_position", "vit_proper_col_multiplicative_sinusoidal"),
    ]
    order_rows = []
    for label, learned_name, fixed_name in order_models:
        order_rows.append([label, format_acc(summary[learned_name]), format_acc(summary[fixed_name])])
    add_caption(doc, 4, "Order and assignment comparison for learned and multiplicative PE.")
    add_table(doc, ["Convention", "Learned PE accuracy (%)", "Multiplicative PE accuracy (%)"], order_rows, [3000, 3000, 3000])
    add_body(doc, """
        The deterministic mapping test passes for normal_row, normal_col, proper_row and proper_col and records all 64 physical patch coordinates. Learned PE means remain within a narrow band across the four conventions. Multiplicative PE changes more substantially under the serpentine assignments, particularly for the proper_col condition. The result is recorded as an interaction with the verified patch-to-position assignment; it is not described as evidence that changing token order alone destroys adjacency in full self-attention.
    """)
    doc.add_paragraph("4.5 Low-Data Regime", style="Heading 2")
    add_body(doc, """
        This suite tests whether the ranking between learned and multiplicative PE changes when the training subset is reduced to 1,000, 5,000 or 10,000 examples. The reduced-data runs use the same five seeds and selected-checkpoint protocol as the main experiment. At the time of this structured draft, the protocol and aggregation code are fixed but the complete 30-run result set has not yet passed the completeness check. No partial numerical conclusion is reported here.
    """)
    doc.add_paragraph("4.6 Second-Dataset Generalisation", style="Heading 2")
    add_body(doc, """
        The cross-dataset suite compares no PE, learned PE, shifted additive PE and shifted multiplicative PE using five seeds. CIFAR-100 is the default dataset unless the CADB Scene gate is satisfied. The implementation supports CIFAR-100 and its 100-class output interface. Because a complete five-seed table is not yet available, the current evidence is explicitly pending and is not mixed with the completed CIFAR-10 results.
    """)
    doc.add_paragraph("4.7 Hybrid, Fusion and Squared Extensions", style="Heading 2")
    extension_names = [
        "vit_normal_col_learnable_position",
        "vit_normal_col_learnable_multiplicative_sinusoidal",
        "vit_squared_multiplicative_sinusoidal",
        "vit_squared_multiplicative_sinusoidal_shifted",
        "vit_row_col_latent_fusion",
        "vit_row_col_mean_fusion",
        "vit_row_col_mean_mlp_fusion",
        "vit_row_col_cross_attention_fusion",
        "vit_row_col_cross_attention_mlp_head_fusion",
    ]
    extension_rows = [[LABELS.get(name, name), format_acc(summary[name]), format_loss(summary[name]), f"{int(summary[name]['parameter_count']):,}"] for name in extension_names]
    add_caption(doc, 5, "Exploratory hybrid, squared and fusion results.")
    add_table(doc, ["Configuration", "Accuracy (%) ± 95% CI", "Loss ± 95% CI", "Parameters"], extension_rows, [3300, 2100, 1900, 1500])
    add_body(doc, f"""
        The order-matched hybrid differs from normal_col learned PE by {float(paired['order_matched_learnable_to_hybrid']['mean_test_acc_delta_pp']):.3f} points on average, with a 95% interval from {float(paired['order_matched_learnable_to_hybrid']['ci95_lower_test_acc_delta_pp']):.3f} to {float(paired['order_matched_learnable_to_hybrid']['ci95_upper_test_acc_delta_pp']):.3f}. The learned fixed_pos_scale is small and varies in sign across the five selected checkpoints, ranging from −0.0403 to 0.0492. The squared designs do not exceed shifted multiplicative PE.

        Cross-attention + MLP head is the strongest fusion model at {format_acc(fusion)}% accuracy. Relative to the learned single-branch reference, its paired mean difference is {float(paired['learnable_to_best_fusion']['mean_test_acc_delta_pp']):.3f} points, defined as fusion minus learned. The more complex fusion therefore does not deliver an improvement commensurate with its parameter increase.
    """)

    add_numbered_heading(doc, "5", "Analysis and Discussion")
    doc.add_paragraph("5.1 RQ1: Is Positional Encoding Necessary?", style="Heading 2")
    add_body(doc, """
        The approximately seven-point paired gap between learned PE and no PE is the most robust result in the completed experiment. Patch content contains local visual evidence, but it does not uniquely specify spatial origin. A class token attending globally can aggregate content without PE, which explains why the baseline remains functional, yet the weaker result shows that content statistics do not fully replace a coordinate frame. The conclusion is deliberately local: positional information is necessary for the strongest performance of this compact CIFAR-10 ViT under the tested recipe, not for every possible Transformer classifier.

        Learned absolute PE is the strongest core reference and has a narrow confidence interval. It should therefore remain the default comparator for the remaining analyses. The result does not show that learned PE is universally superior to fixed PE, because the architecture and optimiser were not tuned separately for each encoding and the evidence currently covers one completed dataset.
    """)
    doc.add_paragraph("5.2 RQ2: What Matters in the Fixed Construction?", style="Heading 2")
    add_body(doc, """
        Single-axis PE is weaker than the leading two-axis designs, consistent with the loss of one coordinate dimension. Additive PE retains both axes but combines them linearly. Multiplicative PE introduces channel-wise interactions and performs better on average. The shifted multiplicative variant is the strongest fixed design and unusually stable across seeds. A plausible interpretation is that offset axis frequencies reduce redundant alignment between row and column components, but classification accuracy alone does not identify the mechanism.

        Radial PE compresses (r,c) to distance from the upper-left origin. Distinct directions may therefore share the same radial coordinate. Its result above no PE but below the leading axis-aware designs is compatible with partial spatial information loss. Squaring the multiplicative vector removes sign information and does not improve the best fixed result. These negative results are useful design boundaries: extra nonlinearity or symmetry does not automatically produce a better positional signal.
    """)
    doc.add_paragraph("5.3 RQ3: Assignment, Data Regime and Generalisation", style="Heading 2")
    add_body(doc, """
        Learned absolute embeddings can adapt independently at each sequence slot, which explains their small variation across the tested mappings. Fixed encodings constrain the relation between a physical patch and its coordinate vector, so assignment changes can alter the information available even when the encoder retains global attention. The deterministic mapping test is therefore not merely software QA; it defines the causal interpretation of the ablation.

        The low-data and second-dataset suites are required before making a broader statement about inductive bias. A fixed encoding could be relatively useful when learned positional parameters receive less data, but that is a hypothesis rather than a current finding. Similarly, agreement between CIFAR-10 and CIFAR-100 would strengthen the ranking claim, whereas a reversal would show dataset dependence. Until those runs are complete, the dissertation’s strongest conclusions remain CIFAR-10-specific.
    """)
    doc.add_paragraph("5.4 RQ4: Hybrid and Fusion Complexity", style="Heading 2")
    add_body(doc, """
        The hybrid is best described as an exploratory combination. Its mean is close to the order-matched learned model, its paired interval includes both small gains and losses, and fixed_pos_scale varies around zero. This evidence does not justify the terms optimised PE or significant improvement. It instead suggests that the learned model can tolerate access to the fixed vector without deriving a consistent large benefit from it.

        Fusion provides a clearer capacity lesson. Bidirectional cross-attention is the best of the fusion mechanisms, but it remains below the learned single-branch reference. Because fusion duplicates encoder components, it is not a parameter-matched test of representation quality. The correct claim is that increased row–column branch complexity did not produce an accuracy gain commensurate with the added capacity under this training recipe. It would be invalid to conclude that all cross-attention fusion is ineffective.
    """)
    doc.add_paragraph("5.5 Threats to Validity", style="Heading 2")
    add_body(doc, """
        Internal validity is limited by hyperparameters shared across models: a controlled recipe improves fairness but may be suboptimal for individual PE families. The hybrid and fusion suites contain deliberate architectural changes and are interpreted as exploratory. Five seeds reveal run-to-run variation but provide limited statistical power. Confidence intervals over training seeds describe stochastic repeatability rather than population uncertainty over datasets.

        External validity is limited by one completed low-resolution dataset, one compact architecture and training from scratch. Results may change with image resolution, pretraining, augmentation, relative attention or a larger model. Construct validity is limited because classification accuracy and loss measure downstream success, not spatial information directly. Additional diagnostics such as position-vector similarity, attention displacement or resolution transfer could test proposed mechanisms.

        The repeated development of models creates test-set researcher degrees of freedom even when each run uses a validation-selected checkpoint. This dissertation mitigates the risk by reporting the full prespecified family, retaining negative results, generating tables automatically and avoiding claims based on a single favourable seed. A future confirmatory study should freeze hypotheses and a holdout dataset before additional model iteration.
    """)
    doc.add_paragraph("5.6 Future Work", style="Heading 2")
    add_body(doc, """
        The immediate priorities are to complete the five-seed low-data matrix and the four-model second-dataset evaluation. The first asks whether multiplicative fixed PE becomes relatively stronger with fewer labels; the second tests whether the main ordering transfers beyond CIFAR-10. Further work should parameter-match fusion branches, evaluate higher resolutions, and compare the absolute input-level encodings with relative, conditional and rotary baselines. Mechanistic analysis should be specified before inspecting the results and should connect any claim about geometry to a measurable property of the positional vectors or attention patterns.
    """)

    add_numbered_heading(doc, "6", "Conclusion")
    add_body(doc, f"""
        This dissertation evaluates positional encoding as a controlled design variable in compact Vision Transformers. The completed CIFAR-10 experiment contains 32 configurations, five seeds per configuration and a validation-selected checkpoint protocol. RQ1 is answered most clearly: positional information is important in this setting, with learned absolute PE reaching {float(learned['mean_test_acc_pct']):.3f}% mean test accuracy compared with {float(baseline['mean_test_acc_pct']):.3f}% for no PE. Learned PE remains the strongest core reference.

        For RQ2, the results show that two-dimensional construction matters. Shifted multiplicative PE is the strongest fixed design at {float(shifted_mult['mean_test_acc_pct']):.3f}%, while single-axis, radial and squared variants do not exceed it. The evidence supports a promising fixed alternative, not a universal claim that fixed PE outperforms learned PE. For RQ3, deterministic mapping tests and the order suite show that learned embeddings are comparatively robust to the tested assignments, whereas fixed multiplicative PE is more sensitive to the patch-to-position relation. The low-data and second-dataset components remain necessary to establish generality.

        For RQ4, neither additional mechanism produces a compelling gain. The hybrid differs only slightly from its order-matched learned comparator, and the learned fixed scale varies around zero. The best cross-attention fusion remains below learned PE despite greater capacity. The overall contribution is therefore a reproducible empirical map of positional choices, their interactions and their limitations. Within its stated scope, learned absolute PE is a strong default and shifted multiplicative PE is the most competitive fixed construction; broader recommendations await data-regime and cross-dataset evidence.
    """)

    doc.add_page_break()
    doc.add_paragraph("References", style="Heading 1")
    for index, reference in enumerate(REFERENCES, 1):
        paragraph = doc.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Cm(0.9)
        paragraph.paragraph_format.first_line_indent = Cm(-0.9)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.add_run(f"[{index}] {reference}")

    doc.core_properties.title = "Positional Encoding in Compact Vision Transformers: A Controlled Evaluation"
    doc.core_properties.author = "Yikai Zhao"
    doc.core_properties.subject = "UCL MSc Dissertation"
    doc.core_properties.comments = "Fixed working draft. No figures or figure placeholders."

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            black_arial(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        black_arial(run)
    return doc


def main() -> None:
    for required in (SUMMARY_CSV, CORE_CSV, PAIRED_CSV):
        if not required.exists():
            raise FileNotFoundError(required)
    THESIS_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(TEMP_OUTPUT)
    try:
        os.replace(TEMP_OUTPUT, OUTPUT)
    except PermissionError as error:
        raise RuntimeError(
            f"Cannot update {OUTPUT} because it appears to be open in Word. "
            "Close the document and run this script again."
        ) from error
    print(OUTPUT)


if __name__ == "__main__":
    main()
