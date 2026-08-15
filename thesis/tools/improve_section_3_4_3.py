from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document

from standardize_math_notation import (
    add_body_text,
    add_ref_field,
    append_inline_math,
    clear_paragraph_content,
    fraction,
    math_run,
    subscript,
)


ROOT = Path(__file__).resolve().parents[2]
DOCX = ROOT / "thesis" / "Yikai_Zhao_MSc_Dissertation.docx"
BUILDING = ROOT / "thesis" / ".Yikai_Zhao_MSc_Dissertation.3_4_3.tmp.docx"


def field_counts(document: Document) -> dict[str, int]:
    codes = [
        (element.text or "").strip()
        for element in document.element.body.xpath(".//w:instrText")
    ]
    return {
        "zotero": sum(code.startswith("ADDIN ZOTERO") for code in codes),
        "seq": sum(code.startswith("SEQ ") for code in codes),
        "ref": sum(code.startswith("REF ") for code in codes),
    }


def find_section_paragraphs(document: Document):
    paragraphs = document.paragraphs
    headings = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip()
        == "3.4.3 Additive, Multiplicative and Shifted Encodings"
    ]
    if len(headings) != 1:
        raise RuntimeError(f"Expected one Section 3.4.3 heading, found {len(headings)}")

    heading = headings[0]
    if paragraphs[heading + 7].text.strip() != "3.4.4 Squared and Radial Extensions":
        raise RuntimeError("Unexpected Section 3.4.3 paragraph structure")
    return paragraphs[heading + 1], paragraphs[heading + 3], paragraphs[heading + 6]


def rewrite_combination_paragraph(paragraph) -> None:
    clear_paragraph_content(paragraph)
    add_body_text(
        paragraph,
        "The additive and multiplicative encodings combine the row and column "
        "vectors into a single positional vector without changing the "
        "128-dimensional embedding size. The additive variant sums corresponding "
        "components, while the multiplicative variant takes their element-wise "
        "product, as defined in Equation (",
    )
    add_ref_field(paragraph, "eq7", "7")
    add_body_text(paragraph, "). Here, ")
    append_inline_math(paragraph, math_run("⊙"))
    add_body_text(paragraph, " denotes element-wise multiplication.")


def rewrite_shift_paragraph(paragraph) -> None:
    clear_paragraph_content(paragraph)
    add_body_text(
        paragraph,
        "The shifted variants retain these two combination rules but assign "
        "different sinusoidal frequency schedules to the row and column axes. "
        "For sine–cosine pair ",
    )
    append_inline_math(paragraph, math_run("i"))
    add_body_text(paragraph, ", the row schedule uses exponent ")
    append_inline_math(paragraph, math_run("−"), fraction("2i", "d"))
    add_body_text(paragraph, ", whereas the column schedule uses ")
    append_inline_math(paragraph, math_run("−"), fraction("2i+1", "d"))
    add_body_text(paragraph, ", as defined in Equation (")
    add_ref_field(paragraph, "eq8", "8")
    add_body_text(
        paragraph,
        "). Within each axis, the sine and cosine channels of a pair share that "
        "frequency. Equation (",
    )
    add_ref_field(paragraph, "eq9", "9")
    add_body_text(paragraph, ") then combines the axis-specific mappings ")
    append_inline_math(paragraph, subscript("S", "r"), math_run("(r)"))
    add_body_text(paragraph, " and ")
    append_inline_math(paragraph, subscript("S", "c"), math_run("(c)"))
    add_body_text(paragraph, " by addition or element-wise multiplication.")


def rewrite_scope_paragraph(paragraph) -> None:
    clear_paragraph_content(paragraph)
    add_body_text(
        paragraph,
        "The term shifted therefore refers only to the one-index offset between "
        "the row and column frequency schedules. It leaves the patch coordinates, "
        "token order and patch-to-position assignment unchanged; those factors are "
        "examined separately in Section 3.5.",
    )


def validate(document: Document, expected_fields: dict[str, int]) -> None:
    first, second, third = find_section_paragraphs(document)
    if not first.text.startswith(
        "The additive and multiplicative encodings combine the row and column vectors"
    ):
        raise RuntimeError("The first revised paragraph was not found")
    if not second.text.startswith(
        "The shifted variants retain these two combination rules"
    ):
        raise RuntimeError("The second revised paragraph was not found")
    if not third.text.startswith("The term shifted therefore refers only"):
        raise RuntimeError("The scope clarification was not found")

    if len(second._p.xpath(".//m:f")) != 2:
        raise RuntimeError("Shifted exponents must remain stacked fractions")
    if "/" in second.text:
        raise RuntimeError("A slash-form fraction remains in Section 3.4.3")

    section_refs = [
        (element.text or "").strip()
        for paragraph in (first, second, third)
        for element in paragraph._p.xpath(".//w:instrText")
    ]
    if section_refs != ["REF eq7 \\h", "REF eq8 \\h", "REF eq9 \\h"]:
        raise RuntimeError(f"Unexpected Section 3.4.3 references: {section_refs}")

    for bookmark in ("eq7", "eq8", "eq9"):
        matches = [
            element
            for element in document.element.body.xpath(".//w:bookmarkStart")
            if element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name")
            == bookmark
        ]
        if len(matches) != 1:
            raise RuntimeError(f"Equation bookmark {bookmark} was changed")

    if field_counts(document) != expected_fields:
        raise RuntimeError("Document field counts changed")
    if len(document.inline_shapes) != 10 or len(document.tables) != 11:
        raise RuntimeError("Figure or table counts changed")


def main() -> None:
    document = Document(str(DOCX))
    expected_fields = field_counts(document)
    first, second, third = find_section_paragraphs(document)

    rewrite_combination_paragraph(first)
    rewrite_shift_paragraph(second)
    rewrite_scope_paragraph(third)
    validate(document, expected_fields)

    document.save(str(BUILDING))
    reopened = Document(str(BUILDING))
    validate(reopened, expected_fields)
    BUILDING.replace(DOCX)

    with ZipFile(DOCX) as archive:
        if archive.testzip() is not None:
            raise RuntimeError("The saved DOCX package failed its ZIP integrity test")

    print(f"Updated: {DOCX}")
    print("Section 3.4.3 paragraphs revised: 3")
    print(f"Fields preserved: {expected_fields}")


if __name__ == "__main__":
    main()
